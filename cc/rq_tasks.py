import base64
import datetime
import io
import json
import shutil
import threading
import uuid
from datetime import time, timedelta
import csv
import ffmpeg
import pandas as pd
import webvtt
from PIL import Image
from asgiref.sync import async_to_sync
from bs4 import BeautifulSoup
from channels.layers import get_channel_layer
from django.contrib.auth.models import User
from django.core.files import File
from django.core.management import call_command
from django.core.signing import TimestampSigner
from django.db import transaction
from django.db.models import Q, QuerySet
from django.utils import timezone
from django_rq import job
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re
import logging
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from openpyxl.reader.excel import load_workbook
from openpyxl.styles import PatternFill, Border, Side, Alignment
from openpyxl.utils import get_column_letter
from openpyxl.workbook import Workbook
from openpyxl.worksheet.datavalidation import DataValidation
from pytesseract import pytesseract
from rest_framework.exceptions import ValidationError
from sdrf_pipelines.sdrf.sdrf import SdrfDataFrame
import time
from pathlib import Path
from cc.models import Annotation, ProtocolModel, ProtocolStep, StepVariation, ProtocolSection, Session, \
    AnnotationFolder, Reagent, ProtocolReagent, StepReagent, ProtocolTag, StepTag, Tag, Project, MetadataColumn, \
    InstrumentJob, SubcellularLocation, Species, MSUniqueVocabularies, Unimod, FavouriteMetadataOption, InstrumentUsage, \
    LabGroup, Tissue, StorageObject, ReagentAction, StoredReagent, Instrument, SiteSettings, SamplePool
from django.conf import settings
import numpy as np
import subprocess
import os
import select
from cc.serializers import AnnotationSerializer, ProtocolModelSerializer, ProtocolSectionSerializer, \
    StepVariationSerializer, ProtocolStepSerializer, SessionSerializer, AnnotationFolderSerializer, \
    ReagentSerializer, ProtocolReagentSerializer, StepReagentSerializer, TagSerializer, ProtocolTagSerializer, \
    StepTagSerializer, ProjectSerializer
import docx
import re

from docx.shared import Inches, RGBColor
import re

from cc.improved_docx_generator import EnhancedDocxGenerator, DocxGenerationError
from cc.models import SiteSettings
from cc.utils import user_metadata, staff_metadata, required_metadata_name, identify_barcode_format
from cc.utils.user_data_export_revised import export_user_data_revised, export_protocol_data, export_session_data
from cc.utils.user_data_import_revised import dry_run_import_user_data, import_user_data_revised
from mcp_server.tools.protocol_analyzer import ProtocolAnalyzer
from cc.models import ProtocolStep, ProtocolStepSuggestionCache
from mcp_server.tools.sdrf_generator import SDRFMetadataGenerator

capture_language = re.compile(r"auto-detected language: (\w+)")

@job('transcribe', timeout='1h')
def transcribe_audio(audio_path: str, model_path: str, step_annotation_id: int, language: str = "auto", translate: bool = False, custom_id: str = None):
    """
    Convert audio from webm to wav using ffmpeg, then store the wave file as temporary file and transcribe it using the whisper model using subprocess and whispercpp main binary and base.en model before deleting the temporary file
    :param audio_path:
    :param model_path:
    :param step_annotation_id:
    :return:
    """
    if audio_path.endswith(".webm"):
        wav_path = audio_path.replace(".webm", ".wav")
    elif audio_path.endswith(".m4a"):
        wav_path = audio_path.replace(".m4a", ".wav")
    else:
        wav_path = audio_path + ".wav"
    #ffmpeg.input(audio_path).output(wav_path, format="s16le", acodec="pcm_s16le", ac=1, ar=44100).run(cmd=['ffmpeg', '-nostdin'], capture_stdout=True, capture_stderr=True)
    subprocess.run(["ffmpeg", "-y", "-i", audio_path, "-vn", "-ar", "16000", wav_path])
    # Transcribe audio using whisper model
    whispercpp_bin_path = settings.WHISPERCPP_PATH
    temporary_vtt_path = wav_path + ".vtt"
    thread_count = settings.WHISPERCPP_THREAD_COUNT
    # Run whispercpp binary
    print(f"Running whispercpp binary: {whispercpp_bin_path} -m {model_path} -f {wav_path} -ovtt -t {thread_count}")
    cmd = ['stdbuf', '-oL', whispercpp_bin_path, "-m", model_path, "-f", wav_path, "-ovtt", '-t', thread_count, '-l', language]
    #if translate:
    #    cmd.append("-tr")
    #process = subprocess.run(cmd)
    process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    # capture whisper detected language

    annotation = Annotation.objects.get(id=step_annotation_id)
    for line in process.stdout:
        line = line.decode("utf-8")
        print(line)
        if "auto-detected language" in line:
            s = capture_language.search(line)
            if s:
                annotation.language = s.group(1)
    if not annotation.language:
        for line in process.stderr:
            line = line.decode("utf-8")
            print(line)
            if "auto-detected language" in line:
                s = capture_language.search(line)
                if s:
                    annotation.language = s.group(1)

    #for line in process.stderr.decode("utf-8").split("\n"):
    #    print(f"stderr: {line}")
    with open(temporary_vtt_path, "rt") as f:
        annotation.transcription = f.read()
        annotation.transcribed = True
        annotation.save()

    if translate:
        if annotation.language and annotation.language != "en":
            print("Translating from" + annotation.language + " to en")
            cmd.append("-tr")
            process = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            with open(temporary_vtt_path, "rt") as f:
                annotation.translation = f.read()
    annotation.save()

    os.remove(wav_path)
    os.remove(temporary_vtt_path)
    session_id = annotation.session.unique_id
    channel_layer = get_channel_layer()
    data = AnnotationSerializer(annotation, many=False).data

    async_to_sync(channel_layer.group_send)(
        f"transcription_{session_id}",
        {
            "type": "transcription_message",
            "message": data,
        },
    )
    return annotation.transcription





@job('transcribe', timeout='1h')
def transcribe_audio_from_video(video_path: str, model_path: str, step_annotation_id: int, language: str = "auto", translate: bool = False, custom_id: str = None):
    """
    Convert audio from webm video to wav using ffmpeg, then store the wave file as temporary file and transcribe it using the whisper model using subprocess and whispercpp main binary and base.en model before deleting the temporary file
    :param video_path:
    :param model_path:
    :param step_annotation_id:
    :return:
    """

    # Convert audio from webm video to wav using ffmpeg specify the kHz to 16 kHz
    if video_path.endswith(".webm"):
        wav_path = video_path.replace(".webm", ".wav")
    elif video_path.endswith(".mp4"):
        wav_path = video_path.replace(".mp4", ".wav")
    else:
        wav_path = video_path + ".wav"
    #ffmpeg.input(video_path).output(wav_path, format="s16le", acodec="pcm_s16le", ac=1, ar=44100).run(cmd=['ffmpeg', '-nostdin'], capture_stdout=True, capture_stderr=True)
    #subprocess.run(["ffmpeg", "-i", video_path, "-vn", "-ar", "44100", wav_path])
    subprocess.run(["ffmpeg", "-y", "-i", video_path, "-vn", "-ar", "16000", wav_path])
    # Transcribe audio using whisper model
    whispercpp_bin_path = settings.WHISPERCPP_PATH
    temporary_vtt_path = wav_path+".vtt"
    thread_count = settings.WHISPERCPP_THREAD_COUNT
    # Run whispercpp binary
    print(f"Running whispercpp binary: {whispercpp_bin_path} -m {model_path} -f {wav_path} -ovtt -t {thread_count}")
    cmd = [whispercpp_bin_path, "-m", model_path, "-f", wav_path, "-ovtt", '-t', thread_count, '-l', language]
    process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    # capture whisper detected language

    annotation = Annotation.objects.get(id=step_annotation_id)
    for line in process.stdout:
        line = line.decode("utf-8")
        print(line)
        if "auto-detected language" in line:
            s = capture_language.search(line)
            if s:
                annotation.language = s.group(1)
    if not annotation.language:
        for line in process.stderr:
            line = line.decode("utf-8")
            print(line)
            if "auto-detected language" in line:
                s = capture_language.search(line)
                if s:
                    annotation.language = s.group(1)

    with open(temporary_vtt_path, "rt") as f:
        annotation.transcription = f.read()
        annotation.transcribed = True

    if translate:
        if annotation.language and annotation.language != "en":
            print("Translating from" + annotation.language + " to en")
            cmd.append("-tr")
            process = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            with open(temporary_vtt_path, "rt") as f:
                annotation.translation = f.read()

    annotation.save()

    os.remove(wav_path)
    os.remove(temporary_vtt_path)
    session_id = annotation.session.unique_id
    channel_layer = get_channel_layer()
    data = AnnotationSerializer(annotation, many=False).data
    async_to_sync(channel_layer.group_send)(
        f"transcription_{session_id}",
        {
            "type": "transcription_message",
            "message": data,
        },
    )
    return annotation.transcription

@job('export', timeout='2h')
def create_docx(protocol_id: int, session_id: str = None, user_id: int = None, instance_id: str = None):
    """
    Enhanced DOCX creation using the improved EnhancedDocxGenerator
    
    Args:
        protocol_id: ID of the protocol to export
        session_id: Optional session ID for session-specific annotations
        user_id: User ID for notifications
        instance_id: Instance ID for tracking
        
    Returns:
        str: Path to generated DOCX file
    """
    logger = logging.getLogger(__name__)
    
    try:
        # Get protocol with error handling
        try:
            protocol = ProtocolModel.objects.get(id=protocol_id)
        except ProtocolModel.DoesNotExist:
            error_msg = f"Protocol {protocol_id} not found"
            logger.error(error_msg)
            if user_id:
                _notify_user_error(user_id, instance_id, error_msg)
            raise ValueError(error_msg)
        
        # Get session if provided
        session = None
        if session_id:
            try:
                session = Session.objects.get(unique_id=session_id)
                logger.info(f"Using session: {session.unique_id}")
            except Session.DoesNotExist:
                logger.warning(f"Session {session_id} not found, proceeding without session")
        
        logger.info(f"Starting enhanced DOCX generation for protocol {protocol.id}: {protocol.protocol_title}")
        
        # Use the enhanced DOCX generator
        try:
            generator = EnhancedDocxGenerator(protocol, session, user_id, instance_id)
            docx_filepath = generator.generate_document()
            
            logger.info(f"Successfully created enhanced DOCX file: {docx_filepath}")
            
            # Get generation statistics
            stats = generator.get_statistics()
            logger.info(f"Document generation stats: {stats}")
            
            # Notify user of completion with enhanced information
            signer = TimestampSigner()
            filename = os.path.basename(docx_filepath)
            signed_value = signer.sign(filename)
            
            if user_id:
                channel_layer = get_channel_layer()
                async_to_sync(channel_layer.group_send)(
                    f"user_{user_id}",
                    {
                        "type": "download_message",
                        "message": {
                            "signed_value": signed_value,
                            "user_download": True,
                            "instance_id": instance_id,
                            "filename": filename,
                            "stats": stats,
                            "enhanced": True
                        },
                    },
                )
            
            # Schedule cleanup (20 minutes)
            threading.Timer(60*20, remove_file, args=[docx_filepath]).start()
            
            return docx_filepath
            
        except DocxGenerationError as e:
            # Handle specific DOCX generation errors
            error_msg = f"Enhanced DOCX generation failed: {str(e)}"
            logger.error(error_msg, exc_info=True)
            
            if user_id:
                _notify_user_error(user_id, instance_id, error_msg)
            
            # Fallback to basic generation if enhanced fails
            logger.info("Attempting fallback to basic DOCX generation...")
            return _create_docx_fallback(protocol, session, user_id, instance_id)
            
    except Exception as e:
        error_msg = f"DOCX generation failed: {str(e)}"
        logger.error(error_msg, exc_info=True)
        
        if user_id:
            _notify_user_error(user_id, instance_id, error_msg)
        
        raise


def _create_docx_fallback(protocol, session, user_id, instance_id):
    """
    Fallback DOCX generation using basic approach
    """
    logger = logging.getLogger(__name__)
    
    try:
        logger.info(f"Using fallback DOCX generation for protocol {protocol.id}")
        
        # Basic document creation
        doc = docx.Document()
        
        # Simple title
        doc.add_heading(protocol.protocol_title, level=0)
        
        # Add basic protocol information
        if hasattr(protocol, 'protocol_description') and protocol.protocol_description:
            doc.add_heading('Description', level=1)
            doc.add_paragraph(protocol.protocol_description)
        
        # Add steps in simple format
        steps = protocol.get_step_in_order() if hasattr(protocol, 'get_step_in_order') else []
        if steps:
            doc.add_heading('Steps', level=1)
            for i, step in enumerate(steps, 1):
                doc.add_heading(f'Step {i}', level=2)
                if hasattr(step, 'step_description'):
                    doc.add_paragraph(step.step_description)
        
        # Generate filename and save
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        protocol_name_clean = re.sub(r'[^\w\-_.]', '_', protocol.protocol_title)[:50]
        filename = f"fallback_{protocol_name_clean}_{timestamp}.docx"
        
        # Ensure temp directory exists
        temp_dir = os.path.join(settings.MEDIA_ROOT, "temp")
        os.makedirs(temp_dir, exist_ok=True)
        
        docx_filepath = os.path.join(temp_dir, filename)
        doc.save(docx_filepath)
        
        logger.info(f"Fallback DOCX created successfully: {docx_filepath}")
        return docx_filepath
        
    except Exception as e:
        logger.error(f"Fallback DOCX generation also failed: {str(e)}")
        raise


def _setup_document_styles(doc):
    """Setup custom styles for better document formatting"""
    try:
        styles = doc.styles
        
        # Protocol title style
        if 'Protocol Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Protocol Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Arial'
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Step header style
        if 'Step Header' not in [s.name for s in styles]:
            step_style = styles.add_style('Step Header', WD_STYLE_TYPE.PARAGRAPH)
            step_style.font.name = 'Arial'
            step_style.font.size = Pt(12)
            step_style.font.bold = True
            step_style.font.color.rgb = RGBColor(51, 51, 51)
            step_style.paragraph_format.space_before = Pt(6)
            step_style.paragraph_format.space_after = Pt(3)
            
    except Exception as e:
        # If style creation fails, continue without custom styles
        pass


def _add_document_header(doc, protocol):
    """Add improved document header"""
    try:
        # Protocol title
        title_para = doc.add_paragraph()
        title_para.style = 'Protocol Title' if 'Protocol Title' in [s.name for s in doc.styles] else 'Title'
        title_para.add_run(remove_html_tags(getattr(protocol, 'protocol_name', 'Untitled Protocol')))
        
        # Protocol ID if available
        if protocol.protocol_id:
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(f"Protocol ID: {protocol.protocol_id}")
            subtitle_run.font.size = Pt(11)
            subtitle_run.italic = True
        
        # Generation timestamp
        timestamp_para = doc.add_paragraph()
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp_run = timestamp_para.add_run(
            f"Generated: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}"
        )
        timestamp_run.font.size = Pt(9)
        timestamp_run.font.color.rgb = RGBColor(102, 102, 102)
        
        doc.add_page_break()
        
    except Exception as e:
        # Fallback to simple header
        doc.add_heading(remove_html_tags(getattr(protocol, 'protocol_name', 'Protocol')), level=1)


def _add_protocol_metadata(doc, protocol):
    """Add protocol metadata section"""
    try:
        doc.add_heading('Protocol Information', level=1)
        
        # Create metadata table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'
        
        metadata_items = [
            ('Authors', protocol.protocol_authors),
            ('Created', protocol.protocol_created_on.strftime('%B %d, %Y') if protocol.protocol_created_on else None),
            ('Last Modified', protocol.protocol_modified_on.strftime('%B %d, %Y') if protocol.protocol_modified_on else None),
            ('Version', str(protocol.protocol_version) if protocol.protocol_version else None),
            ('DOI', protocol.protocol_doi),
            ('URL', protocol.protocol_url)
        ]
        
        for label, value in metadata_items:
            if value:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = str(value)
                row_cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_page_break()
        
    except Exception as e:
        # If metadata table fails, just add a simple section
        doc.add_paragraph("Protocol Information")


def _add_protocol_step(doc, step, step_number, session, logger):
    """Add a single protocol step with improved formatting"""
    try:
        # Step divider
        divider_para = doc.add_paragraph('─' * 80)
        divider_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Step header
        step_header_para = doc.add_paragraph()
        step_header_para.style = 'Step Header' if 'Step Header' in [s.name for s in doc.styles] else 'Heading 2'
        step_header_para.add_run(f"Step {step_number}: {remove_html_tags(step.step_title or 'Untitled Step')}")
        
        # Process step description with reagent substitutions
        description = step.step_description or ""
        
        # Handle reagent substitutions
        try:
            step_reagents = StepReagent.objects.filter(step=step)
            
            for reagent in step_reagents:
                replacements = {
                    f"{reagent.id}.name": reagent.reagent.reagent_name,
                    f"{reagent.id}.quantity": str(reagent.quantity_required or ''),
                    f"{reagent.id}.unit": reagent.unit or '',
                    f"{reagent.id}.scaled_quantity": str((reagent.quantity_required or 0) * (reagent.scalable_factor or 1))
                }
                
                for placeholder, replacement in replacements.items():
                    description = description.replace(placeholder, replacement)
                    
        except Exception as e:
            logger.warning(f"Failed to process reagent substitutions for step {step.id}: {e}")
        
        # Add processed description
        html_to_docx(description, doc)
        
        # Add step duration
        if hasattr(step, 'step_duration') and step.step_duration:
            duration_para = doc.add_paragraph(f"Duration: {convert_seconds_to_time(step.step_duration)}")
            duration_para.runs[0].italic = True
        
        # Add step annotations
        _add_step_annotations(doc, step, session, logger)
        
        # Add safety information if available
        if hasattr(step, 'step_safety_information') and step.step_safety_information:
            safety_para = doc.add_paragraph()
            safety_run = safety_para.add_run('⚠️ Safety Information: ')
            safety_run.bold = True
            safety_run.font.color.rgb = RGBColor(255, 0, 0)
            html_to_docx(step.step_safety_information, doc)
        
        doc.add_paragraph()  # Add spacing
        
    except Exception as e:
        logger.error(f"Failed to add step {step.id}: {e}")
        # Add basic step info as fallback
        doc.add_paragraph(f"Step {step_number}: {step.step_title or 'Step'}")


def _add_step_annotations(doc, step, session, logger):
    """Add step annotations with improved error handling"""
    try:
        annotations = step.annotations.all()
        if session:
            annotations = annotations.filter(session=session)
        
        if not annotations.exists():
            return
        
        # Annotations header
        annotations_para = doc.add_paragraph()
        annotations_para.add_run('Annotations:').bold = True
        
        for annotation in annotations:
            try:
                _add_single_annotation(doc, annotation, logger)
            except Exception as e:
                logger.error(f"Failed to add annotation {annotation.id}: {e}")
                # Add basic annotation info as fallback
                doc.add_paragraph(f"Annotation ({annotation.annotation_type}): {annotation.annotation or 'No content'}")
                
    except Exception as e:
        logger.error(f"Failed to add annotations for step {step.id}: {e}")


def _add_single_annotation(doc, annotation, logger):
    """Add a single annotation with proper error handling"""
    try:
        # Annotation header
        annotation_para = doc.add_paragraph()
        type_run = annotation_para.add_run(f"[{annotation.annotation_type.upper()}] ")
        type_run.bold = True
        type_run.font.color.rgb = RGBColor(0, 102, 204)
        
        if annotation.annotation_name:
            annotation_para.add_run(f"{annotation.annotation_name}: ")
        
        # Handle different annotation types with better error handling
        if annotation.annotation_type == "image":
            _add_image_annotation_safe(doc, annotation, logger)
        elif annotation.annotation_type == "sketch":
            _add_sketch_annotation_safe(doc, annotation, logger)
        elif annotation.annotation_type == "table":
            _add_table_annotation_safe(doc, annotation, logger)
        elif annotation.annotation_type == "checklist":
            _add_checklist_annotation_safe(doc, annotation, logger)
        elif annotation.annotation_type == "alignment":
            _add_alignment_annotation_safe(doc, annotation, logger)
        else:
            # Generic annotation
            if annotation.annotation:
                doc.add_paragraph(annotation.annotation)
        
        # Add transcription, translation, summary
        if annotation.transcribed and annotation.transcription:
            _add_transcription_safe(doc, annotation.transcription, logger)
        
        if annotation.translation:
            _add_translation_safe(doc, annotation.translation, logger)
        
        if annotation.summary:
            summary_para = doc.add_paragraph()
            summary_para.add_run('Summary: ').bold = True
            summary_para.add_run(annotation.summary)
            
    except Exception as e:
        logger.error(f"Failed to process annotation {annotation.id}: {e}")


def _add_image_annotation_safe(doc, annotation, logger):
    """Safely add image annotation"""
    try:
        if not annotation.file or not annotation.file.path:
            return
        
        if not os.path.exists(annotation.file.path):
            logger.warning(f"Image file not found: {annotation.file.path}")
            doc.add_paragraph("Image file not available")
            return
        
        # Get image dimensions safely
        try:
            image_dimensions = ffmpeg.probe(annotation.file.path, show_entries="stream=width,height")
            width = int(image_dimensions["streams"][0]["width"])
            height = int(image_dimensions["streams"][0]["height"])
            
            # Add image with proper sizing
            if width > height:
                doc.add_picture(annotation.file.path, width=Inches(5))
            else:
                doc.add_picture(annotation.file.path, height=Inches(6))
                
        except Exception as e:
            logger.warning(f"Failed to get image dimensions, using default size: {e}")
            doc.add_picture(annotation.file.path, width=Inches(4))
            
    except Exception as e:
        logger.error(f"Failed to add image: {e}")
        doc.add_paragraph("Failed to load image")


def _add_sketch_annotation_safe(doc, annotation, logger):
    """Safely add sketch annotation"""
    try:
        if not annotation.file:
            return
        
        with open(annotation.file.path, 'r') as f:
            load_json = json.load(f)
        
        if "png" in load_json:
            data = load_json["png"].split('base64,')
            if len(data) > 1:
                image_bytes = base64.b64decode(data[1])
                image_file = io.BytesIO(image_bytes)
                
                # Calculate size
                pixel_width = load_json.get("width", 400)
                pixel_height = load_json.get("height", 300)
                
                if pixel_width > pixel_height:
                    doc.add_picture(image_file, width=Inches(4))
                else:
                    doc.add_picture(image_file, height=Inches(5))
                    
    except Exception as e:
        logger.error(f"Failed to add sketch: {e}")
        doc.add_paragraph("Failed to load sketch")


def _add_table_annotation_safe(doc, annotation, logger):
    """Safely add table annotation"""
    try:
        data = json.loads(annotation.annotation)
        if not data:
            return
        
        # Add table name
        if "name" in data:
            doc.add_paragraph(data["name"]).runs[0].bold = True
        
        # Create table
        rows = data.get("nRow", 0)
        cols = data.get("nCol", 0)
        
        if rows > 0 and cols > 0:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            content = data.get("content", [])
            tracking_map = data.get("trackingMap", {})
            
            for n, row in enumerate(content):
                if n >= rows:
                    break
                row_cells = table.rows[n].cells
                for nc, c in enumerate(row):
                    if nc >= cols:
                        break
                    row_cells[nc].text = str(c)
                    
                    # Apply highlighting
                    if f"{n},{nc}" in tracking_map and tracking_map[f"{n},{nc}"]:
                        try:
                            shading_elm = parse_xml(r'<w:shd {} w:fill="B3D9FF"/>'.format(nsdecls('w')))
                            row_cells[nc]._tc.get_or_add_tcPr().append(shading_elm)
                        except:
                            pass  # Skip highlighting if it fails
                            
    except Exception as e:
        logger.error(f"Failed to add table: {e}")
        doc.add_paragraph("Failed to load table")


def _add_checklist_annotation_safe(doc, annotation, logger):
    """Safely add checklist annotation"""
    try:
        data = json.loads(annotation.annotation)
        if not data:
            return
        
        # Add checklist name
        if "name" in data:
            doc.add_paragraph(data["name"]).runs[0].bold = True
        
        # Add checklist items
        for n, c in enumerate(data.get("checkList", []), 1):
            content = c.get('content', f'Item {n}') if isinstance(c, dict) else str(c)
            checked = c.get('checked', False) if isinstance(c, dict) else False
            
            checkbox = "☑️" if checked else "☐"
            doc.add_paragraph(f"{checkbox} {content}")
            
    except Exception as e:
        logger.error(f"Failed to add checklist: {e}")
        doc.add_paragraph("Failed to load checklist")


def _add_alignment_annotation_safe(doc, annotation, logger):
    """Safely add alignment annotation"""
    try:
        data = json.loads(annotation.annotation)
        if not data:
            return
        
        # Add main image
        if "dataURL" in data:
            main = data["dataURL"].split('base64,')
            if len(main) > 1:
                image_bytes = base64.b64decode(main[1])
                image_file = io.BytesIO(image_bytes)
                doc.add_picture(image_file, width=Inches(5))
        
        # Add extracted segments
        for extracted in data.get("extractedSegments", []):
            if "dataURL" in extracted:
                doc.add_paragraph(f"Segment: {extracted.get('start', 'N/A')}-{extracted.get('end', 'N/A')}")
                data_parts = extracted["dataURL"].split('base64,')
                if len(data_parts) > 1:
                    image_bytes = base64.b64decode(data_parts[1])
                    image_file = io.BytesIO(image_bytes)
                    doc.add_picture(image_file, width=Inches(4))
                    
    except Exception as e:
        logger.error(f"Failed to add alignment: {e}")
        doc.add_paragraph("Failed to load alignment")


def _add_transcription_safe(doc, transcription, logger):
    """Safely add transcription"""
    try:
        transcription_para = doc.add_paragraph()
        transcription_para.add_run('🎙️ Transcription: ').bold = True
        
        if transcription.startswith("WEBVTT"):
            try:
                for i in webvtt.read_buffer(io.StringIO(transcription)):
                    doc.add_paragraph(f"({i.start} - {i.end}) {i.text}")
            except:
                doc.add_paragraph(transcription)
        else:
            doc.add_paragraph(transcription)
            
    except Exception as e:
        logger.error(f"Failed to add transcription: {e}")


def _add_translation_safe(doc, translation, logger):
    """Safely add translation"""
    try:
        translation_para = doc.add_paragraph()
        translation_para.add_run('🌐 Translation: ').bold = True
        
        if translation.startswith("WEBVTT"):
            try:
                for i in webvtt.read_buffer(io.StringIO(translation)):
                    doc.add_paragraph(f"({i.start} - {i.end}) {i.text}")
            except:
                doc.add_paragraph(translation)
        else:
            doc.add_paragraph(translation)
            
    except Exception as e:
        logger.error(f"Failed to add translation: {e}")


def _add_reagents_appendix(doc, protocol):
    """Add reagents appendix"""
    try:
        protocol_reagents = ProtocolReagent.objects.filter(protocol=protocol)
        step_reagents = StepReagent.objects.filter(step__protocol=protocol)
        
        if not protocol_reagents.exists() and not step_reagents.exists():
            return
        
        doc.add_page_break()
        doc.add_heading('Reagents and Materials', level=1)
        
        # Add protocol reagents table
        if protocol_reagents.exists():
            doc.add_heading('Protocol Reagents', level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Headers
            headers = ['Reagent', 'Quantity', 'Unit', 'Notes']
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            
            # Data
            for reagent in protocol_reagents:
                row_cells = table.add_row().cells
                row_cells[0].text = reagent.reagent.reagent_name
                row_cells[1].text = str(reagent.quantity_required or '')
                row_cells[2].text = reagent.unit or ''
                row_cells[3].text = reagent.notes or ''
                
    except Exception as e:
        # If appendix fails, just skip it
        pass


def _add_session_summary(doc, session):
    """Add session summary if available"""
    try:
        if not session:
            return
        
        doc.add_heading('Session Information', level=2)
        
        # Session info table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'
        
        session_info = [
            ('Session ID', session.unique_id),
            ('Session Name', session.name),
            ('Started', session.started_at.strftime('%B %d, %Y at %I:%M %p') if session.started_at else 'Not started'),
            ('Status', 'Enabled' if session.enabled else 'Disabled')
        ]
        
        for label, value in session_info:
            if value:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = str(value)
                row_cells[0].paragraphs[0].runs[0].bold = True
                
    except Exception as e:
        pass


def _notify_user_error(user_id, instance_id, error_message):
    """Notify user of error via WebSocket"""
    try:
        channel_layer = get_channel_layer()
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}",
            {
                "type": "error_message",
                "message": {
                    "error": error_message,
                    "instance_id": instance_id
                },
            },
        )
    except Exception as e:
        # If notification fails, just log it
        print(f"Failed to notify user of error: {e}")

def remove_file(file_path):
    if os.path.exists(file_path):
        os.remove(file_path)

def remove_html_tags(text):
    clean = re.compile('<.*?>')
    return re.sub(clean, '', text)

def convert_seconds_to_time(seconds):
    second_time = datetime.timedelta(seconds=seconds)
    # convert second to hour, minute, second
    return time(second_time.seconds // 3600, (second_time.seconds // 60) % 60, second_time.seconds % 60).strftime("%H:%M:%S")


def html_to_docx(html, doc: docx.Document):
    soup = BeautifulSoup(html, 'html.parser')
    for tag in soup:
        if tag.name == 'p':
            # Create a new paragraph in the docx file
            para = doc.add_paragraph()
            for subtag in tag:
                if subtag.name == 'b':
                    # Add a run with bold style
                    run = para.add_run(subtag.text)
                    run.bold = True
                elif subtag.name == 'i':
                    # Add a run with italic style
                    run = para.add_run(subtag.text)
                    run.italic = True
                elif subtag.name == 'span' and 'style' in subtag.attrs:
                    # Add a run with color style
                    if "rgb" in subtag['style']:
                        color = subtag['style'].split('(')[1].split(')')[0].split(',')
                        r, g, b = int(color[0]), int(color[1]), int(color[2])
                    else:
                        color = subtag['style'].split(':')[1]  # Assuming style="color: #RRGGBB"
                        r, g, b = int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
                    run = para.add_run(subtag.text)
                    run.font.color.rgb = RGBColor(r, g, b)
                else:
                    # Add a run with normal style
                    para.add_run(subtag.text)
        else:
            # Add a run with normal style
            doc.add_paragraph(tag.text)

@job('llama', timeout='1h')
def llama_summary(prompt: str, user_id: int, target: dict = None, instance_id: str = None):
    """
    Generate completion using llama model
    :param prompt:
    :param user_id:
    :return:
    """
    llama_bin_path = settings.LLAMA_BIN_PATH
    cmd = [llama_bin_path, "-p", prompt, "-m", settings.LLAMA_DEFAULT_MODEL, "-t", "8", "-c", "2048", "--repeat_penalty", "1.0", "--log-disable", "-n", "4096"]

    process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    finished = False
    started = False
    for line in process.stdout:
        line = line.decode("utf-8")
        print(line)
        if "<|im_start|>assistant" in line:
            started = True
        if "<|im_end|>" in line and started:
            finished = True
            break
        channel_layer = get_channel_layer()
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}_summary",
            {
                "type": "summary_message",
                "message": {
                    "data": line,
                    "type": "step",
                    "finished": finished,
                    "target": target,
                    "instance_id": instance_id
                },
            },
        )

    if finished:
        channel_layer = get_channel_layer()
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}_summary",
            {
                "type": "summary_message",
                "message": {
                    "data": line.split("<|im_end|>")[0],
                    "type": "step",
                    "finished": finished,
                    "target": target,
                    "instance_id": instance_id
                },
            },
        )


@job('llama', timeout='1h')
def llama_summary_transcript(prompt: str, user_id: int, target: dict = None, instance_id: str = None):
    """
    Provide summary from webtt transcript using llama model
    :param prompt:
    :param user_id:
    :param target:
    :return:
    """
    llama_bin_path = settings.LLAMA_BIN_PATH
    cmd = [llama_bin_path, "-p", prompt, "-m", settings.LLAMA_DEFAULT_MODEL, "-t", "8", "-c", "2048", "--repeat_penalty", "1.0", "--log-disable", "-n", "4096"]
    process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    started = False
    finished = False
    for line in process.stdout:
        line = line.decode("utf-8")
        print(line)
        if "<|im_start|>assistant" in line:
            started = True
        if "<|im_end|>" in line and started:
            finished = True
            break
        channel_layer = get_channel_layer()
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}_summary",
            {
                "type": "summary_message",
                "message": {
                    "data": line,
                    "type": "annotation",
                    "finished": finished,
                    "target": target,
                    "instance_id": instance_id
                },
            },
        )
    if finished:
        channel_layer = get_channel_layer()
        annotation = Annotation.objects.get(id=target["annotation"])
        annotation.summary = line.split("<|im_end|>")[0]
        annotation.save()
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}_summary",
            {
                "type": "summary_message",
                "message": {
                    "data": line.split("<|im_end|>")[0],
                    "type": "annotation",
                    "finished": finished,
                    "target": target,
                    "instance_id": instance_id
                },
            },
        )


@job('ocr', timeout='1h')
def ocr_b64_image(image_b64: str, annotation_id: int, session_id: str, instance_id: str = None):
    """
    Perform OCR on base64 image
    :param image_b64:
    :param user_id:
    :return:
    """

    imgstring = image_b64.split('base64,')[-1].strip()
    image_string = io.BytesIO(base64.b64decode(imgstring))
    image = Image.open(image_string)
    bg = Image.new("RGB", image.size, (255, 255, 255))
    bg.paste(image, image)
    data = pytesseract.image_to_string(bg)
    annotation = Annotation.objects.get(id=annotation_id)
    annotation.transcription = data
    annotation.transcribed = True
    annotation.save()
    print(data)
    data = AnnotationSerializer(annotation, many=False).data
    channel_layer = get_channel_layer()
    async_to_sync(channel_layer.group_send)(
        f"transcription_{session_id}",
        {
            "type": "transcription_message",
            "message": data,
        },
    )

@job('export', timeout='3h')
def export_sqlite(user_id:int, session_id: str = None, instance_id: str = None):
    """
    Export user data to sqlite database
    :param user_id:
    :return
    """
    #check if user created this session
    session = Session.objects.get(unique_id=session_id)
    if session.user.id != user_id:
        return
    #create a temporary folder in media temp folder
    media_root = settings.MEDIA_ROOT
    filename = uuid.uuid4().hex
    user_folder = os.path.join(media_root, "temp", filename)
    call_command("export_data_sqlite", session_id, user_folder)
    signer = TimestampSigner()
    value = signer.sign(f"{filename}.cupcake")
    channel_layer = get_channel_layer()
    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}",
        {
            "type": "download_message",
            "message": {
                "signed_value": value,
                "user_download": True,
                "instance_id": instance_id
            },
        },
    )

@job('export', timeout='3h')
def export_data(user_id: int, protocol_ids: list[int] = None, session_ids: list[int] = None, instance_id: str = None, export_type: str = "complete", format_type: str = "zip"):
    """
    Export user data using the revised export system with options for complete, protocol-specific, or session-specific exports
    
    :param user_id: ID of the user to export data for
    :param protocol_ids: Optional list of protocol IDs to limit export to specific protocols
    :param session_ids: Optional list of session IDs to limit export to specific sessions  
    :param instance_id: Optional instance ID for tracking the export job
    :param export_type: Type of export - "complete", "protocol", or "session"
    :param format_type: Archive format - "zip" or "tar.gz"
    """
    user = User.objects.get(id=user_id)
    channel_layer = get_channel_layer()
    
    # Send initial notification
    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}",
        {
            "type": "export_progress",
            "message": {
                "instance_id": instance_id,
                "progress": 0,
                "status": "starting",
                "message": f"Starting {export_type} export in {format_type} format...",
                "export_type": export_type
            },
        },
    )
    
    # Create export directory in Django media temp folder
    media_temp_dir = os.path.join(settings.MEDIA_ROOT, "temp")
    os.makedirs(media_temp_dir, exist_ok=True)
    
    # Clean up old export files (older than 7 days) to prevent storage buildup
    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}",
        {
            "type": "export_progress",
            "message": {
                "instance_id": instance_id,
                "progress": 5,
                "status": "preparing",
                "message": "Preparing export environment...",
                "export_type": export_type
            },
        },
    )
    _cleanup_old_exports(media_temp_dir)
    
    try:
        # Send progress notification for data collection
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}",
            {
                "type": "export_progress",
                "message": {
                    "instance_id": instance_id,
                    "progress": 10,
                    "status": "collecting",
                    "message": "Collecting and exporting data...",
                    "export_type": export_type
                },
            },
        )
        
        # Create progress callback
        def progress_callback(progress, message, status="processing"):
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}",
                {
                    "type": "export_progress",
                    "message": {
                        "instance_id": instance_id,
                        "progress": progress,
                        "status": status,
                        "message": message,
                        "export_type": export_type
                    },
                },
            )
        
        if export_type == "protocol" and protocol_ids:
            filename = export_protocol_data(user, protocol_ids, export_dir=media_temp_dir, format_type=format_type, progress_callback=progress_callback)
        elif export_type == "session" and session_ids:
            filename = export_session_data(user, session_ids, export_dir=media_temp_dir, format_type=format_type, progress_callback=progress_callback)
        else:
            filename = export_user_data_revised(user, export_dir=media_temp_dir, format_type=format_type, progress_callback=progress_callback)
        
        # Send progress notification for file creation
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}",
            {
                "type": "export_progress",
                "message": {
                    "instance_id": instance_id,
                    "progress": 90,
                    "status": "finalizing",
                    "message": "Creating archive and generating hash...",
                    "export_type": export_type
                },
            },
        )

        base_filename = os.path.basename(filename)
        
        # Create relative path for media URL
        relative_path = os.path.relpath(filename, settings.MEDIA_ROOT)
        
        signer = TimestampSigner()
        value = signer.sign(base_filename)
        
        # Send completion notification with download information
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}",
            {
                "type": "export_progress",
                "message": {
                    "instance_id": instance_id,
                    "progress": 100,
                    "status": "completed",
                    "message": f"Export completed successfully! File ready for download.",
                    "export_type": export_type,
                    "file_size": _get_file_size_mb(filename)
                },
            },
        )
        
        # Send download message
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}",
            {
                "type": "download_message",
                "message": {
                    "signed_value": value,
                    "user_download": True,
                    "instance_id": instance_id,
                    "export_type": export_type,
                    "export_path": filename,
                    "download_url": f"/media/{relative_path}",
                    "file_size_mb": _get_file_size_mb(filename)
                },
            },
        )
        
    except Exception as e:
        # Send error notification with progress information
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}",
            {
                "type": "export_progress",
                "message": {
                    "instance_id": instance_id,
                    "progress": -1,  # Indicates error
                    "status": "error",
                    "message": f"Export failed: {str(e)}",
                    "export_type": export_type,
                    "error_details": str(e)
                },
            },
        )
        
        # Also send the legacy error format for backward compatibility
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}",
            {
                "type": "export_error",
                "message": {
                    "error": str(e),
                    "instance_id": instance_id,
                    "export_type": export_type
                },
            },
        )

def export_user_data(user_id: int, filename: str = None, protocol_ids: list[int] = None):
    """
    A function that would export user's protocol, protocolstep, annotation, variation, and section data to individual jsons and archive into a tar.gz file
    :param user_id:
    :return:
    """
    user = User.objects.get(id=user_id)
    protocols = ProtocolModel.objects.filter(Q(user=user)|Q(steps__annotations__user=user))
    if protocol_ids:
        protocols = protocols.filter(id__in=protocol_ids)
    print(protocol_ids)
    protocols = protocols.distinct()
    protocolsSteps = ProtocolStep.objects.filter(protocol__in=protocols)
    variations = StepVariation.objects.filter(step__protocol__user=user)
    sections = ProtocolSection.objects.filter(protocol__user=user)
    if protocol_ids:
        sections = sections.filter(protocol__in=protocols)
    sessions = user.sessions.all()
    if protocol_ids:
        sessions = Session.objects.filter(annotations__step__protocol__id__in=protocol_ids).distinct()
    projects = []
    for session in sessions:
        for p in session.projects.filter(owner=user):
            if p not in projects:
                projects.append(p)
    annotations = Annotation.objects.filter(session__user=user)
    folders = AnnotationFolder.objects.filter(session__in=sessions)
    protocol_reagents = ProtocolReagent.objects.filter(protocol__in=protocols)
    step_reagents = StepReagent.objects.filter(step__in=protocolsSteps)
    reagents = Reagent.objects.filter(Q(protocolreagent__in=protocol_reagents) | Q(stepreagent__in=step_reagents))
    protocol_tags = ProtocolTag.objects.filter(protocol__in=protocols)
    step_tags = StepTag.objects.filter(step__in=protocolsSteps)
    tag = Tag.objects.filter(Q(protocoltag__in=protocol_tags) | Q(steptag__in=step_tags))

    protocols_data = ProtocolModelSerializer(protocols, many=True).data
    protocols_steps_data = ProtocolStepSerializer(protocolsSteps, many=True).data
    annotations_data = AnnotationSerializer(annotations, many=True).data
    variations_data = StepVariationSerializer(variations, many=True).data
    sections_data = ProtocolSectionSerializer(sections, many=True).data
    session_data = SessionSerializer(sessions, many=True).data
    folders_data = AnnotationFolderSerializer(folders, many=True).data
    reagents_data = ReagentSerializer(reagents, many=True).data
    protocol_reagents_data = ProtocolReagentSerializer(protocol_reagents, many=True).data
    step_reagents_data = StepReagentSerializer(step_reagents, many=True).data
    tags_data = TagSerializer(tag, many=True).data
    protocol_tags_data = ProtocolTagSerializer(protocol_tags, many=True).data
    step_tags_data = StepTagSerializer(step_tags, many=True).data
    projects_data = ProjectSerializer(projects, many=True).data



    media_root = settings.MEDIA_ROOT
    user_folder = os.path.join(media_root, "temp", f"{user.username}")
    os.makedirs(user_folder, exist_ok=True)
    os.makedirs(os.path.join(user_folder, "projects"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "protocols"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "steps"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "annotations"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "variations"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "sections"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "sessions"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "folders"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "reagents"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "protocol_reagents"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "step_reagents"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "tags"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "protocol_tags"), exist_ok=True)
    os.makedirs(os.path.join(user_folder, "step_tags"), exist_ok=True)

    for project in projects_data:
        with open(os.path.join(user_folder, "projects", f"{project['id']}.json"), "wt") as f:
            json.dump(
                {
                    "id": project["id"],
                    "project_name": project["project_name"],
                    "project_description": project["project_description"]
                }, f)

    for protocol in protocols_data:
        with open(os.path.join(user_folder, "protocols", f"{protocol['id']}.json"), "wt") as f:
            json.dump(protocol, f)
    for step in protocols_steps_data:
        with open(os.path.join(user_folder, "steps", f"{step['id']}.json"), "wt") as f:
            print(step)
            json.dump(step, f)
    for annotation in annotations_data:
        with open(os.path.join(user_folder, "annotations", f"{annotation['id']}.json"), "wt") as f:
            json.dump(annotation, f)
    for variation in variations_data:
        with open(os.path.join(user_folder, "variations", f"{variation['id']}.json"), "wt") as f:
            json.dump(variation, f)
    for section in sections_data:
        with open(os.path.join(user_folder, "sections", f"{section['id']}.json"), "wt") as f:
            json.dump(section, f)
    for session in session_data:
        with open(os.path.join(user_folder, "sessions", f"{session['id']}.json"), "wt") as f:
            json.dump(session, f)
    for folder in folders_data:
        with open(os.path.join(user_folder, "folders", f"{folder['id']}.json"), "wt") as f:
            json.dump(folder, f)
    for reagent in reagents_data:
        with open(os.path.join(user_folder, "reagents", f"{reagent['id']}.json"), "wt") as f:
            json.dump(reagent, f)
    for protocol_ingredient in protocol_reagents_data:
        with open(os.path.join(user_folder, "protocol_reagents", f"{protocol_ingredient['id']}.json"), "wt") as f:
            json.dump(protocol_ingredient, f)
    for step_ingredient in step_reagents_data:
        with open(os.path.join(user_folder, "step_reagents", f"{step_ingredient['id']}.json"), "wt") as f:
            json.dump(step_ingredient, f)
    for tag in tags_data:
        with open(os.path.join(user_folder, "tags", f"{tag['id']}.json"), "wt") as f:
            json.dump(tag, f)
    for protocol_tag in protocol_tags_data:
        with open(os.path.join(user_folder, "protocol_tags", f"{protocol_tag['id']}.json"), "wt") as f:
            json.dump(protocol_tag, f)

    for step_tag in step_tags_data:
        with open(os.path.join(user_folder, "step_tags", f"{step_tag['id']}.json"), "wt") as f:
            json.dump(step_tag, f)

    os.makedirs(os.path.join(user_folder, "media", "annotations"), exist_ok=True)
    for i in annotations_data:
        if i["file"]:
            correct_path = i["file"].replace("/media/app/", "/app/")
            print(i)
            print(os.path.exists(correct_path))
            if not os.path.exists(correct_path):
                correct_path = "/app" + i["file"]
                print(os.path.exists(correct_path))
            new_path = os.path.join(user_folder, "media", "annotations", os.path.basename(correct_path))
            print(correct_path)
            print(new_path)
            shutil.copy(correct_path, new_path)
    if not filename:
        filename = str(uuid.uuid4())
        shutil.make_archive(os.path.join(media_root, "temp", f"{filename}"), 'gztar', user_folder)
    else:
        shutil.make_archive(filename, 'gztar', user_folder)
    shutil.rmtree(user_folder)
    return filename

@job('import-data', timeout='3h')
def import_data(user_id: int, archive_file: str, instance_id: str = None, import_options: dict = None, storage_object_mappings: dict = None, bulk_transfer_mode: bool = False):
    """
    Import user data with progress tracking and selective import options
    
    :param user_id: ID of the user to import data for
    :param archive_file: Path to the archive file (zip or tar.gz)
    :param instance_id: Optional instance ID for tracking the import job
    :param import_options: Optional dict specifying what to import (protocols, sessions, etc.)
    :param storage_object_mappings: Optional dict mapping original storage IDs to nominated storage IDs
    :param bulk_transfer_mode: If True, import everything as-is without user-centric modifications
    """
    user = User.objects.get(id=user_id)
    channel_layer = get_channel_layer()
    
    # Check site settings to filter import options
    try:
        site_settings = SiteSettings.objects.filter(is_active=True).first()
        if site_settings and import_options:
            # Filter import options based on site settings (unless user is staff)
            filtered_options = site_settings.filter_import_options(import_options, user)
            if filtered_options != import_options:
                # Send notification about restricted options
                restricted_items = [k for k in import_options if k not in filtered_options or not filtered_options[k]]
                if restricted_items:
                    async_to_sync(channel_layer.group_send)(
                        f"user_{user_id}",
                        {
                            "type": "import_progress",
                            "message": {
                                "instance_id": instance_id,
                                "progress": 0,
                                "status": "warning",
                                "message": f"Some import options restricted by site settings: {', '.join(restricted_items)}",
                                "import_type": "user_data"
                            },
                        },
                    )
            import_options = filtered_options
    except Exception as e:
        print(f"Warning: Could not check site settings for import restrictions: {e}")
    
    # Send initial notification
    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}",
        {
            "type": "import_progress",
            "message": {
                "instance_id": instance_id,
                "progress": 0,
                "status": "starting",
                "message": "Starting import process...",
                "import_type": "user_data"
            },
        },
    )
    
    try:
        # Create progress callback
        def progress_callback(progress, message, status="processing"):
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}",
                {
                    "type": "import_progress",
                    "message": {
                        "instance_id": instance_id,
                        "progress": progress,
                        "status": status,
                        "message": message,
                        "import_type": "user_data"
                    },
                },
            )
        
        # Perform the import with progress tracking
        result = import_user_data_revised(
            user, 
            archive_file, 
            import_options=import_options,
            progress_callback=progress_callback,
            storage_object_mappings=storage_object_mappings,
            bulk_transfer_mode=bulk_transfer_mode
        )
        
        if result['success']:
            # Send completion notification
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}",
                {
                    "type": "import_progress",
                    "message": {
                        "instance_id": instance_id,
                        "progress": 100,
                        "status": "completed",
                        "message": f"Import completed successfully! {result['stats']['models_imported']} models imported.",
                        "import_type": "user_data",
                        "stats": result['stats']
                    },
                },
            )
        else:
            # Send error notification
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}",
                {
                    "type": "import_progress",
                    "message": {
                        "instance_id": instance_id,
                        "progress": -1,
                        "status": "error",
                        "message": f"Import failed: {result.get('error', 'Unknown error')}",
                        "import_type": "user_data",
                        "error_details": result.get('error', 'Unknown error')
                    },
                },
            )
        
    except Exception as e:
        # Send error notification
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}",
            {
                "type": "import_progress",
                "message": {
                    "instance_id": instance_id,
                    "progress": -1,
                    "status": "error",
                    "message": f"Import failed: {str(e)}",
                    "import_type": "user_data",
                    "error_details": str(e)
                },
            },
        )


@job('import-data', timeout='1h')
def dry_run_import_data(user_id: int, archive_file: str, instance_id: str = None, import_options: dict = None, bulk_transfer_mode: bool = False):
    """
    Perform a dry run analysis of user data import without making any changes
    
    :param user_id: ID of the user to analyze import for
    :param archive_file: Path to the archive file (zip or tar.gz)
    :param instance_id: Optional instance ID for tracking the analysis job
    :param import_options: Optional dict specifying what to analyze for import
    """
    
    user = User.objects.get(id=user_id)
    channel_layer = get_channel_layer()
    
    # Check site settings to filter import options (same as actual import)
    try:
        site_settings = SiteSettings.objects.filter(is_active=True).first()
        if site_settings and import_options:
            filtered_options = site_settings.filter_import_options(import_options, user)
            if filtered_options != import_options:
                restricted_items = [k for k in import_options if k not in filtered_options or not filtered_options[k]]
                if restricted_items:
                    async_to_sync(channel_layer.group_send)(
                        f"user_{user_id}",
                        {
                            "type": "import_progress",
                            "message": {
                                "instance_id": instance_id,
                                "progress": 0,
                                "status": "warning",
                                "message": f"Some import options will be restricted: {', '.join(restricted_items)}",
                                "import_type": "dry_run_analysis"
                            },
                        },
                    )
            import_options = filtered_options
    except Exception as e:
        print(f"Warning: Could not check site settings for dry run: {e}")
    
    # Send initial notification
    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}",
        {
            "type": "import_progress",
            "message": {
                "instance_id": instance_id,
                "progress": 0,
                "status": "analyzing",
                "message": "Starting import analysis...",
                "import_type": "dry_run_analysis"
            },
        },
    )
    
    try:
        # Create progress callback
        def progress_callback(progress, message, status="analyzing"):
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}",
                {
                    "type": "import_progress",
                    "message": {
                        "instance_id": instance_id,
                        "progress": progress,
                        "status": status,
                        "message": message,
                        "import_type": "dry_run_analysis"
                    },
                },
            )
        
        # Perform the dry run analysis
        result = dry_run_import_user_data(
            user, 
            archive_file, 
            import_options=import_options,
            progress_callback=progress_callback
        )
        
        if result['success']:
            # Send completion notification with analysis report
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}",
                {
                    "type": "import_progress", 
                    "message": {
                        "instance_id": instance_id,
                        "progress": 100,
                        "status": "completed",
                        "message": "Import analysis completed successfully!",
                        "import_type": "dry_run_analysis",
                        "analysis_report": result['analysis_report'],
                        "metadata": result.get('metadata', {}),
                        "archive_format": result.get('archive_format', 'unknown')
                    },
                },
            )
        else:
            # Send error notification
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}",
                {
                    "type": "import_progress",
                    "message": {
                        "instance_id": instance_id,
                        "progress": -1,
                        "status": "error",
                        "message": f"Import analysis failed: {result.get('error', 'Unknown error')}",
                        "import_type": "dry_run_analysis",
                        "error_details": result.get('error', 'Unknown error'),
                        "analysis_report": result.get('analysis_report', {})
                    },
                },
            )
        
    except Exception as e:
        # Send error notification
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}",
            {
                "type": "import_progress",
                "message": {
                    "instance_id": instance_id,
                    "progress": -1,
                    "status": "error",
                    "message": f"Import analysis failed: {str(e)}",
                    "import_type": "dry_run_analysis",
                    "error_details": str(e)
                },
            },
        )


def import_user_data(user_id: int, tar_file: str, instance_id: str = None):
    """
    A function that would import user's protocol, protocolstep, annotation, variation, and section data from individual jsons and archive into a tar.gz file
    :param user_id:
    :param tar_file:
    :return:
    """
    user = User.objects.get(id=user_id)
    media_root = settings.MEDIA_ROOT
    user_folder = os.path.join(media_root, "temp", f"{user.username}")
    annotation_folder = os.path.join(media_root, "annotations")
    os.makedirs(user_folder, exist_ok=True)
    shutil.unpack_archive(tar_file, user_folder, format="gztar")
    projects_folder = os.path.join(user_folder, "projects")
    protocols_folder = os.path.join(user_folder, "protocols")
    steps_folder = os.path.join(user_folder, "steps")
    annotations_folder = os.path.join(user_folder, "annotations")
    variations_folder = os.path.join(user_folder, "variations")
    sections_folder = os.path.join(user_folder, "sections")
    sessions_folder = os.path.join(user_folder, "sessions")
    annotationfolder_folder = os.path.join(user_folder, "folders")
    reagents_folder = os.path.join(user_folder, "reagents")
    protocol_reagents_folder = os.path.join(user_folder, "protocol_reagents")
    step_reagents_folder = os.path.join(user_folder, "step_reagents")
    tags_folder = os.path.join(user_folder, "tags")
    protocol_tags_folder = os.path.join(user_folder, "protocol_tags")
    step_tags_folder = os.path.join(user_folder, "step_tags")
    protocol_map = {}

    project_map = {}
    for project in os.listdir(projects_folder):
        with open(os.path.join(projects_folder, project), "rt") as f:
            data = json.load(f)
            project = Project()
            project.project_name = "[imported]" + data["project_name"] + str(datetime.datetime.now())
            project.project_description = data["project_description"]
            project.user = user
            project.save()
            project_map[data["id"]] = project

    for file in os.listdir(protocols_folder):
        with open(os.path.join(protocols_folder, file), "rt") as f:
            data = json.load(f)
            data["user"] = user_id
            if ProtocolModel.objects.filter(protocol_title=data["protocol_title"]).exists():
                date = datetime.datetime.now()
                data["protocol_title"] = "[imported] " + data["protocol_title"] + " " + str(date)

            protocol = ProtocolModel()
            protocol.protocol_created_on = data["protocol_created_on"]
            protocol.protocol_title = data["protocol_title"]
            protocol.protocol_description = data["protocol_description"]
            protocol.user = user
            protocol.save()
            protocol_map[data["id"]] = protocol
    tag_map = {}
    for file in os.listdir(tags_folder):
        with open(os.path.join(tags_folder, file), "rt") as f:
            data = json.load(f)
            instance = Tag.objects.get_or_create(tag=data["tag"])[0]
            tag_map[data["id"]] = instance
    protocol_tag_map = {}

    for file in os.listdir(protocol_tags_folder):
        with open(os.path.join(protocol_tags_folder, file), "rt") as f:
            data = json.load(f)
            data["protocol"] = protocol_map[data["protocol"]]
            data["tag"] = tag_map[data["tag"]["id"]]
            protocol_tag_data = {k: v for k, v in data.items() if k != "id"}
            instance = ProtocolTag.objects.create(**protocol_tag_data)
            protocol_tag_map[data["id"]] = instance

    section_map = {}
    for file in os.listdir(sections_folder):
        with open(os.path.join(sections_folder, file), "rt") as f:
            data = json.load(f)
            data["protocol"] = protocol_map[data["protocol"]].id
            serializer = ProtocolSectionSerializer(data=data)
            if serializer.is_valid():
                section_data = {k: v for k, v in serializer.validated_data.items() if k != "id" and k != "reagents" and k != "tags"}
                instance = ProtocolSection.objects.create(**section_data)
                instance.remote_id = data["id"]
                section_map[data["id"]] = instance
                instance.save()
    step_map = {}
    for file in os.listdir(steps_folder):
        with open(os.path.join(steps_folder, file), "rt") as f:
            data = json.load(f)
            data["protocol"] = protocol_map[data["protocol"]]
            data["step_section"] = section_map[data["step_section"]]
            step_data = {k: v for k, v in data.items() if k != "id" and k != "previous_step" and k != "next_step" and k != "annotations" and k != "variations" and k != "reagents" and k != "tags"}
            instance = ProtocolStep.objects.create(**step_data)
            step_map[data["id"]] = instance
            if data["previous_step"]:
                if data["previous_step"] in step_map:
                    step_map[data["previous_step"]].next_step.add(instance)
    step_tag_map = {}
    for file in os.listdir(step_tags_folder):
        with open(os.path.join(step_tags_folder, file), "rt") as f:
            data = json.load(f)
            data["step"] = step_map[data["step"]]
            data["tag"] = tag_map[data["tag"]["id"]]
            step_tag_data = {k: v for k, v in data.items() if k != "id"}
            instance = StepTag.objects.create(**step_tag_data)
            step_tag_map[data["id"]] = instance

    reagent_map = {}
    for file in os.listdir(reagents_folder):
        with open(os.path.join(reagents_folder, file), "rt") as f:
            data = json.load(f)
            serializer = ReagentSerializer(data=data)
            if serializer.is_valid():
                reagent_data = {k: v for k, v in serializer.validated_data.items() if k != "id"}
                instance = Reagent.objects.get_or_create(name=reagent_data["name"], unit=reagent_data["unit"])[0]
                reagent_map[data["id"]] = instance

    protocol_reagent_map = {}
    for file in os.listdir(protocol_reagents_folder):
        with open(os.path.join(protocol_reagents_folder, file), "rt") as f:
            data = json.load(f)
            data["protocol"] = protocol_map[data["protocol"]]
            data["reagent"] = reagent_map[data["reagent"]["id"]]
            protocol_reagent_data = {k: v for k, v in data.items() if k != "id"}

            instance = ProtocolReagent.objects.create(**protocol_reagent_data)
            protocol_reagent_map[data["id"]] = instance

    step_reagent_map = {}
    for file in os.listdir(step_reagents_folder):
        with open(os.path.join(step_reagents_folder, file), "rt") as f:
            data = json.load(f)
            data["step"] = step_map[data["step"]]
            data["reagent"] = reagent_map[data["reagent"]["id"]]
            step_reagent_data = {k: v for k, v in data.items() if k != "id"}
            instance = StepReagent.objects.create(**step_reagent_data)
            step_reagent_map[data["id"]] = instance
            description = data["step"].step_description
            for i in [f"%{data['id']}.name%", f"%{data['id']}.quantity%", f"%{data['id']}.unit%", f"%{data['id']}.scaled_quantity%"]:
                if i in description:
                    if i == f"%{data['id']}.name%":
                        description = description.replace(i, f"%{instance.id}.name%")
                    if i == f"%{data['id']}.quantity%":
                        description = description.replace(i, f"%{instance.id}.quantity%")
                    if i == f"%{data['id']}.unit%":
                        description = description.replace(i, f"%{instance.id}.unit%")
                    if i == f"%{data['id']}.scaled_quantity%":
                        description = description.replace(i, f"%{instance.id}.scaled_quantity%")
            data["step"].step_description = description
            data["step"].save()

    session_map = {}
    for file in os.listdir(sessions_folder):
        with open(os.path.join(sessions_folder, file), "rt") as f:
            data = json.load(f)
            data["user"] = user_id
            data["time_keeper"] = []
            data["protocols"] = [protocol_map[i].id for i in data["protocols"] if i in protocol_map]
            session = Session()
            session.user = user
            session.created_at = data["created_at"]
            session.updated_at = data["updated_at"]
            session.remote_id = data["id"]
            session.started_at = data["started_at"]
            session.ended_at = data["ended_at"]
            session.unique_id = uuid.uuid4().hex
            session.name = data["name"]
            session.processing = True
            session.save()
            for protocol in data["protocols"]:
                session.protocols.add(protocol)
            if "projects" in data:
                for project in data["projects"]:
                    session.projects.add(project_map[project])

            session_map[data["id"]] = session


    annotationfolder_map = {}
    annotationfolder_json_map = {}
    for file in os.listdir(annotationfolder_folder):
        with open(os.path.join(annotationfolder_folder, file), "rt") as f:
            data = json.load(f)
            annotationfolder_json_map[data["id"]] = data
            annotationfolder_map[data["id"]] = AnnotationFolder.objects.create(
                folder_name=data["folder_name"],
                session=session_map[data["session"]],
                parent_folder=None,
            )
    for i in annotationfolder_json_map:
        data = annotationfolder_json_map[i]
        if data["parent_folder"]:
            annotationfolder_map[i] = annotationfolder_map[data["parent_folder"]]
            annotationfolder_map[i].save()

    annotation_map = {}
    for file in os.listdir(annotations_folder):
        with open(os.path.join(annotations_folder, file), "rt") as f:

            data = json.load(f)
            folder = None
            if data["session"] not in session_map:
                continue
            if data["step"]:
                if data["step"] in step_map:
                    data["step"] = step_map[data["step"]]
                else:
                    continue
                data["remote_id"] = data["id"]

                if len(data["folder"]) > 0:
                    folder = annotationfolder_map[data["folder"][0]["id"]]
            instance = Annotation.objects.create(
                session=session_map[data["session"]],
                step=data["step"],
                annotation_type=data["annotation_type"],
                remote_id=data["id"],
                folder=folder
            )
            instance.annotation = data["annotation"]
            instance.transcribed = data["transcribed"]
            instance.transcription = data["transcription"]
            instance.language = data["language"]
            instance.translation = data["translation"]
            instance.user = user
            if data["file"]:
                file_path = os.path.join(user_folder, "media", "annotations", os.path.basename(data["file"]))
                #new_file_path = os.path.join(annotation_folder, str(user_id) + os.path.basename(data["file"]))
                #new_file_path = os.path.normpath(new_file_path)
                #shutil.copy(file_path, new_file_path)
                with open(file_path, "rb") as f:
                    instance.file.save(uuid.uuid4().hex + os.path.basename(data["file"]), f)
            instance.created_at = data["created_at"]
            instance.updated_at = data["updated_at"]
            instance.save()
            annotation_map[data["id"]] = instance

    for file in os.listdir(variations_folder):
        with open(os.path.join(variations_folder, file), "rt") as f:
            data = json.load(f)
            if data["step"] in step_map:
                data["step"] = step_map[data["step"]].id

            serializer = StepVariationSerializer(data=data)
            if serializer.is_valid():
                serializer.save()

    for i in session_map:
        session = session_map[i]
        session.processing = False
        session.save()

@job('export', timeout='3h')
def export_instrument_job_metadata(instrument_job_id: int, data_type: str, user_id: int, instance_id: str = None):
    instrument_job = InstrumentJob.objects.get(id=instrument_job_id)
    if data_type == "user_metadata":
        metadata = instrument_job.user_metadata.all()
    elif data_type == "staff_metadata":
        metadata = instrument_job.staff_metadata.all()
    else:
        metadata = list(instrument_job.user_metadata.all()) + list(instrument_job.staff_metadata.all())
    
    # Get reference pools for this instrument job (only reference pools appear in SDRF)
    from cc.models import SamplePool
    pools = SamplePool.objects.filter(instrument_job=instrument_job, is_reference=True)
    
    result, _ = sort_metadata(metadata, instrument_job.sample_number)
    
    # Add reference pool rows to the result if they exist
    if pools.exists():
        result = add_pool_rows_to_sdrf(result, pools, data_type)
    
    # create tsv file from result
    filename = str(uuid.uuid4())
    tsv_filepath = os.path.join(settings.MEDIA_ROOT, "temp", f"{filename}.tsv")
    with open(tsv_filepath, "wt") as f:
        writer = csv.writer(f, delimiter="\t")
        writer.writerow(result[0])
        for row in result[1:]:
            writer.writerow(row)
    signer = TimestampSigner()
    value = signer.sign(f"{filename}.tsv")
    channel_layer = get_channel_layer()
    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}_instrument_job",
        {
            "type": "download_message",
            "message": {
                "signed_value": value,
                "instance_id": instance_id
            },
        }
    )
    return value


def add_pool_rows_to_sdrf(result, pools, data_type):
    """Add pool rows to SDRF export data"""
    headers = result[0]
    data_rows = result[1:]
    
    # Find column indices
    source_name_col = None
    pooled_sample_col = None
    
    for i, header in enumerate(headers):
        header_lower = header.lower()
        if header_lower == "source name":
            source_name_col = i
        elif "pooled sample" in header_lower or "pooled_sample" in header_lower:
            pooled_sample_col = i
    
    # If no pooled sample column exists, add it
    if pooled_sample_col is None:
        headers.append("characteristics[pooled sample]")
        pooled_sample_col = len(headers) - 1
        # Add empty values to existing rows
        for row in data_rows:
            row.append("not pooled")
    
    # Add pool rows
    for pool in pools:
        pool_row = ["" for _ in headers]
        
        # Get pool metadata
        if data_type == "user_metadata":
            pool_metadata = pool.user_metadata.all()
        elif data_type == "staff_metadata":
            pool_metadata = pool.staff_metadata.all()
        else:
            pool_metadata = list(pool.user_metadata.all()) + list(pool.staff_metadata.all())
        
        # Fill pool row with metadata values
        for metadata_column in pool_metadata:
            column_name = metadata_column.name.lower()
            
            # Find matching header
            for i, header in enumerate(headers):
                header_lower = header.lower()
                
                # Match various header formats
                if (header_lower == column_name or 
                    header_lower == f"characteristics[{column_name}]" or
                    header_lower == f"comment[{column_name}]" or
                    header_lower.replace("[", "").replace("]", "") == column_name):
                    
                    pool_row[i] = metadata_column.value
                    break
        
        # Ensure source name is set to pool name
        if source_name_col is not None:
            pool_row[source_name_col] = pool.pool_name
        
        # Ensure pooled sample column has the SDRF value
        if pooled_sample_col is not None:
            pool_row[pooled_sample_col] = pool.sdrf_value
        
        data_rows.append(pool_row)
    
    # Update pooled sample values for regular samples
    for pool in pools:
        # Mark pooled_only_samples as "pooled"
        for sample_index in pool.pooled_only_samples:
            if sample_index - 1 < len(data_rows) and pooled_sample_col is not None:
                data_rows[sample_index - 1][pooled_sample_col] = "pooled"
        
        # Mark pooled_and_independent_samples as "not pooled" (they remain independent)
        for sample_index in pool.pooled_and_independent_samples:
            if sample_index - 1 < len(data_rows) and pooled_sample_col is not None:
                data_rows[sample_index - 1][pooled_sample_col] = "not pooled"
    
    return [headers] + data_rows


def sort_metadata(metadata: list[MetadataColumn]|QuerySet, sample_number: int):
    id_metadata_column_map = {}
    headers = []
    default_columns_list = [{
        "name": "Source name", "type": "", "mandatory": True
    },
        {
            "name": "Organism", "type": "Characteristics", "mandatory": True
        }, {
            "name": "Organism part", "type": "Characteristics", "mandatory": True
        }, {
            "name": "Disease", "type": "Characteristics", "mandatory": True
        }, {
            "name": "Cell type", "type": "Characteristics", "mandatory": True
        }, {
            "name": "Biological replicate", "type": "Characteristics", "mandatory": True
        },
        {
          "name": 'Enrichment process', 'type': 'Characteristics', 'mandatory': True
        },
        {
            "name": "Material type", "type": "", "mandatory": True
        },
        {
            "name": "Assay name", "type": "", "mandatory": True
        }, {
            "name": "Technology type", "type": "", "mandatory": True
        },
        {
            "name": "Proteomics data acquisition method", "type": "Comment", "mandatory": True
        }

        , {"name": "Label", "type": "Comment", "mandatory": True},
        {"name": "Instrument", "type": "Comment", "mandatory": True},
        {"name": "Fraction identifier", "type": "Comment", "mandatory": True},
        {
            "name": "Technical replicate", "type": "Comment", "mandatory": True
        },
        {"name": "Cleavage agent details", "type": "Comment", "mandatory": True},
        {"name": "Modification parameters", "type": "Comment", "mandatory": True},
        {"name": "Dissociation method", "type": "Comment", "mandatory": True},
        {"name": "Precursor mass tolerance", "type": "Comment", "mandatory": True},
        {"name": "Fragment mass tolerance", "type": "Comment", "mandatory": True},
        {"name": "File uri", "type": "Comment", "mandatory": True},
        {"name": "Data file", "type": "Comment", "mandatory": True},
        {"name": "", "type": "Factor value", "mandatory": True},
    ]
    metadata_column_map = {}
    source_name_metadata = None
    assay_name_metadata = None
    material_type_metadata = None
    technology_type_metadata = None
    factor_value_columns = []
    metadata_cache = {}
    for m in metadata:
        if m.name not in metadata_cache:
            metadata_cache[m.name] = {}
        if m.value not in metadata_cache[m.name]:
            metadata_cache[m.name][m.value] = convert_metadata_column_value_to_sdrf(m.name.lower(), m.value)
        m.value = metadata_cache[m.name][m.value]
        if m.modifiers:
            m.modifiers = json.loads(m.modifiers)
            if m.modifiers:
                for n, mod in enumerate(m.modifiers):
                    if mod["value"] not in metadata_cache[m.name]:
                        metadata_cache[m.name][mod["value"]] = convert_metadata_column_value_to_sdrf(m.name.lower(), mod["value"])
                    m.modifiers[n]["value"] = metadata_cache[m.name][mod["value"]]
            else:
                m.modifiers = []
        else:
            m.modifiers = []
        if m.type != "Factor value":
            if m.name not in metadata_column_map:
                metadata_column_map[m.name] = []
            metadata_column_map[m.name].append(m)
        else:
            factor_value_columns.append(m)
    new_metadata = []

    non_default_columns = []
    default_column_map = {}
    for i in default_columns_list:
        default_column_map[i["name"]] = i
        if i["name"] in metadata_column_map and i["name"] != "Assay name" and i["name"] != "Source name" and i["name"] != "Material type" and i["name"] != "Technology type":
            new_metadata.extend(metadata_column_map[i["name"]])
        if i["name"] == "Assay name":
            if "Assay name" in metadata_column_map:
                assay_name_metadata = metadata_column_map[i["name"]][0]
        elif i["name"] == "Source name":
            if "Source name" in metadata_column_map:
                source_name_metadata = metadata_column_map["Source name"][0]
        elif i["name"] == "Material type":
            if "Material type" in metadata_column_map:
                material_type_metadata = metadata_column_map["Material type"][0]
        elif i["name"] == "Technology type":
            if "Technology type" in metadata_column_map:
                technology_type_metadata = metadata_column_map["Technology type"][0]

    for name in metadata_column_map:
        if name not in default_column_map and name != "Assay name" and name != "Source name" and name != "Material type" and name != "Technology type":
            non_default_columns.extend(metadata_column_map[name])

    # render an empty 2d array with number of row equal to job sample_number and number of columns equalt to number of metadata columns
    col_count = len(new_metadata) + len(non_default_columns)
    if source_name_metadata:
        col_count += 1
    if assay_name_metadata:
        col_count += 1
    if len(factor_value_columns) > 0:
        col_count += len(factor_value_columns)
    if material_type_metadata:
        col_count += 1
    if technology_type_metadata:
        col_count += 1
    data = [["" for i in range(col_count)] for j in range(sample_number)]
    # render in order, source name, characteristics, non type, comment and factor values
    # fill first column with source name
    last_characteristics = 0
    if source_name_metadata:
        headers.append("source name")
        if source_name_metadata.modifiers:
            modifiers = source_name_metadata.modifiers
            for m in modifiers:
                samples = parse_sample_indices_from_modifier_string(m["samples"])
                for s in samples:
                    data[s][0] = m["value"]
        for i in range(sample_number):
            if data[i][0] == "":
                data[i][0] = source_name_metadata.value
        id_metadata_column_map[source_name_metadata.id] = {"column": 0, "name": "source name", "type": "", "hidden": source_name_metadata.hidden}
        last_characteristics += 1
    # fill characteristics
    for i in range(0, len(new_metadata)):
        m = new_metadata[i]
        if m.type == "Characteristics":
            if m.name.lower() == 'tissue' or m.name.lower() == "organism part":
                headers.append("characteristics[organism part]")
            else:
                headers.append(f"characteristics[{m.name.lower()}]")
            if m.modifiers:
                modifiers = m.modifiers
                if modifiers:
                    for mod in modifiers:
                        samples = parse_sample_indices_from_modifier_string(mod["samples"])
                        for s in samples:
                            data[s][last_characteristics] = mod["value"]
            for j in range(sample_number):
                if data[j][last_characteristics] == "":
                    data[j][last_characteristics] = m.value
            id_metadata_column_map[m.id] = {"column": last_characteristics, "name": headers[-1], "type": "characteristics", "hidden": m.hidden}
            last_characteristics += 1
    # fill characteristics from non default columns
    for i in range(0, len(non_default_columns)):
        m = non_default_columns[i]
        if m.type == "Characteristics":
            headers.append(f"characteristics[{m.name.lower()}]")
            if m.modifiers:
                modifiers = m.modifiers
                if modifiers:
                    for mod in modifiers:
                        samples = parse_sample_indices_from_modifier_string(mod["samples"])
                        for s in samples:
                            data[s][last_characteristics] = mod["value"]
            for j in range(sample_number):
                if data[j][last_characteristics] == "":
                    data[j][last_characteristics] = m.value
            id_metadata_column_map[m.id] = {"column": last_characteristics, "name": headers[-1], "type": "characteristics", "hidden": m.hidden}
            last_characteristics += 1
    # fill material type column
    last_non_type = last_characteristics
    if material_type_metadata:
        headers.append("material type")
        if material_type_metadata.modifiers:
            modifiers = material_type_metadata.modifiers
            if modifiers:
                for m in modifiers:
                    samples = parse_sample_indices_from_modifier_string(m["samples"])
                    for s in samples:
                        data[s][last_non_type] = m["value"]
        for i in range(sample_number):
            if data[i][last_non_type] == "":
                data[i][last_non_type] = material_type_metadata.value
        id_metadata_column_map[material_type_metadata.id] = {"column": last_non_type, "name": "material type", "type": "", "hidden": material_type_metadata.hidden}
        last_non_type += 1
    # fill assay name column
    if assay_name_metadata:
        headers.append("assay name")
        if assay_name_metadata.modifiers:
            modifiers = assay_name_metadata.modifiers
            if modifiers:
                for m in modifiers:
                    samples = parse_sample_indices_from_modifier_string(m["samples"])
                    for s in samples:
                        data[s][last_non_type] = m["value"]
        for i in range(sample_number):
            if data[i][last_non_type] == "":
                data[i][last_non_type] = assay_name_metadata.value
        id_metadata_column_map[assay_name_metadata.id] = {"column": last_non_type, "name": "assay name", "type": "", "hidden": assay_name_metadata.hidden}
        last_non_type += 1
    # fill technology type column
    if technology_type_metadata:
        headers.append("technology type")
        if technology_type_metadata.modifiers:
            modifiers = technology_type_metadata.modifiers
            if modifiers:
                for m in modifiers:
                    samples = parse_sample_indices_from_modifier_string(m["samples"])
                    for s in samples:
                        data[s][last_non_type] = m["value"]
        for i in range(sample_number):
            if data[i][last_non_type] == "":
                data[i][last_non_type] = technology_type_metadata.value
        id_metadata_column_map[technology_type_metadata.id] = {"column": last_non_type, "name": "technology type", "type": "", "hidden": technology_type_metadata.hidden}
        last_non_type += 1
    # fill non type column
    for i in range(0, len(new_metadata)):
        m = new_metadata[i]
        if m.type == "":
            headers.append(m.name.lower())
            if m.modifiers:
                modifiers = m.modifiers
                if modifiers:
                    for mod in modifiers:
                        samples = parse_sample_indices_from_modifier_string(mod["samples"])
                        for s in samples:
                            data[s][last_non_type] = mod["value"]

            for j in range(sample_number):
                if data[j][last_non_type] == "":
                    data[j][last_non_type] = m.value
            id_metadata_column_map[m.id] = {"column": last_non_type, "name": headers[-1], "type": "", "hidden": m.hidden}
            last_non_type += 1
    # fill non type from non default columns
    for i in range(0, len(non_default_columns)):
        m = non_default_columns[i]
        if m.type == "":
            headers.append(m.name.lower())
            if m.modifiers:
                modifiers = m.modifiers
                if modifiers:
                    for mod in modifiers:
                        samples = parse_sample_indices_from_modifier_string(mod["samples"])
                        for s in samples:
                            data[s][last_non_type] = mod["value"]
            for j in range(sample_number):
                if data[j][last_non_type] == "":
                    data[j][last_non_type] = m.value
            id_metadata_column_map[m.id] = {"column": last_non_type, "name": headers[-1], "type": "", "hidden": m.hidden}
            last_non_type += 1
    # fill comment column
    last_comment = last_non_type
    for i in range(0, len(new_metadata)):
        m = new_metadata[i]
        if m.type == "Comment":
            headers.append(f"comment[{m.name.lower()}]")

            if m.modifiers:
                modifiers = m.modifiers
                if modifiers:
                    for mod in modifiers:
                        samples = parse_sample_indices_from_modifier_string(mod["samples"])
                        for s in samples:
                            data[s][last_comment] = mod["value"]
            for j in range(sample_number):
                if data[j][last_comment] == "":
                    data[j][last_comment] = m.value
            id_metadata_column_map[m.id] = {"column": last_comment, "name": headers[-1], "type": "comment", "hidden": m.hidden}
            last_comment += 1
    # fill comment from non default columns
    for i in range(0, len(non_default_columns)):
        m = non_default_columns[i]
        if m.type == "Comment":
            headers.append(f"comment[{m.name.lower()}]")

            if m.modifiers:
                modifiers = m.modifiers
                if modifiers:
                    for mod in modifiers:
                        samples = parse_sample_indices_from_modifier_string(mod["samples"])
                        for s in samples:
                            data[s][last_comment] = mod["value"]
            for j in range(sample_number):
                if data[j][last_comment] == "":
                    data[j][last_comment] = m.value
            id_metadata_column_map[m.id] = {"column": last_comment, "name": headers[-1], "type": "comment",  "hidden": m.hidden}
            last_comment += 1
    # write factor values

    for i in range(0, len(factor_value_columns)):
        m = factor_value_columns[i]
        if m.name == "Tissue" or m.name == "Organism part":
            m.name = "Organism part"
        headers.append(f"factor value[{m.name.lower()}]")
        if m.modifiers:
            modifiers = m.modifiers
            if modifiers:
                for mod in modifiers:
                    samples = parse_sample_indices_from_modifier_string(mod["samples"])
                    for s in samples:
                        data[s][last_comment] = mod["value"]
        for j in range(sample_number):
            if data[j][last_comment] == "":
                data[j][last_comment] = m.value
        id_metadata_column_map[m.id] = {"column": last_comment, "name": headers[-1], "type": "factor value", "hidden": m.hidden}
        last_comment += 1
    
    # Final cleanup: ensure all empty/null values are properly converted to SDRF standards
    for i in range(len(data)):
        for j in range(len(data[i])):
            if data[i][j] == "" or data[i][j] is None:
                # Get the column name from headers
                if j < len(headers):
                    column_name = headers[j].lower()
                    # Remove SDRF column type prefixes to get the actual column name
                    if column_name.startswith('characteristics[') and column_name.endswith(']'):
                        column_name = column_name[14:-1]  # Remove 'characteristics[' and ']'
                    elif column_name.startswith('comment[') and column_name.endswith(']'):
                        column_name = column_name[8:-1]  # Remove 'comment[' and ']'
                    elif column_name.startswith('factor value[') and column_name.endswith(']'):
                        column_name = column_name[13:-1]  # Remove 'factor value[' and ']'
                    
                    # Apply proper null value based on whether the column is mandatory
                    if column_name in required_metadata_name or column_name == "tissue" or column_name == "organism part":
                        data[i][j] = "not applicable"
                    else:
                        data[i][j] = "not available"
    
    return [headers, *data], id_metadata_column_map

def parse_sample_indices_from_modifier_string(samples: str):
    """
    Sample indices is a string of comma separated integers and ranges of integers separated by hyphen. This will parse it into a sorted list of integers from lowest to highest
    :param samples:
    :return:
    """
    samples = samples.split(",")
    sample_indices = []
    for sample in samples:
        if "-" in sample:
            start, end = sample.split("-")
            sample_indices.extend(range(int(start)-1, int(end)))
        else:
            sample_indices.append(int(sample)-1)
    return sorted(sample_indices)

def convert_metadata_column_value_to_sdrf(column_name: str, value: str):
    """
    Convert metadata column value to SDRF format
    :param column_name:
    :param value:
    :return:
    """
    # Handle null/empty values according to SDRF standards
    if value == "" or value is None:
        if column_name in required_metadata_name or column_name == "tissue" or column_name == "organism part":
            return "not applicable"
        else:
            return "not available"

    if column_name == "subcellular location":
        if value:
            v = SubcellularLocation.objects.filter(location_identifier=value)
            if v.exists():
                if "AC=" in value:
                    return f"NT={value}"
                else:
                    return f"NT={value};AC={v.first().accession}"
            else:
                return f"NT={value}"
    if column_name == "organism":
        if value:
            v = Species.objects.filter(official_name=value)
            if v.exists():
                #return f"http://purl.obolibrary.org/obo/NCBITaxon_{v.first().taxon}"
                return f"{value}"
            else:
                return f"{value}"
    if column_name == "label":
        if value:
            v = MSUniqueVocabularies.objects.filter(name=value, term_type="sample attribute")
            if v.exists():
                return f"NT={value};AC={v.first().accession}"
            else:
                return f"NT={value}"
    if column_name == "instrument":
        if value:
            v = MSUniqueVocabularies.objects.filter(name=value, term_type="instrument")
            if v.exists():
                if "AC=" in value:
                    return f"NT={value}"
                else:
                    return f"NT={value};AC={v.first().accession}"
            else:
                return f"NT={value}"
    if column_name == "dissociation method":
        if value:
            v = MSUniqueVocabularies.objects.filter(name=value, term_type="dissociation method")
            if v.exists():
                if "AC=" in value:
                    return f"NT={value}"
                else:
                    return f"NT={value};AC={v.first().accession}"
            else:
                return f"{value}"
    if column_name == "cleavage agent details":
        if value:
            v = MSUniqueVocabularies.objects.filter(name=value, term_type="cleavage agent")
            if v.exists():
                if "AC=" in value:
                    return f"NT={value}"
                else:
                    return f"NT={value};AC={v.first().accession}"
            else:
                return f"{value}"
    if column_name == "enrichment process":
        if value:
            v = MSUniqueVocabularies.objects.filter(name=value, term_type="enrichment process")
            if v.exists():
                if "AC=" in value:
                    return f"NT={value}"
                else:
                    return f"AC={v.first().accession};NT={value}"
            else:
                return f"{value}"
    if column_name == "fractionation method":
        if value:
            v = MSUniqueVocabularies.objects.filter(name=value, term_type="fractionation method")
            if v.exists():
                if "AC=" in value:
                    return f"NT={value}"
                else:
                    return f"AC={v.first().accession};NT={value}"
            else:
                return value
    if column_name == "proteomics data acquisition method":
        if value:
            v = MSUniqueVocabularies.objects.filter(name=value, term_type="proteomics data acquisition method")
            if v.exists():
                if "AC=" in value:
                    return f"NT={value}"
                else:
                    return f"AC={v.first().accession};NT={value}"
            else:
                return value
    if column_name == "reduction reagent":
        if value:
            v = MSUniqueVocabularies.objects.filter(name=value, term_type="reduction reagent")
            if v.exists():
                if "AC=" in value:
                    return f"NT={value}"
                else:
                    return f"AC={v.first().accession};NT={value}"
            else:
                return value
    if column_name == "alkylation reagent":
        if value:
            v = MSUniqueVocabularies.objects.filter(name=value, term_type="alkylation reagent")
            if v.exists():
                if "AC=" in value:
                    return f"NT={value}"
                else:
                    return f"AC={v.first().accession};NT={value}"
            else:
                return value
    if column_name == "modification parameters":
        if value:
            v = Unimod.objects.filter(name=value.split(";")[0])
            if v.exists():
                if "AC=" in value or "ac=" in value:
                    return f"NT={value}"
                else:
                    return f"AC={v.first().accession};NT={value}"
            else:
                return value
    if column_name == "ms2 analyzer type":
        if value:
            v = MSUniqueVocabularies.objects.filter(name=value, term_type="mass analyzer type")
            if v.exists():
                if "AC=" in value:
                    return f"NT={value}"
                else:
                    return f"AC={v.first().accession};NT={value}"
            else:
                return value
    return value

def read_sdrf_file(file: str):
    """
    Import SDRF file
    :param file_path:
    :return:
    """

    with open(file, "rt") as f:
        reader = csv.reader(f, delimiter="\t")
        headers = next(reader)
        data = []
        for row in reader:
            data.append(row)
    return headers, data

def convert_sdrf_to_metadata(name: str, value: str):
    data = value.split(";")
    if name == "tissue" or name == "organism part":
        for i in data:
            if "NT=" in i:
                metadata_nt = i.split("=")[1]
                v = Tissue.objects.filter(identifier=metadata_nt)
                if v.exists():
                    return v.first().identifier

        return value
    if name == "subcellular location":
        for i in data:
            if "NT=" in i.upper():
                metadata_nt = i.split("=")[1]
                v = SubcellularLocation.objects.filter(location_identifier=metadata_nt)
                if v.exists():
                    return v.first().location_identifier
            if "AC=" in i.upper():
                metadata_ac = i.split("=")[1]
                v = SubcellularLocation.objects.filter(accession=metadata_ac)
                if v.exists():
                    return v.first().location_identifier
        return value
    if name == "organism":
        for i in data:
            if "http" in i:
                metadata_tx = i.split("_")[1]
                v = Species.objects.filter(taxon=metadata_tx)
                if v.exists():
                    return v.first().official_name
            else:
                return value
    if name == "label":
        for i in data:
            if "NT=" in i:
                metadata_nt = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(name=metadata_nt, term_type="sample attribute")
                if v.exists():
                    return v.first().name
            if "AC=" in i:
                metadata_ac = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(accession=metadata_ac, term_type="sample attribute")
                if v.exists():
                    return v.first().name
        return value
    if name == "instrument":
        for i in data:
            if "NT=" in i:
                metadata_nt = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(name=metadata_nt, term_type="instrument")
                if v.exists():
                    return v.first().name
            if "AC=" in i:
                metadata_ac = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(accession=metadata_ac, term_type="instrument")
                if v.exists():
                    return v.first().name
        return value
    if name == "dissociation method":
        for i in data:
            if "NT=" in i:
                metadata_nt = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(name=metadata_nt, term_type="dissociation method")
                if v.exists():
                    return v.first().name
            if "AC=" in i:
                metadata_ac = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(accession=metadata_ac, term_type="dissociation method")
                if v.exists():
                    return v.first().name
        return value
    if name == "cleavage agent details":
        for i in data:
            if "NT=" in i:
                metadata_nt = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(name=metadata_nt, term_type="cleavage agent")
                if v.exists():
                    return v.first().name
            if "AC=" in i:
                metadata_ac = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(accession=metadata_ac, term_type="cleavage agent")
                if v.exists():
                    return v.first().name
        return value
    if name == "enrichment process":
        for i in data:
            if "NT=" in i:
                metadata_nt = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(name=metadata_nt, term_type="enrichment process")
                if v.exists():
                    return v.first().name
            if "AC=" in i:
                metadata_ac = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(accession=metadata_ac, term_type="enrichment process")
                if v.exists():
                    return v.first().name
        return value
    if name == "fractionation method":
        for i in data:
            if "NT=" in i:
                metadata_nt = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(name=metadata_nt, term_type="fractionation method")
                if v.exists():
                    return v.first().name
            if "AC=" in i:
                metadata_ac = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(accession=metadata_ac, term_type="fractionation method")
                if v.exists():
                    return v.first().name
        return value
    if name == "proteomics data acquisition method":
        for i in data:
            if "NT=" in i:
                metadata_nt = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(name=metadata_nt, term_type="proteomics data acquisition method")
                if v.exists():
                    return v.first().name
            if "AC=" in i:
                metadata_ac = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(accession=metadata_ac, term_type="proteomics data acquisition method")
                if v.exists():
                    return v.first().name
        return value
    if name == "reduction reagent":
        for i in data:
            if "NT=" in i:
                metadata_nt = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(name=metadata_nt, term_type="reduction reagent")
                if v.exists():
                    return v.first().name
            if "AC=" in i:
                metadata_ac = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(accession=metadata_ac, term_type="reduction reagent")
                if v.exists():
                    return v.first().name
        return value

    if name == "alkylation reagent":
        for i in data:
            if "NT=" in i:
                metadata_nt = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(name=metadata_nt, term_type="alkylation reagent")
                if v.exists():
                    return v.first().name
            if "AC=" in i:
                metadata_ac = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(accession=metadata_ac, term_type="alkylation reagent")
                if v.exists():
                    return v.first().name
        return value
    if name == "modification parameters":
        for i in data:
            if "NT=" in i:
                metadata_nt = i.split("=")[1]
                v = Unimod.objects.filter(name=metadata_nt)
                if v.exists():
                    all_data = [metadata_nt]+[d for d in data if "NT=" not in d]
                    return ";".join(all_data)
            if "AC=" in i:
                metadata_ac = i.split("=")[1]
                v = Unimod.objects.filter(accession=metadata_ac)
                if v.exists():
                    all_data = [v.first().name] + [d for d in data if "NT=" not in d]
                    return ";".join(all_data)
        return value
    if name == "ms2 analyzer type":
        for i in data:
            if "NT=" in i:
                metadata_nt = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(name=metadata_nt, term_type="mass analyzer type")
                if v.exists():
                    return v.first().name
            if "AC=" in i:
                metadata_ac = i.split("=")[1]
                v = MSUniqueVocabularies.objects.filter(accession=metadata_ac, term_type="mass analyzer type")
                if v.exists():
                    return v.first().name
        return value
    return value

@job('import-data', timeout='3h')
def _synchronize_pools_with_import_data(instrument_job, import_pools_data, metadata_columns, 
                                      pooled_sample_column_index, source_name_column_index,
                                      user, data_type, user_metadata_field_map, staff_metadata_field_map):
    """
    Synchronize pools with import data:
    - Update existing pools with same name
    - Create new pools from import data  
    - Delete pools that don't exist in import data
    """
    from cc.models import SamplePool, MetadataColumn
    
    # Get existing pools for this instrument job
    existing_pools = SamplePool.objects.filter(instrument_job=instrument_job)
    existing_pool_names = {pool.pool_name: pool for pool in existing_pools}
    
    # Track pools found in import data
    import_pool_names = {pool_data['pool_name'] for pool_data in import_pools_data}
    
    # Process each pool from import data
    for pool_data in import_pools_data:
        pool_name = pool_data['pool_name']
        
        if pool_name in existing_pool_names:
            # Update existing pool
            existing_pool = existing_pool_names[pool_name]
            
            # Update pool properties
            existing_pool.pooled_only_samples = pool_data['pooled_only_samples']
            existing_pool.pooled_and_independent_samples = pool_data['pooled_and_independent_samples']
            existing_pool.is_reference = pool_data['is_reference']
            existing_pool.save()
            
            # Clear existing metadata and recreate
            existing_pool.user_metadata.clear()
            existing_pool.staff_metadata.clear()
            
            # Create new metadata for the updated pool
            _create_pool_metadata_from_import(existing_pool, pool_data, metadata_columns,
                                            pooled_sample_column_index, source_name_column_index,
                                            data_type, user_metadata_field_map, staff_metadata_field_map)
        else:
            # Create new pool
            new_pool = SamplePool.objects.create(
                pool_name=pool_name,
                pooled_only_samples=pool_data['pooled_only_samples'],
                pooled_and_independent_samples=pool_data['pooled_and_independent_samples'],
                is_reference=pool_data['is_reference'],
                instrument_job=instrument_job,
                created_by=user
            )
            
            # Create metadata for the new pool
            _create_pool_metadata_from_import(new_pool, pool_data, metadata_columns,
                                            pooled_sample_column_index, source_name_column_index,
                                            data_type, user_metadata_field_map, staff_metadata_field_map)
    
    # Delete pools that are not in import data
    pools_to_delete = [pool for pool_name, pool in existing_pool_names.items() 
                      if pool_name not in import_pool_names]
    
    for pool in pools_to_delete:
        # This will trigger the pool deletion metadata update logic
        pool.delete()

def _create_pool_metadata_from_import(pool, pool_data, metadata_columns,
                                    pooled_sample_column_index, source_name_column_index,
                                    data_type, user_metadata_field_map, staff_metadata_field_map):
    """Create metadata for a pool from import data"""
    from cc.models import MetadataColumn
    
    row = pool_data['metadata_row']
    pool_name = pool_data['pool_name']
    sn_value = pool_data['sn_value']
    
    for col_index, metadata_column in enumerate(metadata_columns):
        if col_index == pooled_sample_column_index:
            # Create pooled sample column with SN= value
            pool_metadata_column = MetadataColumn.objects.create(
                name=metadata_column.name,
                type=metadata_column.type,
                value=sn_value,
                mandatory=metadata_column.mandatory,
                hidden=metadata_column.hidden,
                readonly=metadata_column.readonly
            )
        elif col_index == source_name_column_index:
            # Create source name column with pool name
            pool_metadata_column = MetadataColumn.objects.create(
                name=metadata_column.name,
                type=metadata_column.type,
                value=pool_name,
                mandatory=metadata_column.mandatory,
                hidden=metadata_column.hidden,
                readonly=metadata_column.readonly
            )
        else:
            # Create other metadata columns from the row
            pool_metadata_column = MetadataColumn.objects.create(
                name=metadata_column.name,
                type=metadata_column.type,
                value=row[col_index] if col_index < len(row) else "",
                mandatory=metadata_column.mandatory,
                hidden=metadata_column.hidden,
                readonly=metadata_column.readonly
            )
        
        # Add to pool's metadata based on data type and field mapping
        if data_type == "user_metadata":
            pool.user_metadata.add(pool_metadata_column)
        elif data_type == "staff_metadata":
            pool.staff_metadata.add(pool_metadata_column)
        else:
            if metadata_column.type in user_metadata_field_map:
                if metadata_column.name in user_metadata_field_map[metadata_column.type]:
                    pool.user_metadata.add(pool_metadata_column)
                else:
                    pool.staff_metadata.add(pool_metadata_column)
            else:
                pool.staff_metadata.add(pool_metadata_column)

def import_sdrf_file(annotation_id: int, user_id: int, instrument_job_id: int, instance_id: str = None, data_type: str = "user_metadata"):
    """
    Import SDRF file
    :param completed_chunk_file_id
    :param user_id:
    :param instrument_job_id:
    :param instance_id:
    :return:
    """
    annotation = Annotation.objects.get(id=annotation_id)
    headers, data = read_sdrf_file(annotation.file.path)
    instrument_job = InstrumentJob.objects.get(id=instrument_job_id)
    metadata_columns = []
    user_metadata_field_map = {}
    for i in user_metadata:
        if i['type'] not in user_metadata_field_map:
            user_metadata_field_map[i['type']] = {}
        user_metadata_field_map[i['type']][i['name']] = i
    staff_metadata_field_map = {}
    for i in staff_metadata:
        if i['type'] not in staff_metadata_field_map:
            staff_metadata_field_map[i['type']] = {}
        staff_metadata_field_map[i['type']][i['name']] = i

    # Check for pooled sample column and identify pool data
    pooled_sample_column_index = None
    pooled_rows = []
    sn_rows = []
    
    for i, header in enumerate(headers):
        header_lower = header.lower()
        if "pooled sample" in header_lower or "pooled_sample" in header_lower:
            pooled_sample_column_index = i
            break
    
    # If we found a pooled sample column, process the data
    if pooled_sample_column_index is not None:
        for row_index, row in enumerate(data):
            pooled_value = row[pooled_sample_column_index].strip()
            if pooled_value.startswith("SN="):
                sn_rows.append(row_index)
            elif pooled_value.lower() == "pooled":
                pooled_rows.append(row_index)
    
    for header in headers:
        metadata_column = MetadataColumn()
        header = header.lower()
        #extract type from pattern <type>[<name>]
        if "[" in header:
            type = header.split("[")[0]
            name = header.split("[")[1].replace("]", "")
        else:
            type = ""
            name = header
        #if name == "organism part":
        #    name = "tissue"
        metadata_column.name = name.capitalize().replace("Ms1", "MS1").replace("Ms2", "MS2")
        metadata_column.type = type.capitalize()
        metadata_columns.append(metadata_column)
    # Remove SN= rows from data before general processing
    if pooled_sample_column_index is not None and sn_rows:
        # Store SN= rows for later pool creation
        sn_data = []
        for row_index in sorted(sn_rows, reverse=True):  # Reverse order to maintain indices
            sn_data.append(data[row_index])
            del data[row_index]
    
    if len(data) != instrument_job.sample_number:
        channel_layer = get_channel_layer()
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}_instrument_job",
            {
                "type": "instrument_job_message",
                "message": {
                    "instance_id": instance_id,
                    "status": "warning",
                    "message": "Number of samples in SDRF file does not match the number of samples in the job"
                },
            }
        )
        # extend the number of samples to match the number of samples in the job and fill with empty strings
        if len(data) < instrument_job.sample_number:
            data.extend(
                [
                    ["" for i in range(len(headers))]
                    for j in range(instrument_job.sample_number - len(data))
                ]
            )

        else:
            data = data[:instrument_job.sample_number]

    if data_type == "user_metadata":
        for m in instrument_job.user_metadata.all():
            m.delete()
        instrument_job.user_metadata.clear()
    elif data_type == "staff_metadata":
        for m in instrument_job.staff_metadata.all():
            m.delete()
        instrument_job.staff_metadata.clear()
    else:
        for m in instrument_job.user_metadata.all():
            m.delete()
        instrument_job.user_metadata.clear()
        for m in instrument_job.staff_metadata.all():
            m.delete()
        instrument_job.staff_metadata.clear()
    for i in range(len(metadata_columns)):
        metadata_value_map = {}
        for j in range(len(data)):
            name = metadata_columns[i].name.lower()
            if data[j][i] == "":
                continue
            if data[j][i] == "not applicable":
                metadata_columns[i].not_applicable = True
                continue
            value = convert_sdrf_to_metadata(name, data[j][i])
            if value not in metadata_value_map:
                metadata_value_map[value] = []
            metadata_value_map[value].append(j)
        # get value with the highest count
        max_count = 0
        max_value = None
        for value in metadata_value_map:
            if len(metadata_value_map[value]) > max_count:
                max_count = len(metadata_value_map[value])
                max_value = value
        if max_value:
            metadata_columns[i].value = max_value
            metadata_columns[i].save()
        # calculate modifiers from the rest of the values
        modifiers = []
        for value in metadata_value_map:
            if value != max_value:
                modifier = {"samples": [], "value": value}
                # sort from lowest to highest. add samples index. for continuous samples, add range
                samples = metadata_value_map[value]
                samples.sort()
                start = samples[0]
                end = samples[0]
                for i2 in range(1, len(samples)):
                    if samples[i2] == end + 1:
                        end = samples[i2]
                    else:
                        if start == end:
                            modifier["samples"].append(str(start+1))
                        else:
                            modifier["samples"].append(f"{start+1}-{end+1}")
                        start = samples[i2]
                        end = samples[i2]
                if start == end:
                    modifier["samples"].append(str(start+1))
                else:
                    modifier["samples"].append(f"{start+1}-{end+1}")
                if len(modifier["samples"]) == 1:
                    modifier["samples"] = modifier["samples"][0]
                else:
                    modifier["samples"] = ",".join(modifier["samples"])
                modifiers.append(modifier)
        if modifiers:
            metadata_columns[i].modifiers = json.dumps(modifiers)
        metadata_columns[i].save()
        if data_type == "user_metadata":
            instrument_job.user_metadata.add(metadata_columns[i])
        elif data_type == "staff_metadata":
            instrument_job.staff_metadata.add(metadata_columns[i])
        else:
            if metadata_columns[i].type in user_metadata_field_map:
                if metadata_columns[i].name in user_metadata_field_map[metadata_columns[i].type]:
                    instrument_job.user_metadata.add(metadata_columns[i])
                else:
                    instrument_job.staff_metadata.add(metadata_columns[i])
            else:
                instrument_job.staff_metadata.add(metadata_columns[i])

    # Create pools from SDRF data if pooled samples were found
    if pooled_sample_column_index is not None:
        from cc.models import SamplePool, User
        user = User.objects.get(id=user_id)
        
        # Get the source name column index
        source_name_column_index = None
        for i, header in enumerate(headers):
            header_lower = header.lower()
            if "source name" in header_lower or "source_name" in header_lower:
                source_name_column_index = i
                break
        
        # Pool synchronization: track pools from import data
        import_pools_data = []
        
        if sn_rows and 'sn_data' in locals():
            # Case 1: There are rows with SN= values - create pools from them
            for pool_index, row in enumerate(sn_data):
                sn_value = row[pooled_sample_column_index].strip()
                
                # Extract source names from SN= value
                if sn_value.startswith("SN="):
                    source_names = sn_value[3:].split(",")
                    source_names = [name.strip() for name in source_names]
                    
                    # Get pool name from source name column or use default
                    pool_name = row[source_name_column_index] if source_name_column_index is not None else f"Pool {pool_index + 1}"
                    
                    # Find sample indices that match these source names
                    pooled_only_samples = []
                    pooled_and_independent_samples = []
                    
                    for sample_index, sample_row in enumerate(data):
                        if source_name_column_index is not None:
                            sample_source_name = sample_row[source_name_column_index].strip()
                            if sample_source_name in source_names:
                                # Check if this sample is also marked as "not pooled" or independent
                                sample_pooled_value = sample_row[pooled_sample_column_index].strip().lower() if pooled_sample_column_index < len(sample_row) else ""
                                
                                if sample_pooled_value == "not pooled" or sample_pooled_value == "" or sample_pooled_value == "independent":
                                    # Sample exists both in pool and as independent
                                    pooled_and_independent_samples.append(sample_index + 1)
                                else:
                                    # Sample is only in pool
                                    pooled_only_samples.append(sample_index + 1)
                    
                    # Store pool data for synchronization
                    import_pools_data.append({
                        'pool_name': pool_name,
                        'pooled_only_samples': pooled_only_samples,
                        'pooled_and_independent_samples': pooled_and_independent_samples,
                        'is_reference': True,  # SN= pools are reference pools
                        'metadata_row': row,
                        'sn_value': sn_value
                    })
                    
        elif pooled_rows:
            # Case 2: No SN= rows but there are "pooled" rows - create a pool from them
            # Get source names of all pooled samples
            pooled_source_names = []
            pooled_only_samples = []
            
            for row_index in pooled_rows:
                if source_name_column_index is not None:
                    source_name = data[row_index][source_name_column_index].strip()
                    if source_name:
                        pooled_source_names.append(source_name)
                        pooled_only_samples.append(row_index + 1)
            
            if pooled_source_names:
                # Create SN= value from source names
                sn_value = "SN=" + ",".join(pooled_source_names)
                pool_name = "Pool 1"
                template_row = data[pooled_rows[0]]
                
                # Store pool data for synchronization
                import_pools_data.append({
                    'pool_name': pool_name,
                    'pooled_only_samples': pooled_only_samples,
                    'pooled_and_independent_samples': [],
                    'is_reference': False,  # Pooled rows are not reference pools by default
                    'metadata_row': template_row,
                    'sn_value': sn_value
                })
        
        # Synchronize pools: update existing, create new, delete missing
        _synchronize_pools_with_import_data(instrument_job, import_pools_data, metadata_columns, 
                                          pooled_sample_column_index, source_name_column_index, 
                                          user, data_type, user_metadata_field_map, staff_metadata_field_map)

    channel_layer = get_channel_layer()
    #notify user through channels that it has completed
    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}_instrument_job",
        {
            "type": "instrument_job_message",
            "message": {
                "instance_id": instance_id,
                "status": "completed",
                "message": "Metadata imported successfully"
            },
        }
    )

@job('import-data', timeout='3h')
def validate_sdrf_file(metadata_column_ids: list[int], sample_number: int, user_id: int, instance_id: str):
    """

    :param metadata_column_ids:
    :param user_id:
    :param instance_id:
    :return:
    """

    metadata_column = MetadataColumn.objects.filter(id__in=metadata_column_ids)
    result, _ = sort_metadata(metadata_column, sample_number)
    print(result)
    # check if there is NoneType in the result
    errors = sdrf_validate(result)
    channel_layer = get_channel_layer()
    if errors:
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}_instrument_job",
            {
                "type": "instrument_job_message",
                "message": {
                    "instance_id": instance_id,
                    "status": "error",
                    "message": "Validation failed",
                    "errors": [str(e) for e in errors]
                },
            }
        )
    else:
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}_instrument_job",
            {
                "type": "instrument_job_message",
                "message": {
                    "instance_id": instance_id,
                    "status": "completed",
                    "message": "Validation successful"
                },
            }
        )


def sdrf_validate(result):
    df = pd.DataFrame()
    try:
        df = SdrfDataFrame.parse(io.StringIO("\n".join(["\t".join(i) for i in result])))
    except TypeError:
        errors = ["Invalid data in the SDRF file"]
    except KeyError:
        errors = ["Missing required columns in the SDRF file"]
    if isinstance(df, SdrfDataFrame):
        try:
            errors = df.validate("default", True)
        except Exception as e:
            errors = [str(e)]
        errors = errors + df.validate("mass_spectrometry", True)
        errors = errors + df.validate_experimental_design()
    return errors


@job('export', timeout='3h')
def export_excel_template(user_id: int, instance_id: str, instrument_job_id: int, export_type: str = "user_metadata"):
    """
    Export excel template
    :param user_id:
    :param instance_id:
    :param instrument_job_id:
    :return:
    """
    instrument_job = InstrumentJob.objects.get(id=instrument_job_id)
    field_mask_map = {}
    if instrument_job.selected_template:
        if instrument_job.selected_template.field_mask_mapping:
            for i in json.loads(instrument_job.selected_template.field_mask_mapping):
                field_mask_map[i["name"]] = i["mask"]
    if export_type=="user_metadata":
        metadata = list(instrument_job.user_metadata.all())
    elif export_type=="staff_metadata":
        metadata = list(instrument_job.staff_metadata.all())
    else:
        metadata = list(instrument_job.user_metadata.all()) + list(instrument_job.staff_metadata.all())

    main_metadata = [m for m in metadata if not m.hidden]
    hidden_metadata = [m for m in metadata if m.hidden]
    result_main, id_map_main = sort_metadata(main_metadata, instrument_job.sample_number)
    result_hidden = []
    id_map_hidden = {}
    if hidden_metadata:
        result_hidden, id_map_hidden = sort_metadata(hidden_metadata, instrument_job.sample_number)
    
    # Check for pools and prepare pool data
    pools = list(SamplePool.objects.filter(instrument_job=instrument_job))
    has_pools = len(pools) > 0
    pool_result_main, pool_id_map_main = [], {}
    pool_result_hidden, pool_id_map_hidden = [], {}
    
    if has_pools:
        # Get all pool metadata (only reference pools for export)
        reference_pools = [pool for pool in pools if pool.is_reference]
        if reference_pools:
            # Combine all pool metadata
            all_pool_metadata = []
            for pool in reference_pools:
                if export_type == "user_metadata":
                    all_pool_metadata.extend(list(pool.user_metadata.all()))
                elif export_type == "staff_metadata":
                    all_pool_metadata.extend(list(pool.staff_metadata.all()))
                else:
                    all_pool_metadata.extend(list(pool.user_metadata.all()) + list(pool.staff_metadata.all()))
            
            # Remove duplicates based on name and type
            unique_pool_metadata = []
            seen_metadata = set()
            for m in all_pool_metadata:
                key = (m.name, m.type)
                if key not in seen_metadata:
                    unique_pool_metadata.append(m)
                    seen_metadata.add(key)
            
            pool_main_metadata = [m for m in unique_pool_metadata if not m.hidden]
            pool_hidden_metadata = [m for m in unique_pool_metadata if m.hidden]
            
            if pool_main_metadata:
                pool_result_main, pool_id_map_main = sort_metadata(pool_main_metadata, len(reference_pools))
            if pool_hidden_metadata:
                pool_result_hidden, pool_id_map_hidden = sort_metadata(pool_hidden_metadata, len(reference_pools))

    # get favourites for each metadata column
    favourites = {}
    user_favourite = FavouriteMetadataOption.objects.filter(user_id=user_id, service_lab_group__isnull=True, lab_group__isnull=True)
    facility_recommended = FavouriteMetadataOption.objects.filter(service_lab_group=instrument_job.service_lab_group)
    global_recommendations = FavouriteMetadataOption.objects.filter(is_global=True)
    for r in list(user_favourite):
        if r.name.lower() not in favourites:
            favourites[r.name.lower()] = []
        favourites[r.name.lower()].append(f"{r.display_value}[*]")
    for r in list(facility_recommended):
        if r.name.lower() not in favourites:
            favourites[r.name.lower()] = []
        favourites[r.name.lower()].append(f"{r.display_value}[**]")
        if r.name.lower() == "tissue" or r.name.lower() == "organism part" or r.name.lower() in required_metadata_name:
            favourites[r.name.lower()].append("not applicable")
    for r in list(global_recommendations):
        if r.name.lower() not in favourites:
            favourites[r.name.lower()] = []
        favourites[r.name.lower()].append(f"{r.display_value}[***]")

    # based on column name from result, contruct an excel file with the appropriate rows beside the header row where the cell with the same name in favourite can have preset selection options dropdown related to that column
    wb = Workbook()
    main_ws = wb.active
    main_ws.title = "main"
    hidden_ws = wb.create_sheet(title="hidden")
    id_metadata_column_map_ws = wb.create_sheet(title="id_metadata_column_map")
    
    # Create pool sheets if pools exist
    pool_main_ws = None
    pool_hidden_ws = None
    pool_id_metadata_column_map_ws = None
    
    if has_pools and reference_pools:
        pool_main_ws = wb.create_sheet(title="pool_main")
        pool_hidden_ws = wb.create_sheet(title="pool_hidden") 
        pool_id_metadata_column_map_ws = wb.create_sheet(title="pool_id_metadata_column_map")
    # fill in the id_metadata_column_map_ws with 3 columns: id, name, type
    id_metadata_column_map_ws.append(["id", "column", "name", "type", "hidden"])

    for k, v in id_map_main.items():
        id_metadata_column_map_ws.append([k, v["column"], v["name"], v["type"], v["hidden"]])
    for k, v in id_map_hidden.items():
        id_metadata_column_map_ws.append([k, v["column"], v["name"], v["type"], v["hidden"]])
    
    # Fill pool ID mapping if pools exist
    if has_pools and reference_pools and pool_id_metadata_column_map_ws:
        pool_id_metadata_column_map_ws.append(["id", "column", "name", "type", "hidden"])
        for k, v in pool_id_map_main.items():
            pool_id_metadata_column_map_ws.append([k, v["column"], v["name"], v["type"], v["hidden"]])
        for k, v in pool_id_map_hidden.items():
            pool_id_metadata_column_map_ws.append([k, v["column"], v["name"], v["type"], v["hidden"]])

    fill = PatternFill(start_color="FFFF99", end_color="FFFF99", fill_type="solid")
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'),
                         bottom=Side(style='thin'))

    # Append headers and data to the main worksheet
    main_ws.append(result_main[0])
    main_work_area = f"A1:{get_column_letter(len(result_main[0]))}{instrument_job.sample_number + 1}"

    for row in result_main[1:]:
        main_ws.append(row)
    for row in main_ws[main_work_area]:
        for cell in row:
            cell.fill = fill
            cell.border = thin_border

    for col in main_ws.columns:
        max_length = 0
        column = col[0].column_letter  # Get the column name
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2)
        main_ws.column_dimensions[column].width = adjusted_width

    note_texts = [
        "Note: Cells that are empty will automatically be filled with 'not applicable' or 'no available' depending on the column when submitted.",
        "[*] User-specific favourite options.",
        "[**] Facility-recommended options.",
        "[***] Global recommendations."
    ]

    start_row = instrument_job.sample_number + 2
    for i, note_text in enumerate(note_texts):
        main_ws.merge_cells(start_row=start_row + i, start_column=1, end_row=start_row + i,
                            end_column=len(result_main[0]))
        note_cell = main_ws.cell(row=start_row + i, column=1)
        note_cell.value = note_text
        note_cell.alignment = Alignment(horizontal='left', vertical='center')

    # Append headers and data to the hidden worksheet
    if len(result_hidden) > 0:
        hidden_work_area = f"A1:{get_column_letter(len(result_hidden[0]))}{instrument_job.sample_number + 1}"
        if len(result_hidden) > 1:
            hidden_ws.append(result_hidden[0])
            for row in result_hidden[1:]:
                hidden_ws.append(row)

            for row in hidden_ws[hidden_work_area]:
                for cell in row:
                    cell.fill = fill
                    cell.border = thin_border

            for col in hidden_ws.columns:
                max_length = 0
                column = col[0].column_letter  # Get the column name
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = (max_length + 2)
                hidden_ws.column_dimensions[column].width = adjusted_width

    # Populate pool worksheets if pools exist
    if has_pools and reference_pools and pool_main_ws and len(pool_result_main) > 0:
        # Pool main worksheet
        pool_main_ws.append(pool_result_main[0])
        pool_main_work_area = f"A1:{get_column_letter(len(pool_result_main[0]))}{len(reference_pools) + 1}"
        
        for row in pool_result_main[1:]:
            pool_main_ws.append(row)
            
        for row in pool_main_ws[pool_main_work_area]:
            for cell in row:
                cell.fill = fill
                cell.border = thin_border
                
        for col in pool_main_ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except:
                    pass
            adjusted_width = (max_length + 2)
            pool_main_ws.column_dimensions[column].width = adjusted_width
            
        # Add notes for pool main worksheet
        pool_note_texts = [
            "Note: Pool metadata for reference pools only.",
            "[*] User-specific favourite options.",
            "[**] Facility-recommended options.", 
            "[***] Global recommendations."
        ]
        
        pool_start_row = len(reference_pools) + 2
        for i, note_text in enumerate(pool_note_texts):
            pool_main_ws.merge_cells(start_row=pool_start_row + i, start_column=1, 
                                   end_row=pool_start_row + i, end_column=len(pool_result_main[0]))
            note_cell = pool_main_ws.cell(row=pool_start_row + i, column=1)
            note_cell.value = note_text
            note_cell.alignment = Alignment(horizontal='left', vertical='center')
            
        # Pool hidden worksheet
        if len(pool_result_hidden) > 0:
            pool_hidden_work_area = f"A1:{get_column_letter(len(pool_result_hidden[0]))}{len(reference_pools) + 1}"
            if len(pool_result_hidden) > 1:
                pool_hidden_ws.append(pool_result_hidden[0])
                for row in pool_result_hidden[1:]:
                    pool_hidden_ws.append(row)
                    
                for row in pool_hidden_ws[pool_hidden_work_area]:
                    for cell in row:
                        cell.fill = fill
                        cell.border = thin_border
                        
                for col in pool_hidden_ws.columns:
                    max_length = 0
                    column = col[0].column_letter
                    for cell in col:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(cell.value)
                        except:
                            pass
                    adjusted_width = (max_length + 2)
                    pool_hidden_ws.column_dimensions[column].width = adjusted_width

    for i, header in enumerate(result_main[0]):
        name_splitted = result_main[0][i].split("[")
        required_column = False
        if len(name_splitted) > 1:
            name = name_splitted[1].replace("]", "")
        else:
            name = name_splitted[0]
        if name in required_metadata_name:
            required_column = True
        name_capitalized = name.capitalize().replace("Ms1", "MS1").replace("Ms2", "MS2")
        if name_capitalized in field_mask_map:
            name = field_mask_map[name_capitalized]
            if len(name_splitted) > 1:
                main_ws.cell(row=1, column=i + 1).value = result_main[0][i].replace(name_splitted[1].rstrip("]"), name.lower())
            else:
                main_ws.cell(row=1, column=i + 1).value =  name.lower()
        option_list = []
        if required_column:
            option_list.append(f"not applicable")
        else:
            option_list.append("not available")

        if name.lower() in favourites:
            option_list = option_list + favourites[name.lower()]
        dv = DataValidation(
            type="list",
            formula1=f'"{",".join(option_list)}"',
            showDropDown=False
        )
        col_letter = get_column_letter(i + 1)
        main_ws.add_data_validation(dv)
        dv.add(f"{col_letter}2:{col_letter}{instrument_job.sample_number + 1}")
    if len(result_hidden) > 1:
        for i, header in enumerate(result_hidden[0]):
            name_splitted = result_hidden[0][i].split("[")
            if len(name_splitted) > 1:
                name = name_splitted[1].replace("]", "")
            else:
                name = name_splitted[0]
            required_column = False
            if name in required_metadata_name:
                required_column = True
            name_capitalized = name.capitalize().replace("Ms1", "MS1").replace("Ms2", "MS2")
            if name_capitalized in field_mask_map:
                name = field_mask_map[name_capitalized]
                if len(name_splitted) > 1:
                    hidden_ws.cell(row=1, column=i + 1).value = result_hidden[0][i].replace(name_splitted[1].rstrip("]"), name.lower())
                else:
                    hidden_ws.cell(row=1, column=i + 1).value = name.lower()

            option_list = []
            if required_column:
                option_list.append(f"not applicable")
            else:
                option_list.append("not available")
            if name.lower() in favourites:
                option_list = option_list + favourites[name.lower()]
            dv = DataValidation(
                type="list",
                formula1=f'"{",".join(option_list)}"',
                showDropDown=False
            )
            col_letter = get_column_letter(i + 1)
            hidden_ws.add_data_validation(dv)
            dv.add(f"{col_letter}2:{col_letter}{instrument_job.sample_number + 1}")

    # Add dropdown validation for pool worksheets
    if has_pools and reference_pools and pool_main_ws and len(pool_result_main) > 0:
        # Pool main worksheet dropdowns
        for i, header in enumerate(pool_result_main[0]):
            name_splitted = pool_result_main[0][i].split("[")
            required_column = False
            if len(name_splitted) > 1:
                name = name_splitted[1].replace("]", "")
            else:
                name = name_splitted[0]
            if name in required_metadata_name:
                required_column = True
            name_capitalized = name.capitalize().replace("Ms1", "MS1").replace("Ms2", "MS2")
            if name_capitalized in field_mask_map:
                name = field_mask_map[name_capitalized]
                if len(name_splitted) > 1:
                    pool_main_ws.cell(row=1, column=i + 1).value = pool_result_main[0][i].replace(name_splitted[1].rstrip("]"), name.lower())
                else:
                    pool_main_ws.cell(row=1, column=i + 1).value = name.lower()
            option_list = []
            if required_column:
                option_list.append(f"not applicable")
            else:
                option_list.append("not available")

            if name.lower() in favourites:
                option_list = option_list + favourites[name.lower()]
            dv = DataValidation(
                type="list",
                formula1=f'"{",".join(option_list)}"',
                showDropDown=False
            )
            col_letter = get_column_letter(i + 1)
            pool_main_ws.add_data_validation(dv)
            dv.add(f"{col_letter}2:{col_letter}{len(reference_pools) + 1}")
            
        # Pool hidden worksheet dropdowns
        if len(pool_result_hidden) > 1:
            for i, header in enumerate(pool_result_hidden[0]):
                name_splitted = pool_result_hidden[0][i].split("[")
                if len(name_splitted) > 1:
                    name = name_splitted[1].replace("]", "")
                else:
                    name = name_splitted[0]
                required_column = False
                if name in required_metadata_name:
                    required_column = True
                name_capitalized = name.capitalize().replace("Ms1", "MS1").replace("Ms2", "MS2")
                if name_capitalized in field_mask_map:
                    name = field_mask_map[name_capitalized]
                    if len(name_splitted) > 1:
                        pool_hidden_ws.cell(row=1, column=i + 1).value = pool_result_hidden[0][i].replace(name_splitted[1].rstrip("]"), name.lower())
                    else:
                        pool_hidden_ws.cell(row=1, column=i + 1).value = name.lower()

                option_list = []
                if required_column:
                    option_list.append(f"not applicable")
                else:
                    option_list.append("not available")
                if name.lower() in favourites:
                    option_list = option_list + favourites[name.lower()]
                dv = DataValidation(
                    type="list",
                    formula1=f'"{",".join(option_list)}"',
                    showDropDown=False
                )
                col_letter = get_column_letter(i + 1)
                pool_hidden_ws.add_data_validation(dv)
                dv.add(f"{col_letter}2:{col_letter}{len(reference_pools) + 1}")

    # save the file
    filename = str(uuid.uuid4())
    xlsx_filepath = os.path.join(settings.MEDIA_ROOT, "temp", f"{filename}.xlsx")

    wb.save(xlsx_filepath)
    signer = TimestampSigner()
    value = signer.sign(f"{filename}.xlsx")

    channel_layer = get_channel_layer()
    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}_instrument_job",
        {
            "type": "download_message",
            "message": {
                "signed_value": value,
                "instance_id": instance_id
            },
        }
    )

def _validate_header_consistency(main_headers, pool_main_headers, hidden_headers=None, pool_hidden_headers=None):
    """
    Validate that pool and main data headers are consistent for Excel import.
    
    Args:
        main_headers: List of headers from main sheet
        pool_main_headers: List of headers from pool_main sheet
        hidden_headers: List of headers from hidden sheet (optional)
        pool_hidden_headers: List of headers from pool_hidden sheet (optional)
    
    Raises:
        ValueError: If headers are inconsistent
    """
    
    # Remove None values and convert to lowercase for comparison
    main_headers_clean = [h.lower().strip() if h else '' for h in main_headers if h is not None]
    pool_main_headers_clean = [h.lower().strip() if h else '' for h in pool_main_headers if h is not None]
    
    # Check main headers consistency
    if main_headers_clean != pool_main_headers_clean:
        # Find specific differences
        main_only = set(main_headers_clean) - set(pool_main_headers_clean)
        pool_only = set(pool_main_headers_clean) - set(main_headers_clean)
        
        error_parts = []
        if main_only:
            error_parts.append(f"Main data has columns not in pool data: {', '.join(sorted(main_only))}")
        if pool_only:
            error_parts.append(f"Pool data has columns not in main data: {', '.join(sorted(pool_only))}")
        
        raise ValueError(f"Header inconsistency between main and pool data. {' | '.join(error_parts)}")
    
    # Check hidden headers consistency if both exist
    if hidden_headers and pool_hidden_headers:
        hidden_headers_clean = [h.lower().strip() if h else '' for h in hidden_headers if h is not None]
        pool_hidden_headers_clean = [h.lower().strip() if h else '' for h in pool_hidden_headers if h is not None]
        
        if hidden_headers_clean != pool_hidden_headers_clean:
            # Find specific differences
            hidden_main_only = set(hidden_headers_clean) - set(pool_hidden_headers_clean)
            hidden_pool_only = set(pool_hidden_headers_clean) - set(hidden_headers_clean)
            
            error_parts = []
            if hidden_main_only:
                error_parts.append(f"Main hidden data has columns not in pool hidden data: {', '.join(sorted(hidden_main_only))}")
            if hidden_pool_only:
                error_parts.append(f"Pool hidden data has columns not in main hidden data: {', '.join(sorted(hidden_pool_only))}")
            
            raise ValueError(f"Header inconsistency between main and pool hidden data. {' | '.join(error_parts)}")

@job('import-data', timeout='3h')
def import_excel(annotation_id: int, user_id: int, instrument_job_id: int, instance_id: str = None, data_type: str = "user_metadata"):
    """
    Import excel file
    :param file:
    :param user_id:
    :param instrument_job_id:
    :param instance_id:
    :return:
    """
    annotation = Annotation.objects.get(id=annotation_id)
    wb = load_workbook(annotation.file.path)
    main_ws = wb["main"]
    main_headers = [cell.value for cell in main_ws[1]]
    main_data = [list(row) for row in main_ws.iter_rows(min_row=2, values_only=True)]
    hidden_ws = None
    hidden_headers = []
    hidden_data = []
    id_metadata_column_map_ws = wb["id_metadata_column_map"]
    id_metadata_column_map_list = [list(row) for row in id_metadata_column_map_ws.iter_rows(min_row=2, values_only=True)]
    id_metadata_column_map = {}
    for row in id_metadata_column_map_list:
        id_metadata_column_map[int(row[0])] = {"column": row[1], "name": row[2], "type": row[3], "hidden": row[4]}
    if "hidden" in wb.sheetnames:
        hidden_ws = wb["hidden"]
        # check if there is any data in the hidden sheet
        if hidden_ws.max_row == 1:
            hidden_ws = None
        else:
            hidden_headers = [cell.value for cell in hidden_ws[1]]
            hidden_data = [list(row) for row in hidden_ws.iter_rows(min_row=2, values_only=True)]

    # Check for pool sheets
    pool_main_ws = None
    pool_hidden_ws = None
    pool_id_metadata_column_map_ws = None
    pool_main_headers = []
    pool_main_data = []
    pool_hidden_headers = []
    pool_hidden_data = []
    pool_id_metadata_column_map_list = []
    pool_id_metadata_column_map = {}
    
    if "pool_main" in wb.sheetnames:
        pool_main_ws = wb["pool_main"]
        if pool_main_ws.max_row > 1:
            pool_main_headers = [cell.value for cell in pool_main_ws[1]]
            pool_main_data = [list(row) for row in pool_main_ws.iter_rows(min_row=2, values_only=True)]
            
    if "pool_hidden" in wb.sheetnames:
        pool_hidden_ws = wb["pool_hidden"]
        if pool_hidden_ws.max_row > 1:
            pool_hidden_headers = [cell.value for cell in pool_hidden_ws[1]]
            pool_hidden_data = [list(row) for row in pool_hidden_ws.iter_rows(min_row=2, values_only=True)]
            
    if "pool_id_metadata_column_map" in wb.sheetnames:
        pool_id_metadata_column_map_ws = wb["pool_id_metadata_column_map"]
        pool_id_metadata_column_map_list = [list(row) for row in pool_id_metadata_column_map_ws.iter_rows(min_row=2, values_only=True)]
        pool_id_metadata_column_map = {}
        for row in pool_id_metadata_column_map_list:
            if row[0] is not None:  # Check for valid ID
                pool_id_metadata_column_map[int(row[0])] = {"column": row[1], "name": row[2], "type": row[3], "hidden": row[4]}

    instrument_job = InstrumentJob.objects.get(id=instrument_job_id)
    
    # Validate header consistency between main and pool data if pool data exists
    try:
        if pool_main_headers and main_headers:
            _validate_header_consistency(main_headers, pool_main_headers, hidden_headers, pool_hidden_headers)
    except ValueError as e:
        # Send error notification to user
        channel_layer = get_channel_layer()
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}_instrument_job",
            {
                "type": "instrument_job_message",
                "message": {
                    "instance_id": instance_id,
                    "status": "error",
                    "message": f"Import failed: {str(e)}"
                },
            }
        )
        return
    
    metadata_columns = []

    user_metadata_field_map = {}
    staff_metadata_field_map = {}
    read_only_metadata_map = {}
    field_mask_map = {}
    if instrument_job.selected_template:
        field_mask_mapping = instrument_job.selected_template.field_mask_mapping
        if field_mask_mapping:
            for i in json.loads(field_mask_mapping):
                field_mask_map[i["mask"]] = i["name"]
        user_columns = list(instrument_job.user_metadata.all())
        staff_columns = list(instrument_job.staff_metadata.all())
        for i in user_columns:
            if i.type not in user_metadata_field_map:
                user_metadata_field_map[i.type] = {}
            if i.name not in user_metadata_field_map[i.type]:
                user_metadata_field_map[i.type][i.name] = []
            um = {"id": i.id, "type": i.type, "name": i.name, "hidden": i.hidden, "value": i.value, "modifiers": i.modifiers}
            user_metadata_field_map[i.type][i.name].append(um)


        for i in staff_columns:
            if i.type not in staff_metadata_field_map:
                staff_metadata_field_map[i.type] = {}
            if i.name not in staff_metadata_field_map[i.type]:
                staff_metadata_field_map[i.type][i.name] = []

            sm = {"id": i.id, "type": i.type, "name": i.name, "hidden": i.hidden, "value": i.value, "modifiers": i.modifiers}
            staff_metadata_field_map[i.type][i.name].append(sm)
    else:
        for i in user_metadata:
            if i['type'] not in user_metadata_field_map:
                user_metadata_field_map[i['type']] = {}
            if i['name'] not in user_metadata_field_map[i['type']]:
                user_metadata_field_map[i['type']][i['name']] = []
            user_metadata_field_map[i['type']][i['name']].append(i)
        for i in staff_metadata:
            if i['type'] not in staff_metadata_field_map:
                staff_metadata_field_map[i['type']] = {}
            if i['name'] not in staff_metadata_field_map[i['type']]:
                staff_metadata_field_map[i['type']][i['name']] = []
            staff_metadata_field_map[i['type']][i['name']].push(i)

    for n, header in enumerate(main_headers):
        id_from_map = 0
        for row in id_metadata_column_map_list:
            row_data = id_metadata_column_map[int(row[0])]
            if row_data["column"] == n and not row_data["hidden"]:
                id_from_map = int(row[0])
                break
        metadata_column = MetadataColumn.objects.filter(id=id_from_map)
        if metadata_column.exists():
            metadata_column = metadata_column.first()
        else:
            metadata_column = MetadataColumn()
            header = header.lower()
            if "[" in header:
                type = header.split("[")[0]
                name = header.split("[")[1].replace("]", "")
            else:
                type = ""
                name = header
            name_capitalized = name.capitalize().replace("Ms1", "MS1").replace("Ms2", "MS2")
            if name_capitalized in field_mask_map:
                name = field_mask_map[name_capitalized]
            metadata_column.name = name
            metadata_column.type = type.capitalize()
            metadata_column.hidden = False
            metadata_column.readonly = False
        metadata_columns.append(metadata_column)

    for n, header in enumerate(hidden_headers):
        id_from_map = 0
        for row in id_metadata_column_map_list:
            row_data = id_metadata_column_map[int(row[0])]
            if row_data["column"] == n and row_data["hidden"]:
                id_from_map = int(row[0])
                break

        if id_from_map != 0:
            metadata_column = MetadataColumn.objects.get(id=id_from_map)
        else:
            metadata_column = MetadataColumn()
            header = header.lower()
            if "[" in header:
                type = header.split("[")[0]
                name = header.split("[")[1].replace("]", "")
            else:
                type = ""
                name = header
            name_capitalized = name.capitalize().replace("Ms1", "MS1").replace("Ms2", "MS2")
            if name_capitalized in field_mask_map:
                name = field_mask_map[name_capitalized]
            metadata_column.name = name
            metadata_column.type = type.capitalize()
            metadata_column.hidden = True
        metadata_columns.append(metadata_column)
    if len(hidden_data) > 0:
        headers = main_headers + hidden_headers
        data = [main_row + hidden_row for main_row, hidden_row in zip(main_data, hidden_data)]
    else:
        headers = main_headers
        data = main_data
    if len(data) != instrument_job.sample_number:
        channel_layer = get_channel_layer()
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}_instrument_job",
            {
                "type": "instrument_job_message",
                "message": {
                    "instance_id": instance_id,
                    "status": "warning",
                    "message": "Number of samples in excel file does not match the number of samples in the job"
                },
            }
        )
        # extend the number of samples to match the number of samples in the job and
        # fill with empty strings

        if len(data) < instrument_job.sample_number:
            data.extend(
                [
                    ["" for i in range(len(headers))]
                    for j in range(instrument_job.sample_number - len(data))
                ]
            )

        else:
            data = data[:instrument_job.sample_number]

    if data_type == "user_metadata":
        instrument_job.user_metadata.clear()
    elif data_type == "staff_metadata":
        instrument_job.staff_metadata.clear()
    else:
        instrument_job.user_metadata.clear()
        instrument_job.staff_metadata.clear()

    for i in range(len(metadata_columns)):
        if metadata_columns[i].readonly:
            # check if the metadata column exists in the read_only_metadata_map then remove the reference from the map
            if data_type == "user_metadata":
                check_and_remove_metadata_from_map(i, metadata_columns, user_metadata_field_map)

            elif data_type == "staff_metadata":
                check_and_remove_metadata_from_map(i, metadata_columns, staff_metadata_field_map)
            else:
                check_and_remove_metadata_from_map(i, metadata_columns, user_metadata_field_map)
                check_and_remove_metadata_from_map(i, metadata_columns, staff_metadata_field_map)
            continue
        metadata_value_map = {}
        for j in range(len(data)):
            name = metadata_columns[i].name.lower()
            if data[j][i] is None:
                data[j][i] = ""
            if data[j][i] == "":
                if name == "tissue" or name in required_metadata_name:
                    data[j][i] = "not applicable"
                else:
                    data[j][i] = "not available"
            if data[j][i] == "not applicable" or data[j][i] == "not available":
                value = data[j][i]
            elif data[j][i].endswith("[*]"):
                value = data[j][i].replace("[*]", "")
                value_query = FavouriteMetadataOption.objects.filter(user_id=user_id, name=name, display_value=value, service_lab_group__isnull=True, lab_group__isnull=True)
                if value_query.exists():
                    value = value_query.first().value
                value = convert_sdrf_to_metadata(name, value)
            elif data[j][i].endswith("[**]"):
                value = data[j][i].replace("[**]", "")
                value_query = FavouriteMetadataOption.objects.filter(name=name, service_lab_group=instrument_job.service_lab_group, display_value=value)
                if value_query.exists():
                    value = value_query.first().value
                value = convert_sdrf_to_metadata(name, value)
            elif data[j][i].endswith("[***]"):
                value = data[j][i].replace("[***]", "")
                value_query = FavouriteMetadataOption.objects.filter(name=name, is_global=True, display_value=value)
                if value_query.exists():
                    value = value_query.first().value
                value = convert_sdrf_to_metadata(name, value)
            else:
                value = convert_sdrf_to_metadata(name, data[j][i])
            if value not in metadata_value_map:
                metadata_value_map[value] = []
            metadata_value_map[value].append(j)
        # get value with the highest count
        max_count = 0
        max_value = None
        for value in metadata_value_map:
            if len(metadata_value_map[value]) > max_count:
                max_count = len(metadata_value_map[value])
                max_value = value
        if max_value:
            metadata_columns[i].value = max_value
            metadata_columns[i].save()
        # calculate modifiers from the rest of the values
        modifiers = []
        for value in metadata_value_map:
            if value != max_value:
                modifier = {"samples": [], "value": value}
                # sort from lowest to highest. add samples index. for continuous samples, add range
                samples = metadata_value_map[value]
                samples.sort()
                start = samples[0]
                end = samples[0]
                for i2 in range(1, len(samples)):
                    if samples[i2] == end + 1:
                        end = samples[i2]
                    else:
                        if start == end:
                            modifier["samples"].append(str(start + 1))
                        else:
                            modifier["samples"].append(f"{start + 1}-{end + 1}")
                        start = samples[i2]
                        end = samples[i2]
                if start == end:
                    modifier["samples"].append(str(start + 1))
                else:
                    modifier["samples"].append(f"{start + 1}-{end + 1}")
                if len(modifier["samples"]) == 1:
                    modifier["samples"] = modifier["samples"][0]
                else:
                    modifier["samples"] = ",".join(modifier["samples"])
                modifiers.append(modifier)
        if modifiers:
            metadata_columns[i].modifiers = json.dumps(modifiers)

        if data_type == "user_metadata":
            if metadata_columns[i].id:
                metadata_columns[i].save()
                check_and_remove_metadata_from_map(i, metadata_columns, user_metadata_field_map)
            else:
                # check if the metadata column exists in the user_metadata_field_map then remove the reference from the map
                check_metadata_column_create_then_remove_from_map(i, instrument_job, metadata_columns,
                                                                  user_metadata_field_map, "user_metadata")
        elif data_type == "staff_metadata":
            if metadata_columns[i].id:
                metadata_columns[i].save()
                check_and_remove_metadata_from_map(i, metadata_columns, staff_metadata_field_map)
            else:
                # check if the metadata column exists in the staff_metadata_field_map then remove the reference from the map
                check_metadata_column_create_then_remove_from_map(i, instrument_job, metadata_columns,
                                                                  staff_metadata_field_map, "staff_metadata")
        else:
            if metadata_columns[i].type in staff_metadata_field_map:
                if metadata_columns[i].name in staff_metadata_field_map[metadata_columns[i].type]:
                    first_column = staff_metadata_field_map[metadata_columns[i].type][metadata_columns[i].name][0]
                    if metadata_columns[i].id:
                        metadata_columns[i].save()
                        instrument_job.staff_metadata.add(metadata_columns[i])
                    else:
                        first_column["value"] = metadata_columns[i].value
                        first_column["modifiers"] = metadata_columns[i].modifiers
                        first_column["hidden"] = metadata_columns[i].hidden
                        first_column["readonly"] = metadata_columns[i].readonly
                        m = MetadataColumn.objects.create(**first_column)
                        instrument_job.staff_metadata.add(m)
                    staff_metadata_field_map[metadata_columns[i].type][metadata_columns[i].name].pop(0)
                    if len(staff_metadata_field_map[metadata_columns[i].type][metadata_columns[i].name]) == 0:
                        del staff_metadata_field_map[metadata_columns[i].type][metadata_columns[i].name]
                else:
                    check_metadata_column_create_then_remove_from_map(i, instrument_job, metadata_columns,
                                                                      user_metadata_field_map, "user_metadata")
            else:
                check_metadata_column_create_then_remove_from_map(i, instrument_job, metadata_columns,
                                                                  user_metadata_field_map, "user_metadata")
    # check if there are any metadata columns left in the user_metadata_field_map and staff_metadata_field_map
    for d_type in user_metadata_field_map:
        for name in user_metadata_field_map[d_type]:
            for i in user_metadata_field_map[d_type][name]:
                i.delete()
    for d_type in staff_metadata_field_map:
        for name in staff_metadata_field_map[d_type]:
            for i in staff_metadata_field_map[d_type][name]:
                i.delete()

    # Import pool data if pool sheets exist
    if pool_main_headers and pool_main_data:
        _import_pool_data_from_excel(instrument_job, pool_main_headers, pool_main_data, 
                                   pool_hidden_headers, pool_hidden_data,
                                   pool_id_metadata_column_map, data_type, 
                                   field_mask_map, user_id)

    channel_layer = get_channel_layer()
    # notify user through channels that it has completed
    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}_instrument_job",
        {
            "type": "instrument_job_message",
            "message": {
                "instance_id": instance_id,
                "status": "completed",
                "message": "Metadata imported successfully"
            },
        }
    )


def _import_pool_data_from_excel(instrument_job, pool_main_headers, pool_main_data,
                               pool_hidden_headers, pool_hidden_data,
                               pool_id_metadata_column_map, data_type, 
                               field_mask_map, user_id):
    """
    Import pool data from Excel sheets
    """
    from cc.models import SamplePool, MetadataColumn, User, FavouriteMetadataOption
    
    user = User.objects.get(id=user_id)
    
    # Combine pool headers and data
    pool_headers = pool_main_headers + pool_hidden_headers
    if len(pool_hidden_data) > 0:
        pool_data = [main_row + hidden_row for main_row, hidden_row in zip(pool_main_data, pool_hidden_data)]
    else:
        pool_data = pool_main_data
    
    # Build pool metadata columns from headers and ID mapping
    pool_metadata_columns = []
    
    for n, header in enumerate(pool_headers):
        id_from_map = 0
        # Find ID from mapping
        for map_id, row_data in pool_id_metadata_column_map.items():
            if row_data["column"] == n:
                id_from_map = map_id
                break
                
        # Try to get existing metadata column or create new one
        metadata_column = MetadataColumn.objects.filter(id=id_from_map).first()
        if not metadata_column:
            metadata_column = MetadataColumn()
            header = header.lower() if header else ""
            if "[" in header:
                type_part = header.split("[")[0]
                name = header.split("[")[1].replace("]", "")
            else:
                type_part = ""
                name = header
            
            name_capitalized = name.capitalize().replace("Ms1", "MS1").replace("Ms2", "MS2")
            if name_capitalized in field_mask_map:
                name = field_mask_map[name_capitalized]
                
            metadata_column.name = name
            metadata_column.type = type_part.capitalize()
            metadata_column.hidden = n >= len(pool_main_headers)  # Hidden if in hidden section
            metadata_column.readonly = False
            
        pool_metadata_columns.append(metadata_column)
    
    # Extract pool information and create/update pools
    pools_to_import = []
    
    for row_index, row in enumerate(pool_data):
        pool_info = {
            'pool_name': None,
            'pooled_only_samples': [],
            'pooled_and_independent_samples': [],
            'is_reference': True,  # Pool sheets only contain reference pools
            'metadata': {}
        }
        
        # Extract metadata from row
        for col_index, value in enumerate(row):
            if col_index < len(pool_metadata_columns):
                column = pool_metadata_columns[col_index]
                
                # Handle special columns
                if column.name.lower() == 'source name':
                    pool_info['pool_name'] = value if value else f"Pool {row_index + 1}"
                elif column.name.lower() == 'pooled sample' and value and value.startswith('SN='):
                    # Extract sample information from SN= value
                    source_names = value[3:].split(',')
                    source_names = [name.strip() for name in source_names if name.strip()]
                    
                    # Map source names to sample indices (this requires finding samples by source name)
                    for source_name in source_names:
                        sample_index = _find_sample_index_by_source_name(instrument_job, source_name)
                        if sample_index:
                            pool_info['pooled_only_samples'].append(sample_index)
                
                # Store metadata value
                if value is not None and value != "":
                    # Process favourite values like main import
                    if isinstance(value, str):
                        if value.endswith("[*]"):
                            processed_value = value.replace("[*]", "")
                            value_query = FavouriteMetadataOption.objects.filter(
                                user_id=user_id, name=column.name.lower(), 
                                display_value=processed_value, service_lab_group__isnull=True, 
                                lab_group__isnull=True
                            )
                            if value_query.exists():
                                processed_value = value_query.first().value
                            pool_info['metadata'][col_index] = processed_value
                        elif value.endswith("[**]"):
                            processed_value = value.replace("[**]", "")
                            value_query = FavouriteMetadataOption.objects.filter(
                                name=column.name.lower(), 
                                service_lab_group=instrument_job.service_lab_group, 
                                display_value=processed_value
                            )
                            if value_query.exists():
                                processed_value = value_query.first().value
                            pool_info['metadata'][col_index] = processed_value
                        elif value.endswith("[***]"):
                            processed_value = value.replace("[***]", "")
                            value_query = FavouriteMetadataOption.objects.filter(
                                name=column.name.lower(), is_global=True, 
                                display_value=processed_value
                            )
                            if value_query.exists():
                                processed_value = value_query.first().value
                            pool_info['metadata'][col_index] = processed_value
                        else:
                            pool_info['metadata'][col_index] = value
                    else:
                        pool_info['metadata'][col_index] = value
        
        # Set default pool name if not found
        if not pool_info['pool_name']:
            pool_info['pool_name'] = f"Pool {row_index + 1}"
            
        pools_to_import.append(pool_info)
    
    # Use the existing synchronization logic to update pools
    import_pools_data = []
    for pool_info in pools_to_import:
        import_pools_data.append({
            'pool_name': pool_info['pool_name'],
            'pooled_only_samples': pool_info['pooled_only_samples'],
            'pooled_and_independent_samples': pool_info['pooled_and_independent_samples'],
            'is_reference': pool_info['is_reference'],
            'metadata_row': [pool_info['metadata'].get(i, "") for i in range(len(pool_metadata_columns))],
            'sn_value': _generate_sn_value_from_samples(instrument_job, pool_info['pooled_only_samples'])
        })
    
    # Use existing synchronization function
    _synchronize_pools_with_import_data(instrument_job, import_pools_data, pool_metadata_columns,
                                       None, None, user, data_type, {}, {})

def _find_sample_index_by_source_name(instrument_job, source_name):
    """Find sample index by source name from instrument job metadata"""
    # This would need to look through the instrument job's metadata to find the sample
    # For now, we'll implement a basic version that extracts from source name if it contains sample number
    import re
    
    # Try to extract sample number from source name
    match = re.search(r'sample[_\s]*(\d+)', source_name.lower())
    if match:
        sample_index = int(match.group(1))
        if 1 <= sample_index <= instrument_job.sample_number:
            return sample_index
    
    # Fallback: try to find by checking source name column in metadata
    source_name_columns = instrument_job.user_metadata.filter(name__icontains='source').first() or \
                         instrument_job.staff_metadata.filter(name__icontains='source').first()
    
    if source_name_columns and source_name_columns.modifiers:
        import json
        modifiers = json.loads(source_name_columns.modifiers) if isinstance(source_name_columns.modifiers, str) else source_name_columns.modifiers
        for modifier in modifiers:
            if modifier.get('value') == source_name:
                # Parse sample range from modifier
                samples_str = modifier.get('samples', '')
                if samples_str:
                    # Parse ranges like "1-3,5"
                    sample_indices = []
                    for part in samples_str.split(','):
                        if '-' in part:
                            start, end = part.split('-')
                            sample_indices.extend(range(int(start), int(end) + 1))
                        else:
                            sample_indices.append(int(part))
                    return sample_indices[0] if sample_indices else None
    
    return None

def _generate_sn_value_from_samples(instrument_job, sample_indices):
    """Generate SN= value from sample indices"""
    if not sample_indices:
        return "not pooled"
    
    # Get source names for the samples
    source_names = []
    for sample_index in sample_indices:
        # Try to get actual source name from metadata
        source_name = f"Sample {sample_index}"  # Default fallback
        
        # Look for source name in metadata
        source_name_column = instrument_job.user_metadata.filter(name__icontains='source').first() or \
                           instrument_job.staff_metadata.filter(name__icontains='source').first()
        
        if source_name_column:
            # Check if there are modifiers for this sample
            if source_name_column.modifiers:
                import json
                modifiers = json.loads(source_name_column.modifiers) if isinstance(source_name_column.modifiers, str) else source_name_column.modifiers
                for modifier in modifiers:
                    samples_str = modifier.get('samples', '')
                    if str(sample_index) in samples_str.split(','):
                        source_name = modifier.get('value', source_name)
                        break
            else:
                # Use default value if no modifiers
                source_name = source_name_column.value or source_name
        
        source_names.append(source_name)
    
    return f"SN={','.join(source_names)}"

def check_metadata_column_create_then_remove_from_map(i, instrument_job, metadata_columns, metadata_field_map, field_type):
    if metadata_columns[i].type in metadata_field_map:
        if metadata_columns[i].name in metadata_field_map[metadata_columns[i].type]:
            first_column = metadata_field_map[metadata_columns[i].type][metadata_columns[i].name][0]
            if metadata_columns[i].id:
                metadata_columns[i].save()
                if field_type == "user_metadata":
                    instrument_job.user_metadata.add(metadata_columns[i])
                else:
                    instrument_job.staff_metadata.add(metadata_columns[i])
            else:
                first_column["value"] = metadata_columns[i].value
                first_column["modifiers"] = metadata_columns[i].modifiers
                first_column["hidden"] = metadata_columns[i].hidden
                first_column["readonly"] = metadata_columns[i].readonly
                m = MetadataColumn.objects.create(**first_column)
                if field_type == "user_metadata":
                    instrument_job.user_metadata.add(m)
                else:
                    instrument_job.staff_metadata.add(m)
            metadata_field_map[metadata_columns[i].type][metadata_columns[i].name].pop(0)
            if len(metadata_field_map[metadata_columns[i].type][metadata_columns[i].name]):
                del metadata_field_map[metadata_columns[i].type][metadata_columns[i].name]
        else:
            metadata_columns[i].save()
            if field_type == "user_metadata":
                instrument_job.user_metadata.add(metadata_columns[i])
            else:
                instrument_job.staff_metadata.add(metadata_columns[i])
    else:
        metadata_columns[i].save()
        if field_type == "user_metadata":
            instrument_job.user_metadata.add(metadata_columns[i])
        else:
            instrument_job.user_metadata.add(metadata_columns[i])


def check_and_remove_metadata_from_map(i, metadata_columns, metadata_field_map):
    if metadata_columns[i].type in metadata_field_map:
        if metadata_columns[i].name in metadata_field_map[metadata_columns[i].type]:
            first_column = metadata_field_map[metadata_columns[i].type][metadata_columns[i].name][0]
            metadata_field_map[metadata_columns[i].type][metadata_columns[i].name].pop(0)
            if not metadata_field_map[metadata_columns[i].type][metadata_columns[i].name]:
                del metadata_field_map[metadata_columns[i].type][metadata_columns[i].name]

@job('export', timeout='3h')
def export_instrument_usage(instrument_ids: list[int], lab_group_ids: list[int], user_ids: list[int], mode: str, instance_id: str, time_started: str = None, time_ended: str = None, calculate_duration_with_cutoff: bool = False, user_id: int = 0, file_format: str = "xlsx", includes_maintenance: bool = False, approved_only: bool = True):
    instrument_usages = InstrumentUsage.objects.filter(instrument__id__in=instrument_ids)
    channel_layer = get_channel_layer()
    if mode == 'service_lab_group':
        lab_group = LabGroup.objects.filter(id__in=lab_group_ids, can_perform_ms_analysis=True)
        if lab_group.exists():
            users = User.objects.filter(lab_groups__in=lab_group)
            instrument_jobs = InstrumentJob.objects.filter(service_lab_group__in=lab_group)
            instrument_usages = instrument_usages.filter(
                Q(user__in=users)|Q(annotation__instrument_jobs__in=instrument_jobs)
            )

        else:
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}_instrument_job",
                {
                    "type": "download_message",
                    "message": {
                        "instance_id": instance_id,
                        "status": "error",
                        "message": "Lab group not found"
                    },
                }
            )
            return
    elif mode == 'lab_group':
        lab_group = LabGroup.objects.filter(lab_group__id__in=lab_group_ids, can_perform_ms_analysis=False)
        if lab_group.exists():
            users = User.objects.filter(lab_group__in=lab_group)
            instrument_usages = instrument_usages.filter(user__in=users)
        else:
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}_instrument_job",
                {
                    "type": "download_message",
                    "message": {
                        "instance_id": instance_id,
                        "status": "error",
                        "message": "Lab group not found"
                    },
                }
            )
            return
    else:
        if user_id != 0:
            instrument_usages = instrument_usages.filter(user__id__in=user_ids)

    instrument_usages = instrument_usages.filter(
        (Q(time_started__range=[time_started, time_ended]) |
        Q(time_ended__range=[time_started, time_ended]))
    )
    if not includes_maintenance:
        instrument_usages = instrument_usages.exclude(maintenance=True)
    if approved_only:
        instrument_usages = instrument_usages.filter(approved=True)

    if instrument_usages.exists():
        filename = str(uuid.uuid4())
        if time_started:
            time_started = timezone.make_aware(datetime.datetime.strptime(time_started, "%Y-%m-%dT%H:%M:%S.%fZ"))
        if time_ended:
            time_ended = timezone.make_aware(datetime.datetime.strptime(time_ended, "%Y-%m-%dT%H:%M:%S.%fZ"))
        if file_format == "xlsx":
            wb = Workbook()
            ws = wb.active
            ws.title = "instrument_usage"
            filename = filename + ".xlsx"
            filepath = os.path.join(settings.MEDIA_ROOT, "temp", filename)
        elif file_format == "csv" or file_format == "tsv":
            filename = filename + f".{file_format}"
            filepath = os.path.join(settings.MEDIA_ROOT, "temp", filename)
            infile = open(filepath, "wt", encoding="utf-8", newline="")
            if file_format == "csv":
                writer = csv.writer(infile)
            else:
                writer = csv.writer(infile, delimiter="\t")

        headers = ["Instrument", "User", "Time Started", "Time Ended", "Duration", "Description", "Associated Jobs", "Is Maintenance", "Is Approved"]
        if file_format == "xlsx":
            ws.append(headers)
        elif file_format == "csv" or file_format == "tsv":
            writer.writerow(headers)
        for i in instrument_usages:
            # check if instrument_usage time_started and time_ended are not splitted by the time_started and time_ended of the function
            duration = i.time_ended - i.time_started
            if calculate_duration_with_cutoff:
                if time_ended:
                    if i.time_ended > time_ended:
                        # calculate the duration within the boundary
                        if time_started:
                            if i.time_started <= time_started:
                                duration += time_ended - time_started
                            else:
                                duration += time_ended - i.time_started
                        else:
                            duration += time_ended - i.time_started
                    else:
                        if time_started:
                            if i.time_started <= time_started:
                                duration += i.time_ended - time_started
                            else:
                                duration += i.time_ended - i.time_started
                        else:
                            duration += i.time_ended - i.time_started
                else:
                    if time_started:
                        if i.time_started <= time_started:
                            duration += i.time_ended - time_started
                        else:
                            duration += i.time_ended - i.time_started
            associated_jobs_information = []
            if i.annotation:
                associated_jobs = i.annotation.assigned_instrument_jobs.all()
                for j in associated_jobs:
                    associated_jobs_information.append(f"{j.submitted_at} {j.job_name} ({j.user.username})")

            # convert time to string for excel
            exported_time_started = i.time_started.strftime("%Y-%m-%d")
            exported_time_ended = i.time_ended.strftime("%Y-%m-%d")
            exported_duration = duration.days+1
            if file_format == "xlsx":
                ws.append([i.instrument.instrument_name, i.user.username, exported_time_started, exported_time_ended, exported_duration, i.description, ";\n".join(associated_jobs_information), i.maintenance, i.approved])
            if file_format == "csv" or file_format == "tsv":
                writer.writerow([i.instrument.instrument_name, i.user.username, exported_time_started, exported_time_ended, exported_duration, i.description, ";".join(associated_jobs_information), i.maintenance, i.approved])
        if file_format == "xlsx":
            for col in ws.columns:
                max_length = 0
                column = col[0].column_letter
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[column].width = adjusted_width

        if file_format == "xlsx":
            wb.save(filepath)
        signer = TimestampSigner()
        value = signer.sign(filename)
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}_instrument_job",
            {
                "type": "download_message",
                "message": {
                    "signed_value": value,
                    "instance_id": instance_id
                },
            }
        )
    else:
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}_instrument_job",
            {
                "type": "download_message",
                "message": {
                    "instance_id": instance_id,
                    "status": "error",
                    "message": "No instrument usage found"
                },
            }
        )

@job('export', timeout='3h')
def export_reagent_actions(start_date=None, end_date=None, storage_object_id=None, stored_reagent_ids=None,
                          user_id=None, export_format='csv', instance_id=None):
    """
    Export reagent actions within a specified time period with optional storage location filtering

    Args:
        start_date (datetime): Start date for filtering actions
        end_date (datetime): End date for filtering actions
        storage_object_id (int): Optional ID of storage object to filter by
        user_id (int): Optional user ID who requested the export
        export_format (str): Export format (currently only 'csv' supported)
        instance_id (str): Optional instance ID for tracking the job

    Returns:
        str: Signed filename for secure download
    """
    if not end_date:
        end_date = timezone.now()
    if not start_date:
        start_date = end_date - timedelta(days=30)

    actions_query = ReagentAction.objects.filter(
        created_at__gte=start_date,
        created_at__lte=end_date
    ).select_related('reagent', 'reagent__storage_object', 'user')

    if storage_object_id:
        try:
            storage_object = StorageObject.objects.get(id=storage_object_id)

            all_storage_objects = [storage_object]
            children = storage_object.get_all_children()
            all_storage_objects.extend(children)
            storage_ids = [obj.id for obj in all_storage_objects]

            actions_query = actions_query.filter(reagent__storage_object_id__in=storage_ids)
        except StorageObject.DoesNotExist:
            pass

    if stored_reagent_ids:
        actions_query = actions_query.filter(reagent__id__in=stored_reagent_ids)

    filename = f"reagent_actions_export_{uuid.uuid4().hex}.csv"
    export_dir = os.path.join(settings.MEDIA_ROOT, 'temp')

    os.makedirs(export_dir, exist_ok=True)
    filepath = os.path.join(export_dir, filename)

    with open(filepath, 'w', newline='') as csvfile:
        writer = csv.writer(csvfile)

        writer.writerow([
            'Date', 'Action Type', 'Item Name', 'Barcode', 'Barcode Format', 'Quantity',
            'Storage Location', 'Storage Path', 'User', 'Notes'
        ])

        for action in actions_query:
            storage_path = ""
            storage_obj = action.reagent.storage_object
            if storage_obj:
                path = storage_obj.get_path_to_root()
                storage_path = " > ".join([item["name"] for item in path])
            barcode_format = identify_barcode_format(action.reagent.barcode)
            writer.writerow([
                action.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                action.action_type,
                action.reagent.reagent.name,
                action.reagent.barcode or "",
                barcode_format,
                action.quantity if action.quantity else "",
                action.reagent.storage_object.object_name if action.reagent.storage_object else "",
                storage_path,
                action.user.username if action.user else "",
                action.notes or ""
            ])

    signer = TimestampSigner()
    signed_filename = signer.sign(filename)

    if user_id:
        channel_layer = get_channel_layer()
        async_to_sync(channel_layer.group_send)(
            f"user_{user_id}",
            {
                "type": "download_message",
                "message": {
                    "signed_value": signed_filename,
                    "instance_id": instance_id
                },
            }
        )


@job('import-data', timeout='3h')
def import_reagents_from_file(file_path: str, storage_object_id: int, user_id: int, column_mapping=None,
                     instance_id: str = None):
    """
    Import reagents from a file into the system. The file should be in a format that can be processed
    :param file_path:
    :param storage_object_id:
    :param user_id:
    :param column_mapping:
    :param create_missing_reagents:
    :param instance_id:
    :return:
    """
    channel_layer = get_channel_layer()
    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}",
        {
            "type": "import_message",
            "message": {
                "instance_id": instance_id,
                "status": "started",
                "message": "Importing reagents from file"
            },
        }
    )
    default_mapping = {
        'item_name': 'name',
        'unit': 'unit',
        'quantity': 'quantity',
        'notes': 'notes',
        'barcode': 'barcode',
        'expiration_date': 'expiration_date'
    }

    user = User.objects.get(id=user_id)

    if column_mapping is None:
        column_mapping = default_mapping

    try:
        storage_object = StorageObject.objects.get(id=storage_object_id)
    except StorageObject.DoesNotExist:
        raise ValidationError("Storage object not found")

    file_name = os.path.basename(file_path)
    if file_name.endswith('.csv'):
        df = pd.read_csv(file_path)
    if file_name.endswith('.txt'):
        df = pd.read_csv(file_path, sep="\t")
    elif file_name.endswith(('.xls', '.xlsx')):
        df = pd.read_excel(file_path)
    else:
        raise ValidationError("Unsupported file type. Please provide CSV or Excel file.")

    required_columns = ['item_name', 'unit', 'quantity']
    missing_columns = [col for col in required_columns if col not in df.columns]
    if missing_columns:
        raise ValidationError(f"Missing required columns: {', '.join(missing_columns)}")

    results = {
        'success_count': 0,
        'error_count': 0,
        'errors': []
    }

    allowed_units = ["ng", "ug", "mg", "g", "nL", "uL", "mL", "L", "uM", "mM", "nM", "M", "ea", "other"]

    with transaction.atomic():
        for index, row in df.iterrows():
            try:
                reagent_name = row['item_name'].strip()
                unit = row['unit'].strip()
                # check if unit is within the allowed units
                if unit not in allowed_units:
                    results["error_count"] += 1
                    results["errors"].append(f"Row {index + 1}: Invalid unit '{unit}'. Allowed units are: {', '.join(allowed_units)}")
                    continue

                reagent, created = Reagent.objects.get_or_create(
                    name=reagent_name,
                    unit=unit
                )

                stored_reagent = StoredReagent(
                    reagent=reagent,
                    storage_object=storage_object,
                    quantity=float(row['quantity']),
                    user=user,
                    created_at=timezone.now()
                )

                for file_col, model_field in column_mapping.items():
                    if file_col in df.columns and file_col not in required_columns:
                        if row[file_col] is not None and not pd.isna(row[file_col]):
                            if model_field == 'expiration_date' and row[file_col]:
                                try:
                                    setattr(stored_reagent, model_field, pd.to_datetime(row[file_col]).date())
                                except Exception as e:
                                    results['errors'].append(f"Row {index + 1}: Invalid date format: {str(e)}")
                                    continue
                            else:
                                setattr(stored_reagent, model_field, row[file_col])

                stored_reagent.save()
                stored_reagent.create_default_folders()

                results['success_count'] += 1

            except Exception as e:
                results['error_count'] += 1
                results['errors'].append(f"Row {index + 1}: {str(e)}")
    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}",
        {
            "type": "import_message",
            "message": {
                "instance_id": instance_id,
                "status": "completed",
                "message": f"Import completed. {results['success_count']} rows imported successfully, {results['error_count']} errors.",
                "errors": results['errors']
            },
        }
    )


@job('maintenance', timeout='1h')
def check_instrument_warranty_maintenance(instrument_ids: list[int], days_before_warranty_warning=30, days_before_maintenance_warning=15,
                                          user_id: int = None, instance_id: str = None, send_email: bool = True):
    if not instrument_ids:
        instruments = Instrument.objects.all()
    else:
        instruments = Instrument.objects.filter(id__in=instrument_ids)
    if not instruments.exists():
        return {"status": "error", "message": "No instruments found for the provided IDs."}
    channel_layer = get_channel_layer()
    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}",
        {
            "type": "maintenance_message",
            "message": {
                "instance_id": instance_id,
                "status": "started",
                "message": "Checking instrument warranty and maintenance status"
            },
        }
    )

    task_ran_count ={
        "warranty": 0,
        "maintenance": 0
    }

    for i in instruments:
        if i.check_upcoming_maintenance(days_before_maintenance_warning):
            task_ran_count["warranty"] += 1
        if i.check_warranty_expiration(days_before_warranty_warning):
            task_ran_count["maintenance"] += 1

    async_to_sync(channel_layer.group_send)(
        f"user_{user_id}",
        {
            "type": "maintenance_message",
            "message": {
                "instance_id": instance_id,
                "status": "completed",
                "message": f"Instrument warranty and maintenance check completed with {task_ran_count['warranty']} warranty warnings and {task_ran_count['maintenance']} maintenance warnings.",
            },
        }
    )


def _get_file_size_mb(file_path: str) -> float:
    """
    Get file size in megabytes.
    
    Args:
        file_path: Path to the file
        
    Returns:
        File size in MB, rounded to 2 decimal places
    """
    try:
        size_bytes = os.path.getsize(file_path)
        size_mb = size_bytes / (1024 * 1024)
        return round(size_mb, 2)
    except Exception:
        return 0.0


def _cleanup_old_exports(temp_dir: str, max_age_days: int = 7):
    """
    Clean up old export files from the temp directory to prevent storage buildup.
    
    Args:
        temp_dir: Path to the temp directory containing export files
        max_age_days: Maximum age in days before files are deleted (default: 7)
    """

    
    try:
        current_time = time.time()
        max_age_seconds = max_age_days * 24 * 60 * 60
        
        for item in Path(temp_dir).iterdir():
            if item.is_file() and (item.name.startswith('cupcake_export_') or 
                                 item.name.endswith(('.zip', '.tar.gz', '.sha256'))):
                file_age = current_time - item.stat().st_mtime
                if file_age > max_age_seconds:
                    try:
                        item.unlink()
                        print(f"Cleaned up old export file: {item.name}")
                    except Exception as e:
                        print(f"Failed to delete old export file {item.name}: {e}")
            elif item.is_dir() and item.name.startswith('cupcake_export_'):
                dir_age = current_time - item.stat().st_mtime
                if dir_age > max_age_seconds:
                    try:
                        shutil.rmtree(item)
                        print(f"Cleaned up old export directory: {item.name}")
                    except Exception as e:
                        print(f"Failed to delete old export directory {item.name}: {e}")
                        
    except Exception as e:
        print(f"Error during export cleanup: {e}")
        # Don't raise the exception to avoid breaking the export process


@job('mcp', timeout='30m')
def analyze_protocol_step_task(task_id: str, step_id: int, use_anthropic: bool = False, user_id: int = None):
    """
    Analyze a protocol step for SDRF suggestions asynchronously.
    
    Args:
        task_id: Unique task identifier for WebSocket updates
        step_id: Protocol step ID to analyze
        use_anthropic: Whether to use Anthropic Claude analysis
        user_id: User ID for WebSocket group
    """

    
    channel_layer = get_channel_layer()
    
    def send_progress(status: str, message: str, progress: int = 0, data: dict = None):
        """Send progress update via WebSocket"""
        if user_id:
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}_mcp_analysis",
                {
                    "type": "mcp_analysis_update",
                    "message": {
                        "task_id": task_id,
                        "step_id": step_id,
                        "status": status,
                        "message": message,
                        "progress": progress,
                        "data": data or {}
                    }
                }
            )
    
    try:
        # Send initial progress
        send_progress("started", "Starting protocol step analysis...", 10)
        
        # Get the protocol step
        try:
            step = ProtocolStep.objects.get(id=step_id)
        except ProtocolStep.DoesNotExist:
            send_progress("error", f"Protocol step {step_id} not found", 0)
            return {"success": False, "error": f"Protocol step {step_id} not found"}
        
        send_progress("processing", "Checking cache...", 20)
        
        # Check cache first
        analyzer_type = 'mcp_claude' if use_anthropic else 'standard_nlp'
        cached_result = ProtocolStepSuggestionCache.get_cached_suggestions(step_id, analyzer_type)
        
        if cached_result:
            send_progress("completed", "Analysis completed (cached)", 100, {
                "cached": True,
                "result": cached_result
            })
            return cached_result
        
        send_progress("processing", "Initializing analyzer...", 30)
        
        # Initialize analyzer
        analyzer = ProtocolAnalyzer(use_anthropic=use_anthropic)
        
        send_progress("processing", "Analyzing step content...", 50)
        
        # Perform analysis
        result = analyzer.analyze_protocol_step(step.id)
        
        send_progress("processing", "Caching results...", 80)
        
        # Cache the results
        if result.get('success'):
            ProtocolStepSuggestionCache.cache_suggestions(step_id, analyzer_type, result)
        
        send_progress("completed", "Analysis completed successfully", 100, {
            "cached": False,
            "result": result
        })
        
        return result
        
    except Exception as e:
        error_msg = f"Analysis failed: {str(e)}"
        send_progress("error", error_msg, 0)
        return {"success": False, "error": error_msg}


@job('mcp', timeout='60m')
def analyze_full_protocol_task(task_id: str, protocol_id: int, use_anthropic: bool = False, user_id: int = None):
    """
    Analyze all steps in a protocol for SDRF suggestions asynchronously.
    
    Args:
        task_id: Unique task identifier for WebSocket updates
        protocol_id: Protocol ID to analyze
        use_anthropic: Whether to use Anthropic Claude analysis
        user_id: User ID for WebSocket group
    """
    
    channel_layer = get_channel_layer()
    
    def send_progress(status: str, message: str, progress: int = 0, data: dict = None):
        """Send progress update via WebSocket"""
        if user_id:
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}_mcp_analysis",
                {
                    "type": "mcp_analysis_update",
                    "message": {
                        "task_id": task_id,
                        "protocol_id": protocol_id,
                        "status": status,
                        "message": message,
                        "progress": progress,
                        "data": data or {}
                    }
                }
            )
    
    try:
        send_progress("started", "Starting full protocol analysis...", 5)
        
        # Get the protocol
        try:
            protocol = ProtocolModel.objects.get(id=protocol_id)
        except ProtocolModel.DoesNotExist:
            send_progress("error", f"Protocol {protocol_id} not found", 0)
            return {"success": False, "error": f"Protocol {protocol_id} not found"}
        
        # Get all steps
        steps = protocol.get_step_in_order()
        total_steps = len(steps)
        
        if total_steps == 0:
            send_progress("error", "No steps found in protocol", 0)
            return {"success": False, "error": "No steps found in protocol"}
        
        send_progress("processing", f"Found {total_steps} steps to analyze", 10)
        
        # Initialize analyzer
        analyzer = ProtocolAnalyzer(use_anthropic=use_anthropic)
        analyzer_type = 'mcp_claude' if use_anthropic else 'standard_nlp'
        
        results = []
        
        for i, step in enumerate(steps):
            current_progress = 10 + (i * 80 // total_steps)
            send_progress("processing", f"Analyzing step {i+1}/{total_steps}: {step.step_name}", current_progress)
            
            try:
                # Check cache first
                cached_result = ProtocolStepSuggestionCache.get_cached_suggestions(step.id, analyzer_type)
                
                if cached_result:
                    result = cached_result
                    result['cached'] = True
                else:
                    # Perform analysis
                    result = analyzer.analyze_protocol_step(step.id)
                    result['cached'] = False
                    
                    # Cache the results
                    if result.get('success'):
                        ProtocolStepSuggestionCache.cache_suggestions(step.id, analyzer_type, result)
                
                results.append({
                    'step_id': step.id,
                    'step_name': step.step_name,
                    'result': result
                })
                
            except Exception as e:
                results.append({
                    'step_id': step.id,
                    'step_name': step.step_name,
                    'result': {"success": False, "error": str(e)}
                })
        
        send_progress("completed", f"Analysis completed for all {total_steps} steps", 100, {
            "results": results
        })
        
        return {
            "success": True,
            "protocol_id": protocol_id,
            "total_steps": total_steps,
            "results": results
        }
        
    except Exception as e:
        error_msg = f"Full protocol analysis failed: {str(e)}"
        send_progress("error", error_msg, 0)
        return {"success": False, "error": error_msg}


def generate_pooled_sample_metadata(instrument_job_id):
    """
    Generate pooled sample metadata columns based on pool configuration.
    
    This function creates/updates the 'characteristics[pooled sample]' metadata column
    with appropriate modifiers for pooled samples according to SDRF standards.
    
    Args:
        instrument_job_id (int): The ID of the instrument job
    """
    try:
        instrument_job = InstrumentJob.objects.get(id=instrument_job_id)
        pools = SamplePool.objects.filter(instrument_job=instrument_job)
        
        # Check if pooled sample column already exists for this instrument
        pooled_column = MetadataColumn.objects.filter(
            instrument=instrument_job.instrument,
            name="Pooled sample",
            type="Characteristics"
        ).first()
        
        if not pooled_column:
            # Create the pooled sample metadata column
            pooled_column = MetadataColumn.objects.create(
                instrument=instrument_job.instrument,
                name="Pooled sample",
                type="Characteristics",
                value="not pooled",  # Default value for non-pooled samples
                mandatory=False
            )
        
        # Generate modifiers for pooled samples
        modifiers = []
        
        for pool in pools:
            # Only create modifiers for pooled-only samples
            # Pooled+independent samples will use the default "not pooled" value
            if pool.pooled_only_samples:
                modifier = {
                    "samples": ",".join([str(i) for i in pool.pooled_only_samples]),
                    "value": pool.sdrf_value
                }
                modifiers.append(modifier)
        
        # Update the metadata column with modifiers
        pooled_column.modifiers = json.dumps(modifiers) if modifiers else None
        pooled_column.save()
        
        # Log the update
        print(f"Updated pooled sample metadata for instrument job {instrument_job_id}")
        print(f"Created {len(modifiers)} modifiers for {len(pools)} pools")
        
        return {
            "success": True,
            "column_id": pooled_column.id,
            "modifiers_count": len(modifiers),
            "pools_count": len(pools)
        }
        
    except InstrumentJob.DoesNotExist:
        error_msg = f"Instrument job {instrument_job_id} not found"
        print(f"Error: {error_msg}")
        return {"success": False, "error": error_msg}
    
    except Exception as e:
        error_msg = f"Failed to generate pooled sample metadata: {str(e)}"
        print(f"Error: {error_msg}")
        return {"success": False, "error": error_msg}


def get_sample_pooling_status(instrument_job_id, sample_index):
    """
    Get the pooling status of a specific sample.
    
    Args:
        instrument_job_id (int): The ID of the instrument job
        sample_index (int): The sample index (1-based)
    
    Returns:
        dict: Status information for the sample
    """
    try:
        instrument_job = InstrumentJob.objects.get(id=instrument_job_id)
        pools = SamplePool.objects.filter(instrument_job=instrument_job)
        
        sample_pools = []
        is_pooled_only = False
        is_pooled_and_independent = False
        
        for pool in pools:
            if sample_index in pool.pooled_only_samples:
                sample_pools.append({
                    'pool_name': pool.pool_name,
                    'status': 'pooled_only',
                    'sdrf_value': pool.sdrf_value
                })
                is_pooled_only = True
            elif sample_index in pool.pooled_and_independent_samples:
                sample_pools.append({
                    'pool_name': pool.pool_name,
                    'status': 'pooled_and_independent',
                    'sdrf_value': 'not pooled'  # Independent samples get "not pooled"
                })
                is_pooled_and_independent = True
        
        # Determine overall status
        if is_pooled_only and not is_pooled_and_independent:
            status = "Pooled Only"
            sdrf_value = sample_pools[0]['sdrf_value'] if sample_pools else "not pooled"
        elif is_pooled_and_independent and not is_pooled_only:
            status = "Mixed"
            sdrf_value = "not pooled"
        elif is_pooled_only and is_pooled_and_independent:
            status = "Mixed"
            sdrf_value = "not pooled"
        else:
            status = "Independent"
            sdrf_value = "not pooled"
        
        return {
            "sample_index": sample_index,
            "status": status,
            "sdrf_value": sdrf_value,
            "pools": sample_pools
        }
        
    except InstrumentJob.DoesNotExist:
        return {
            "sample_index": sample_index,
            "status": "Error",
            "sdrf_value": "not pooled",
            "pools": [],
            "error": f"Instrument job {instrument_job_id} not found"
        }
    
    except Exception as e:
        return {
            "sample_index": sample_index,
            "status": "Error",
            "sdrf_value": "not pooled",
            "pools": [],
            "error": str(e)
        }


def validate_sample_pools(instrument_job_id):
    """
    Validate all sample pools for an instrument job.
    
    Args:
        instrument_job_id (int): The ID of the instrument job
    
    Returns:
        dict: Validation results
    """
    try:
        instrument_job = InstrumentJob.objects.get(id=instrument_job_id)
        pools = SamplePool.objects.filter(instrument_job=instrument_job)
        
        validation_results = {
            "valid": True,
            "errors": [],
            "warnings": [],
            "pools_validated": len(pools)
        }
        
        for pool in pools:
            # Validate sample indices are within range
            all_samples = pool.all_pooled_samples
            max_sample = instrument_job.sample_number
            
            invalid_samples = [s for s in all_samples if s < 1 or s > max_sample]
            if invalid_samples:
                validation_results["valid"] = False
                validation_results["errors"].append(
                    f"Pool '{pool.pool_name}' contains invalid sample indices: {invalid_samples}. "
                    f"Valid range is 1-{max_sample}"
                )
            
            # Check for empty pools
            if not all_samples:
                validation_results["warnings"].append(
                    f"Pool '{pool.pool_name}' is empty"
                )
        
        # Check for overlapping pooled-only samples across pools
        pooled_only_samples = {}
        for pool in pools:
            for sample in pool.pooled_only_samples:
                if sample in pooled_only_samples:
                    validation_results["valid"] = False
                    validation_results["errors"].append(
                        f"Sample {sample} is marked as 'pooled only' in both "
                        f"'{pooled_only_samples[sample]}' and '{pool.pool_name}' pools"
                    )
                else:
                    pooled_only_samples[sample] = pool.pool_name
        
        return validation_results
        
    except InstrumentJob.DoesNotExist:
        return {
            "valid": False,
            "errors": [f"Instrument job {instrument_job_id} not found"],
            "warnings": [],
            "pools_validated": 0
        }
    
    except Exception as e:
        return {
            "valid": False,
            "errors": [f"Validation failed: {str(e)}"],
            "warnings": [],
            "pools_validated": 0
        }

