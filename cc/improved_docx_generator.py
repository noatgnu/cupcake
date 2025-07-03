"""
Improved DOCX Generator for CUPCAKE LIMS

This module provides an enhanced version of the create_docx function with:
- Better error handling and logging
- Improved document formatting and styling
- Support for more annotation types
- Template-based generation
- Progress tracking
- Better media handling
"""
import base64
import datetime
import io
import json
import logging
import os
import re
import threading
import uuid
from datetime import time, timedelta
from typing import Dict, List, Optional, Any

import docx
import ffmpeg
import webvtt
from PIL import Image
from asgiref.sync import async_to_sync
from bs4 import BeautifulSoup
from channels.layers import get_channel_layer
from django.conf import settings
from django.core.signing import TimestampSigner
from django.utils import timezone
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from django_rq import job

from cc.models import (
    ProtocolModel, ProtocolStep, ProtocolSection, Session, 
    Annotation, MetadataColumn, Tag, Project
)

# Configure logging
logger = logging.getLogger(__name__)


class DocxGenerationError(Exception):
    """Custom exception for DOCX generation errors"""
    pass


class EnhancedDocxGenerator:
    """
    Enhanced DOCX generator with better formatting, error handling, and features
    """
    
    def __init__(self, protocol: ProtocolModel, session: Optional[Session] = None, 
                 user_id: Optional[int] = None, instance_id: Optional[str] = None):
        self.protocol = protocol
        self.session = session
        self.user_id = user_id
        self.instance_id = instance_id
        self.doc = Document()
        self.temp_files = []
        
        # Statistics for progress tracking
        self.stats = {
            'total_steps': 0,
            'processed_steps': 0,
            'total_annotations': 0,
            'processed_annotations': 0,
            'images_added': 0,
            'tables_added': 0,
            'errors': []
        }
        
        # Initialize document styles
        self._setup_document_styles()
    
    def _setup_document_styles(self):
        """Setup custom styles for the document"""
        try:
            styles = self.doc.styles
            
            # Create custom heading styles
            if 'Protocol Title' not in [s.name for s in styles]:
                title_style = styles.add_style('Protocol Title', WD_STYLE_TYPE.PARAGRAPH)
                title_style.font.name = 'Arial'
                title_style.font.size = Pt(16)
                title_style.font.bold = True
                title_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_style.paragraph_format.space_after = Pt(12)
            
            # Create step header style
            if 'Step Header' not in [s.name for s in styles]:
                step_style = styles.add_style('Step Header', WD_STYLE_TYPE.PARAGRAPH)
                step_style.font.name = 'Arial'
                step_style.font.size = Pt(12)
                step_style.font.bold = True
                step_style.font.color.rgb = RGBColor(51, 51, 51)  # Dark gray
                step_style.paragraph_format.space_before = Pt(6)
                step_style.paragraph_format.space_after = Pt(3)
            
            # Create annotation style
            if 'Annotation' not in [s.name for s in styles]:
                annotation_style = styles.add_style('Annotation', WD_STYLE_TYPE.PARAGRAPH)
                annotation_style.font.name = 'Calibri'
                annotation_style.font.size = Pt(10)
                annotation_style.font.italic = True
                annotation_style.paragraph_format.left_indent = Inches(0.5)
                annotation_style.paragraph_format.space_after = Pt(3)
            
            # Create metadata style
            if 'Metadata' not in [s.name for s in styles]:
                metadata_style = styles.add_style('Metadata', WD_STYLE_TYPE.PARAGRAPH)
                metadata_style.font.name = 'Consolas'
                metadata_style.font.size = Pt(9)
                metadata_style.font.color.rgb = RGBColor(102, 102, 102)  # Gray
                metadata_style.paragraph_format.left_indent = Inches(0.25)
                
        except Exception as e:
            logger.warning(f"Failed to setup custom styles: {e}")
    
    def generate_document(self) -> str:
        """
        Generate the complete DOCX document
        
        Returns:
            str: Path to the generated DOCX file
        """
        try:
            logger.info(f"Starting DOCX generation for protocol {self.protocol.id}")
            
            # Add document header and metadata
            self._add_document_header()
            self._add_protocol_metadata()
            
            # Count total work for progress tracking
            self._calculate_work_stats()
            
            # Add protocol content
            self._add_protocol_description()
            self._add_protocol_sections()
            
            # Add appendices
            self._add_reagents_appendix()
            self._add_tags_appendix()
            self._add_session_summary()
            
            # Save document
            filename = self._generate_filename()
            filepath = self._save_document(filename)
            
            # Notify user and schedule cleanup
            self._notify_completion(filename)
            self._schedule_cleanup(filepath)
            
            logger.info(f"Successfully generated DOCX: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {e}")
            self._notify_error(str(e))
            raise DocxGenerationError(f"DOCX generation failed: {e}")
    
    def get_statistics(self) -> Dict[str, Any]:
        """
        Get generation statistics
        
        Returns:
            Dict containing generation statistics
        """
        return {
            'total_steps': self.stats['total_steps'],
            'processed_steps': self.stats['processed_steps'],
            'total_annotations': self.stats['total_annotations'],
            'processed_annotations': self.stats['processed_annotations'],
            'images_added': self.stats['images_added'],
            'tables_added': self.stats['tables_added'],
            'errors_count': len(self.stats['errors']),
            'errors': self.stats['errors'][:5],  # First 5 errors only
            'completion_rate': (
                self.stats['processed_steps'] / max(self.stats['total_steps'], 1) * 100
            ) if self.stats['total_steps'] > 0 else 100
        }
    
    def _add_document_header(self):
        """Add document header with title and timestamp"""
        try:
            # Protocol title
            title_para = self.doc.add_paragraph()
            title_para.style = 'Protocol Title'
            title_run = title_para.add_run(self._clean_html(getattr(self.protocol, 'protocol_title', getattr(self.protocol, 'protocol_name', 'Untitled Protocol'))))
            
            # Add protocol ID if available
            if self.protocol.protocol_id:
                subtitle_para = self.doc.add_paragraph()
                subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_run = subtitle_para.add_run(f"Protocol ID: {self.protocol.protocol_id}")
                subtitle_run.font.size = Pt(11)
                subtitle_run.font.italic = True
                
            # Add generation timestamp
            timestamp_para = self.doc.add_paragraph()
            timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            timestamp_run = timestamp_para.add_run(
                f"Generated: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}"
            )
            timestamp_run.font.size = Pt(9)
            timestamp_run.font.color.rgb = RGBColor(102, 102, 102)
            
            # Add page break
            self.doc.add_page_break()
            
        except Exception as e:
            logger.error(f"Failed to add document header: {e}")
            self.stats['errors'].append(f"Header generation error: {e}")
    
    def _add_protocol_metadata(self):
        """Add protocol metadata section"""
        try:
            # Metadata section header
            self.doc.add_heading('Protocol Information', level=1)
            
            metadata_items = [
                ('Protocol ID', getattr(self.protocol, 'protocol_id', None)),
                ('Created', self.protocol.protocol_created_on.strftime('%B %d, %Y') if self.protocol.protocol_created_on else 'Unknown'),
                ('Version URI', getattr(self.protocol, 'protocol_version_uri', None)),
                ('DOI', getattr(self.protocol, 'protocol_doi', None)),
                ('URL', getattr(self.protocol, 'protocol_url', None)),
                ('User', str(self.protocol.user) if self.protocol.user else None),
                ('Enabled', 'Yes' if getattr(self.protocol, 'enabled', False) else 'No')
            ]
            
            # Create metadata table
            table = self.doc.add_table(rows=0, cols=2)
            table.style = 'Light Grid Accent 1'
            
            for label, value in metadata_items:
                if value:  # Only add if value exists
                    row_cells = table.add_row().cells
                    row_cells[0].text = label
                    row_cells[1].text = str(value)
                    
                    # Style the cells
                    row_cells[0].paragraphs[0].runs[0].bold = True
            
            # Add tags if available
            tags = Tag.objects.filter(protocoltag__protocol=self.protocol)
            if tags.exists():
                self.doc.add_paragraph()
                tags_para = self.doc.add_paragraph()
                tags_para.add_run('Tags: ').bold = True
                tag_names = ', '.join([getattr(tag, 'tag', str(tag)) for tag in tags])
                tags_para.add_run(tag_names)
            
            self.doc.add_page_break()
            
        except Exception as e:
            logger.error(f"Failed to add protocol metadata: {e}")
            self.stats['errors'].append(f"Metadata generation error: {e}")
    
    def _add_protocol_description(self):
        """Add protocol description"""
        try:
            if hasattr(self.protocol, 'protocol_description') and self.protocol.protocol_description:
                self.doc.add_heading('Description', level=1)
                self._convert_html_to_docx(self.protocol.protocol_description)
                self.doc.add_paragraph()
                
        except Exception as e:
            logger.error(f"Failed to add protocol description: {e}")
            self.stats['errors'].append(f"Description generation error: {e}")
    
    def _add_protocol_sections(self):
        """Add all protocol sections and steps"""
        try:
            sections = self.protocol.get_section_in_order()
            
            if not sections:
                # If no sections, add steps directly
                steps = self.protocol.get_step_in_order()
                if steps:
                    self.doc.add_heading('Protocol Steps', level=1)
                    for i, step in enumerate(steps, 1):
                        self._add_protocol_step(step, i)
                        self.stats['processed_steps'] += 1
            else:
                # Process sections
                for section in sections:
                    self._add_protocol_section(section)
            
        except Exception as e:
            logger.error(f"Failed to add protocol sections: {e}")
            self.stats['errors'].append(f"Sections generation error: {e}")
    
    def _add_protocol_section(self, section: ProtocolSection):
        """Add a single protocol section"""
        try:
            # Section header
            self.doc.add_heading(self._clean_html(section.section_description), level=1)
            
            # Section steps
            steps = section.get_step_in_order()
            for i, step in enumerate(steps, 1):
                self._add_protocol_step(step, i, section_context=section)
                self.stats['processed_steps'] += 1
            
            # Add page break after section
            self.doc.add_page_break()
            
        except Exception as e:
            logger.error(f"Failed to add section {section.id}: {e}")
            self.stats['errors'].append(f"Section {section.id} error: {e}")
    
    def _add_protocol_step(self, step: ProtocolStep, step_number: int, section_context: Optional[ProtocolSection] = None):
        """Add a single protocol step with all its content"""
        try:
            # Step header with divider
            divider_para = self.doc.add_paragraph('─' * 80)
            divider_para.style = 'Annotation'
            
            step_header = self.doc.add_paragraph()
            step_header.style = 'Step Header'
            # Create step title from description or use step number
            step_title = f"Step {step_number}"
            if hasattr(step, 'step_description') and step.step_description:
                # Use first 50 chars of description as title
                desc_short = self._clean_html(step.step_description)[:50]
                if len(desc_short) == 50:
                    desc_short += '...'
                step_title = f"Step {step_number}: {desc_short}"
            step_header.add_run(step_title)
            
            # Process step description with reagent substitutions
            processed_description = self._process_step_description(step)
            self._convert_html_to_docx(processed_description)
            
            # Add step duration if available
            if hasattr(step, 'step_duration') and step.step_duration:
                duration_para = self.doc.add_paragraph()
                duration_para.style = 'Metadata'
                duration_para.add_run(f"Duration: {self._convert_seconds_to_time(step.step_duration)}")
            
            # Add step metadata
            self._add_step_metadata(step)
            
            # Add step annotations
            self._add_step_annotations(step)
            
            # Add expected results if available
            if step.step_expected_results:
                self.doc.add_paragraph()
                results_para = self.doc.add_paragraph()
                results_para.add_run('Expected Results: ').bold = True
                self._convert_html_to_docx(step.step_expected_results)
            
            # Add safety information if available
            if step.step_safety_information:
                self.doc.add_paragraph()
                safety_para = self.doc.add_paragraph()
                safety_run = safety_para.add_run('⚠️ Safety Information: ')
                safety_run.bold = True
                safety_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                self._convert_html_to_docx(step.step_safety_information)
            
            self.doc.add_paragraph()  # Add spacing
            
        except Exception as e:
            logger.error(f"Failed to add step {step.id}: {e}")
            self.stats['errors'].append(f"Step {step.id} error: {e}")
    
    def _add_step_metadata(self, step: ProtocolStep):
        """Add step metadata columns"""
        try:
            metadata_columns = MetadataColumn.objects.filter(protocol_step=step)
            if metadata_columns.exists():
                metadata_para = self.doc.add_paragraph()
                metadata_para.style = 'Metadata'
                metadata_para.add_run('Metadata: ')
                
                for meta in metadata_columns:
                    metadata_para.add_run(f"{meta.name}: {meta.value}; ")
                    
        except Exception as e:
            logger.error(f"Failed to add step metadata for step {step.id}: {e}")
    
    def _add_step_annotations(self, step: ProtocolStep):
        """Add all annotations for a step"""
        try:
            # Get annotations for this step
            annotations = step.annotations.all()
            if self.session:
                annotations = annotations.filter(session=self.session)
            
            if not annotations.exists():
                return
            
            # Annotations header
            annotations_header = self.doc.add_paragraph()
            annotations_header.add_run('Annotations:').bold = True
            
            for annotation in annotations:
                self._add_single_annotation(annotation)
                self.stats['processed_annotations'] += 1
                
        except Exception as e:
            logger.error(f"Failed to add annotations for step {step.id}: {e}")
            self.stats['errors'].append(f"Annotations for step {step.id} error: {e}")
    
    def _add_single_annotation(self, annotation: Annotation):
        """Add a single annotation with proper formatting"""
        try:
            # Annotation container
            annotation_para = self.doc.add_paragraph()
            annotation_para.style = 'Annotation'
            
            # Annotation type badge
            type_run = annotation_para.add_run(f"[{annotation.annotation_type.upper()}] ")
            type_run.bold = True
            type_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue
            
            # Annotation name if available
            if annotation.annotation_name:
                annotation_para.add_run(f"{annotation.annotation_name}: ")
            
            # Handle different annotation types
            if annotation.annotation_type == "text":
                self._add_text_annotation(annotation)
            elif annotation.annotation_type == "image":
                self._add_image_annotation(annotation)
            elif annotation.annotation_type == "sketch":
                self._add_sketch_annotation(annotation)
            elif annotation.annotation_type == "table":
                self._add_table_annotation(annotation)
            elif annotation.annotation_type == "checklist":
                self._add_checklist_annotation(annotation)
            elif annotation.annotation_type == "alignment":
                self._add_alignment_annotation(annotation)
            elif annotation.annotation_type in ["audio", "video"]:
                self._add_media_annotation(annotation)
            else:
                # Generic annotation
                if annotation.annotation:
                    self.doc.add_paragraph(annotation.annotation)
            
            # Add transcription if available
            if annotation.transcribed and annotation.transcription:
                self._add_transcription(annotation.transcription)
            
            # Add translation if available
            if annotation.translation:
                self._add_translation(annotation.translation)
            
            # Add summary if available
            if annotation.summary:
                summary_para = self.doc.add_paragraph()
                summary_para.style = 'Annotation'
                summary_para.add_run('Summary: ').bold = True
                summary_para.add_run(annotation.summary)
            
        except Exception as e:
            logger.error(f"Failed to add annotation {annotation.id}: {e}")
            self.stats['errors'].append(f"Annotation {annotation.id} error: {e}")
    
    def _add_text_annotation(self, annotation: Annotation):
        """Add text annotation"""
        if annotation.annotation:
            self._convert_html_to_docx(annotation.annotation)
    
    def _add_image_annotation(self, annotation: Annotation):
        """Add image annotation with proper sizing"""
        try:
            if not annotation.file or not annotation.file.path:
                return
            
            if not os.path.exists(annotation.file.path):
                logger.warning(f"Image file not found: {annotation.file.path}")
                return
            
            # Get image dimensions
            try:
                image_info = ffmpeg.probe(annotation.file.path, show_entries="stream=width,height")
                width = int(image_info["streams"][0]["width"])
                height = int(image_info["streams"][0]["height"])
                
                # Calculate display size (max 6 inches width, 8 inches height)
                max_width = 6.0
                max_height = 8.0
                
                if width > height:
                    # Landscape orientation
                    display_width = min(max_width, width / 96)  # Convert pixels to inches
                    self.doc.add_picture(annotation.file.path, width=Inches(display_width))
                else:
                    # Portrait orientation
                    display_height = min(max_height, height / 96)
                    self.doc.add_picture(annotation.file.path, height=Inches(display_height))
                
                self.stats['images_added'] += 1
                
            except Exception as e:
                logger.warning(f"Failed to process image dimensions for {annotation.file.path}: {e}")
                # Fallback: add with default size
                self.doc.add_picture(annotation.file.path, width=Inches(4))
                self.stats['images_added'] += 1
                
        except Exception as e:
            logger.error(f"Failed to add image annotation {annotation.id}: {e}")
    
    def _add_sketch_annotation(self, annotation: Annotation):
        """Add sketch annotation from JSON data"""
        try:
            if not annotation.file:
                return
            
            with open(annotation.file.path, 'r') as f:
                sketch_data = json.load(f)
            
            if "png" in sketch_data:
                data_parts = sketch_data["png"].split('base64,')
                if len(data_parts) > 1:
                    image_bytes = base64.b64decode(data_parts[1])
                    image_file = io.BytesIO(image_bytes)
                    
                    # Calculate size based on sketch dimensions
                    pixel_width = sketch_data.get("width", 400)
                    pixel_height = sketch_data.get("height", 300)
                    aspect_ratio = pixel_width / pixel_height
                    
                    # Scale to reasonable size
                    if aspect_ratio > 1:
                        width = min(5.0, pixel_width / 96)
                        self.doc.add_picture(image_file, width=Inches(width))
                    else:
                        height = min(6.0, pixel_height / 96)
                        self.doc.add_picture(image_file, height=Inches(height))
                    
                    self.stats['images_added'] += 1
                    
        except Exception as e:
            logger.error(f"Failed to add sketch annotation {annotation.id}: {e}")
    
    def _add_table_annotation(self, annotation: Annotation):
        """Add table annotation with proper formatting"""
        try:
            table_data = json.loads(annotation.annotation)
            if not table_data:
                return
            
            # Add table name
            if "name" in table_data:
                table_name_para = self.doc.add_paragraph()
                table_name_para.add_run(table_data["name"]).bold = True
            
            # Create table
            rows = table_data.get("nRow", 0)
            cols = table_data.get("nCol", 0)
            
            if rows > 0 and cols > 0:
                table = self.doc.add_table(rows=rows, cols=cols)
                table.style = 'Light Grid Accent 1'
                
                # Fill table content
                content = table_data.get("content", [])
                tracking_map = table_data.get("trackingMap", {})
                
                for row_idx, row_data in enumerate(content):
                    if row_idx >= rows:
                        break
                    
                    row_cells = table.rows[row_idx].cells
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx >= cols:
                            break
                        
                        row_cells[col_idx].text = str(cell_data)
                        
                        # Apply highlighting if tracked
                        cell_key = f"{row_idx},{col_idx}"
                        if cell_key in tracking_map and tracking_map[cell_key]:
                            # Highlight cell
                            shading_elm = parse_xml(r'<w:shd {} w:fill="B3D9FF"/>'.format(nsdecls('w')))
                            row_cells[col_idx]._tc.get_or_add_tcPr().append(shading_elm)
                
                self.stats['tables_added'] += 1
                
        except Exception as e:
            logger.error(f"Failed to add table annotation {annotation.id}: {e}")
    
    def _add_checklist_annotation(self, annotation: Annotation):
        """Add checklist annotation"""
        try:
            checklist_data = json.loads(annotation.annotation)
            if not checklist_data:
                return
            
            # Add checklist name
            if "name" in checklist_data:
                checklist_name_para = self.doc.add_paragraph()
                checklist_name_para.add_run(checklist_data["name"]).bold = True
            
            # Add checklist items
            checklist_items = checklist_data.get("checkList", [])
            for i, item in enumerate(checklist_items, 1):
                item_para = self.doc.add_paragraph()
                
                # Add checkbox symbol
                checked = item.get("checked", False)
                checkbox = "☑️" if checked else "☐"
                item_para.add_run(f"{checkbox} ")
                
                # Add item content
                content = item.get("content", f"Item {i}")
                item_para.add_run(content)
                
        except Exception as e:
            logger.error(f"Failed to add checklist annotation {annotation.id}: {e}")
    
    def _add_alignment_annotation(self, annotation: Annotation):
        """Add alignment annotation with images"""
        try:
            alignment_data = json.loads(annotation.annotation)
            if not alignment_data:
                return
            
            # Add main alignment image
            if "dataURL" in alignment_data:
                data_parts = alignment_data["dataURL"].split('base64,')
                if len(data_parts) > 1:
                    image_bytes = base64.b64decode(data_parts[1])
                    image_file = io.BytesIO(image_bytes)
                    self.doc.add_picture(image_file, width=Inches(5))
                    self.stats['images_added'] += 1
            
            # Add extracted segments
            segments = alignment_data.get("extractedSegments", [])
            for segment in segments:
                if "dataURL" in segment:
                    # Add segment info
                    segment_para = self.doc.add_paragraph()
                    segment_para.add_run(f"Segment: {segment.get('start', 'N/A')} - {segment.get('end', 'N/A')}")
                    
                    # Add segment image
                    data_parts = segment["dataURL"].split('base64,')
                    if len(data_parts) > 1:
                        image_bytes = base64.b64decode(data_parts[1])
                        image_file = io.BytesIO(image_bytes)
                        self.doc.add_picture(image_file, width=Inches(4))
                        self.stats['images_added'] += 1
                        
        except Exception as e:
            logger.error(f"Failed to add alignment annotation {annotation.id}: {e}")
    
    def _add_media_annotation(self, annotation: Annotation):
        """Add media (audio/video) annotation info"""
        try:
            media_para = self.doc.add_paragraph()
            media_para.style = 'Annotation'
            
            if annotation.file:
                file_name = os.path.basename(annotation.file.name)
                media_para.add_run(f"Media file: {file_name}")
                
                # Add file size if available
                try:
                    file_size = os.path.getsize(annotation.file.path)
                    size_mb = file_size / (1024 * 1024)
                    media_para.add_run(f" ({size_mb:.1f} MB)")
                except:
                    pass
            else:
                media_para.add_run("Media file not available")
                
        except Exception as e:
            logger.error(f"Failed to add media annotation {annotation.id}: {e}")
    
    def _add_transcription(self, transcription: str):
        """Add transcription content"""
        try:
            transcription_para = self.doc.add_paragraph()
            transcription_para.style = 'Annotation'
            transcription_para.add_run('🎙️ Transcription: ').bold = True
            
            if transcription.startswith("WEBVTT"):
                # Parse WebVTT format
                try:
                    for caption in webvtt.read_buffer(io.StringIO(transcription)):
                        time_para = self.doc.add_paragraph()
                        time_para.style = 'Annotation'
                        time_para.add_run(f"({caption.start} - {caption.end}) {caption.text}")
                except Exception as e:
                    # Fallback to plain text
                    transcription_para.add_run(transcription)
            else:
                transcription_para.add_run(transcription)
                
        except Exception as e:
            logger.error(f"Failed to add transcription: {e}")
    
    def _add_translation(self, translation: str):
        """Add translation content"""
        try:
            translation_para = self.doc.add_paragraph()
            translation_para.style = 'Annotation'
            translation_para.add_run('🌐 Translation: ').bold = True
            
            if translation.startswith("WEBVTT"):
                # Parse WebVTT format
                try:
                    for caption in webvtt.read_buffer(io.StringIO(translation)):
                        time_para = self.doc.add_paragraph()
                        time_para.style = 'Annotation'
                        time_para.add_run(f"({caption.start} - {caption.end}) {caption.text}")
                except Exception as e:
                    # Fallback to plain text
                    translation_para.add_run(translation)
            else:
                translation_para.add_run(translation)
                
        except Exception as e:
            logger.error(f"Failed to add translation: {e}")
    
    def _add_reagents_appendix(self):
        """Add appendix with reagents information"""
        try:
            # Get all reagents used in protocol
            from cc.models import ProtocolReagent, StepReagent
            
            protocol_reagents = ProtocolReagent.objects.filter(protocol=self.protocol)
            step_reagents = StepReagent.objects.filter(step__protocol=self.protocol)
            
            if not protocol_reagents.exists() and not step_reagents.exists():
                return
            
            self.doc.add_page_break()
            self.doc.add_heading('Reagents and Materials', level=1)
            
            if protocol_reagents.exists():
                self.doc.add_heading('Protocol Reagents', level=2)
                reagents_table = self.doc.add_table(rows=1, cols=3)
                reagents_table.style = 'Light Grid Accent 1'
                
                # Header row
                header_cells = reagents_table.rows[0].cells
                header_cells[0].text = 'Reagent'
                header_cells[1].text = 'Quantity'
                header_cells[2].text = 'Unit'
                header_cells[3].text = 'Notes'
                
                for header_cell in header_cells:
                    header_cell.paragraphs[0].runs[0].bold = True
                
                # Data rows
                for reagent in protocol_reagents:
                    row_cells = reagents_table.add_row().cells
                    reagent_name = getattr(reagent.reagent, 'name', None) or getattr(reagent.reagent, 'reagent_name', 'Unknown')
                    row_cells[0].text = reagent_name
                    row_cells[1].text = str(getattr(reagent, 'quantity', '') or '')
                    row_cells[2].text = reagent.reagent.unit or ''
            
            if step_reagents.exists():
                self.doc.add_heading('Step-Specific Reagents', level=2)
                step_reagents_table = self.doc.add_table(rows=1, cols=5)
                step_reagents_table.style = 'Light Grid Accent 1'
                
                # Header row
                header_cells = step_reagents_table.rows[0].cells
                header_cells[0].text = 'Step'
                header_cells[1].text = 'Reagent'
                header_cells[2].text = 'Quantity'
                header_cells[3].text = 'Unit'
                header_cells[4].text = 'Scalable'
                
                for header_cell in header_cells:
                    header_cell.paragraphs[0].runs[0].bold = True
                
                # Data rows
                for step_reagent in step_reagents:
                    row_cells = step_reagents_table.add_row().cells
                    step_desc = getattr(step_reagent.step, 'step_description', '') or ''
                    step_title = step_desc[:30] + '...' if len(step_desc) > 30 else step_desc
                    row_cells[0].text = step_title or f"Step {getattr(step_reagent.step, 'step_id', 'Unknown')}"
                    reagent_name = getattr(step_reagent.reagent, 'name', None) or getattr(step_reagent.reagent, 'reagent_name', 'Unknown')
                    row_cells[1].text = reagent_name
                    row_cells[2].text = str(getattr(step_reagent, 'quantity', '') or '')
                    row_cells[3].text = step_reagent.reagent.unit or ''
                    row_cells[4].text = 'Yes' if step_reagent.scalable else 'No'
                    
        except Exception as e:
            logger.error(f"Failed to add reagents appendix: {e}")
            self.stats['errors'].append(f"Reagents appendix error: {e}")
    
    def _add_tags_appendix(self):
        """Add appendix with tags information"""
        try:
            from cc.models import ProtocolTag, StepTag
            
            protocol_tags = ProtocolTag.objects.filter(protocol=self.protocol)
            step_tags = StepTag.objects.filter(step__protocol=self.protocol)
            
            if not protocol_tags.exists() and not step_tags.exists():
                return
            
            self.doc.add_heading('Tags and Categories', level=2)
            
            if protocol_tags.exists():
                tags_para = self.doc.add_paragraph()
                tags_para.add_run('Protocol Tags: ').bold = True
                tag_names = ', '.join([pt.tag.tag for pt in protocol_tags])
                tags_para.add_run(tag_names)
            
            if step_tags.exists():
                self.doc.add_paragraph()
                self.doc.add_heading('Step Tags', level=3)
                
                for step_tag in step_tags:
                    step_tag_para = self.doc.add_paragraph()
                    step_desc = getattr(step_tag.step, 'step_description', '') or ''
                    step_title = step_desc[:20] + '...' if len(step_desc) > 20 else step_desc
                    step_title = step_title or f"Step {getattr(step_tag.step, 'step_id', 'Unknown')}"
                    step_tag_para.add_run(f"{step_title}: ").bold = True
                    step_tag_para.add_run(getattr(step_tag.tag, 'tag', str(step_tag.tag)))
                    
        except Exception as e:
            logger.error(f"Failed to add tags appendix: {e}")
    
    def _add_session_summary(self):
        """Add session summary if session is provided"""
        try:
            if not self.session:
                return
            
            self.doc.add_heading('Session Information', level=2)
            
            session_info = [
                ('Session ID', self.session.unique_id),
                ('Session Name', self.session.name),
                ('Started', self.session.started_at.strftime('%B %d, %Y at %I:%M %p') if self.session.started_at else 'Not started'),
                ('Ended', self.session.ended_at.strftime('%B %d, %Y at %I:%M %p') if self.session.ended_at else 'Not ended'),
                ('Status', 'Enabled' if self.session.enabled else 'Disabled'),
                ('Processing', 'Yes' if self.session.processing else 'No')
            ]
            
            # Create session info table
            session_table = self.doc.add_table(rows=0, cols=2)
            session_table.style = 'Light Grid Accent 1'
            
            for label, value in session_info:
                if value:
                    row_cells = session_table.add_row().cells
                    row_cells[0].text = label
                    row_cells[1].text = str(value)
                    row_cells[0].paragraphs[0].runs[0].bold = True
            
            # Add session statistics
            total_annotations = Annotation.objects.filter(session=self.session).count()
            if total_annotations > 0:
                self.doc.add_paragraph()
                stats_para = self.doc.add_paragraph()
                stats_para.add_run(f'Total Annotations: {total_annotations}')
                
        except Exception as e:
            logger.error(f"Failed to add session summary: {e}")
    
    def _process_step_description(self, step: ProtocolStep) -> str:
        """Process step description with reagent substitutions"""
        try:
            description = step.step_description or ""
            
            # Process reagent substitutions
            from cc.models import StepReagent
            step_reagents = StepReagent.objects.filter(step=step)
            
            for reagent in step_reagents:
                replacements = {
                    f"{reagent.id}.name": getattr(reagent.reagent, 'name', None) or getattr(reagent.reagent, 'reagent_name', 'Unknown'),
                    f"{reagent.id}.quantity": str(getattr(reagent, 'quantity', '') or ''),
                    f"{reagent.id}.unit": reagent.reagent.unit or '',
                    f"{reagent.id}.scaled_quantity": str((getattr(reagent, 'quantity', 0) or 0) * (getattr(reagent, 'scalable_factor', 1) or 1))
                }
                
                for placeholder, replacement in replacements.items():
                    description = description.replace(placeholder, replacement)
            
            return description
            
        except Exception as e:
            logger.error(f"Failed to process step description for step {step.id}: {e}")
            return step.step_description or ""
    
    def _convert_html_to_docx(self, html_content: str):
        """Convert HTML content to DOCX paragraphs with proper formatting"""
        try:
            if not html_content:
                return
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            for element in soup.find_all(['p', 'div', 'span', 'br']):
                if element.name == 'br':
                    self.doc.add_paragraph()
                    continue
                
                para = self.doc.add_paragraph()
                self._process_html_element(element, para)
            
            # If no block elements found, add as simple paragraph
            if not soup.find_all(['p', 'div']):
                para = self.doc.add_paragraph()
                self._process_html_element(soup, para)
                
        except Exception as e:
            logger.error(f"Failed to convert HTML to DOCX: {e}")
            # Fallback: add as plain text
            self.doc.add_paragraph(self._clean_html(html_content))
    
    def _process_html_element(self, element, paragraph):
        """Process individual HTML elements and add to paragraph"""
        try:
            if element.string:
                run = paragraph.add_run(element.string)
                self._apply_html_formatting(element, run)
            
            for child in element.children:
                if child.name:
                    child_run = paragraph.add_run(child.get_text())
                    self._apply_html_formatting(child, child_run)
                elif child.string:
                    paragraph.add_run(child.string)
                    
        except Exception as e:
            logger.error(f"Failed to process HTML element: {e}")
    
    def _apply_html_formatting(self, element, run):
        """Apply HTML formatting to DOCX run"""
        try:
            # Bold
            if element.name in ['b', 'strong']:
                run.bold = True
            
            # Italic
            if element.name in ['i', 'em']:
                run.italic = True
            
            # Color from style attribute
            if element.get('style'):
                style = element['style']
                if 'color:' in style:
                    color_match = re.search(r'color:\s*([^;]+)', style)
                    if color_match:
                        color_value = color_match.group(1).strip()
                        if color_value.startswith('#') and len(color_value) == 7:
                            try:
                                r = int(color_value[1:3], 16)
                                g = int(color_value[3:5], 16)
                                b = int(color_value[5:7], 16)
                                run.font.color.rgb = RGBColor(r, g, b)
                            except ValueError:
                                pass
                        elif color_value.startswith('rgb'):
                            rgb_match = re.search(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color_value)
                            if rgb_match:
                                r, g, b = map(int, rgb_match.groups())
                                run.font.color.rgb = RGBColor(r, g, b)
                                
        except Exception as e:
            logger.error(f"Failed to apply HTML formatting: {e}")
    
    def _clean_html(self, text: str) -> str:
        """Remove HTML tags from text"""
        if not text:
            return ""
        clean = re.compile('<.*?>')
        return re.sub(clean, '', text)
    
    def _convert_seconds_to_time(self, seconds: int) -> str:
        """Convert seconds to readable time format"""
        try:
            if not seconds:
                return "00:00:00"
            
            hours = seconds // 3600
            minutes = (seconds % 3600) // 60
            remaining_seconds = seconds % 60
            
            return f"{hours:02d}:{minutes:02d}:{remaining_seconds:02d}"
            
        except Exception as e:
            logger.error(f"Failed to convert seconds to time: {e}")
            return "00:00:00"
    
    def _calculate_work_stats(self):
        """Calculate total work for progress tracking"""
        try:
            self.stats['total_steps'] = self.protocol.steps.count() if hasattr(self.protocol, 'steps') else 0
            self.stats['total_annotations'] = Annotation.objects.filter(
                step__protocol=self.protocol
            ).count()
            
            if self.session:
                self.stats['total_annotations'] = Annotation.objects.filter(
                    step__protocol=self.protocol,
                    session=self.session
                ).count()
                
        except Exception as e:
            logger.error(f"Failed to calculate work stats: {e}")
    
    def _generate_filename(self) -> str:
        """Generate unique filename for the document"""
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        protocol_name = re.sub(r'[^\w\-_.]', '_', self.protocol.protocol_title)[:50]
        return f"{protocol_name}_{timestamp}_{uuid.uuid4().hex[:8]}.docx"
    
    def _save_document(self, filename: str) -> str:
        """Save the document to file system"""
        try:
            # Ensure temp directory exists
            temp_dir = os.path.join(settings.MEDIA_ROOT, "temp")
            os.makedirs(temp_dir, exist_ok=True)
            
            filepath = os.path.join(temp_dir, filename)
            self.doc.save(filepath)
            
            return filepath
            
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            raise DocxGenerationError(f"Failed to save document: {e}")
    
    def _notify_completion(self, filename: str):
        """Notify user of completion via WebSocket"""
        try:
            if not self.user_id:
                return
            
            signer = TimestampSigner()
            signed_value = signer.sign(filename)
            
            channel_layer = get_channel_layer()
            async_to_sync(channel_layer.group_send)(
                f"user_{self.user_id}",
                {
                    "type": "download_message",
                    "message": {
                        "signed_value": signed_value,
                        "user_download": True,
                        "instance_id": self.instance_id,
                        "stats": self.stats
                    },
                },
            )
            
        except Exception as e:
            logger.error(f"Failed to notify completion: {e}")
    
    def _notify_error(self, error_message: str):
        """Notify user of error via WebSocket"""
        try:
            if not self.user_id:
                return
            
            channel_layer = get_channel_layer()
            async_to_sync(channel_layer.group_send)(
                f"user_{self.user_id}",
                {
                    "type": "error_message",
                    "message": {
                        "error": error_message,
                        "instance_id": self.instance_id,
                        "stats": self.stats
                    },
                },
            )
            
        except Exception as e:
            logger.error(f"Failed to notify error: {e}")
    
    def _schedule_cleanup(self, filepath: str):
        """Schedule file cleanup after 20 minutes"""
        try:
            def cleanup():
                try:
                    if os.path.exists(filepath):
                        os.remove(filepath)
                        logger.info(f"Cleaned up temporary file: {filepath}")
                except Exception as e:
                    logger.error(f"Failed to cleanup file {filepath}: {e}")
            
            # Schedule cleanup in 20 minutes
            timer = threading.Timer(60 * 20, cleanup)
            timer.start()
            
        except Exception as e:
            logger.error(f"Failed to schedule cleanup: {e}")


@job('export', timeout='2h')
def create_enhanced_docx(protocol_id: int, session_id: str = None, user_id: int = None, instance_id: str = None):
    """
    Enhanced DOCX creation job with better error handling and features
    
    Args:
        protocol_id: ID of the protocol to export
        session_id: Optional session ID for session-specific annotations
        user_id: User ID for notifications
        instance_id: Instance ID for tracking
        
    Returns:
        str: Path to generated DOCX file
    """
    try:
        # Get protocol
        protocol = ProtocolModel.objects.get(id=protocol_id)
        
        # Get session if provided
        session = None
        if session_id:
            try:
                session = Session.objects.get(unique_id=session_id)
            except Session.DoesNotExist:
                logger.warning(f"Session {session_id} not found, proceeding without session")
        
        # Create generator and generate document
        generator = EnhancedDocxGenerator(
            protocol=protocol,
            session=session,
            user_id=user_id,
            instance_id=instance_id
        )
        
        filepath = generator.generate_document()
        
        logger.info(f"Successfully generated enhanced DOCX: {filepath}")
        logger.info(f"Generation stats: {generator.stats}")
        
        return filepath
        
    except ProtocolModel.DoesNotExist:
        error_msg = f"Protocol {protocol_id} not found"
        logger.error(error_msg)
        if user_id:
            # Notify user of error
            channel_layer = get_channel_layer()
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}",
                {
                    "type": "error_message",
                    "message": {
                        "error": error_msg,
                        "instance_id": instance_id
                    },
                },
            )
        raise DocxGenerationError(error_msg)
        
    except Exception as e:
        error_msg = f"DOCX generation failed: {str(e)}"
        logger.error(error_msg)
        if user_id:
            # Notify user of error
            channel_layer = get_channel_layer()
            async_to_sync(channel_layer.group_send)(
                f"user_{user_id}",
                {
                    "type": "error_message",
                    "message": {
                        "error": error_msg,
                        "instance_id": instance_id
                    },
                },
            )
        raise DocxGenerationError(error_msg)